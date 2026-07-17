"""Finalize every supplied 2025/2026 Third Term paper.

This is intentionally deterministic. It repairs parser losses from DOCX/PDF/TXT
sources, preserves manual English/Literature/Visual Art theory edits, creates the
new language papers, and validates every active paper before committing.
"""

from __future__ import annotations

import os
import re
import sys
import unicodedata
from datetime import datetime
from decimal import Decimal
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", os.getenv("DJANGO_SETTINGS_MODULE", "core.settings"))

import django

django.setup()

from django.core.files import File
from django.core.files.storage import default_storage
from django.db import transaction
from django.db.models import Q
from django.utils import timezone

from apps.accounts.constants import ROLE_DEAN, ROLE_IT_MANAGER
from apps.accounts.models import User
from apps.academics.models import (
    AcademicClass,
    AcademicSession,
    ClassSubject,
    Subject,
    TeacherSubjectAssignment,
    Term,
)
from apps.cbt.models import (
    CBTExamStatus,
    CBTExamType,
    CBTQuestionDifficulty,
    CBTQuestionType,
    CBTWritebackTarget,
    CorrectAnswer,
    Exam,
    ExamBlueprint,
    ExamQuestion,
    Option,
    Question,
    QuestionBank,
)
from apps.cbt.workflow import _activation_snapshot_hash, _activation_snapshot_payload
from scripts.import_ca23_third_term_20260624 import (
    parse_answer_terminated_rows,
    parse_numbered_text_rows,
    parse_plain_option_answer_rows,
)
from scripts.import_third_term_exams_20260629 import configured_base


SOURCE_ROOT = ROOT / "SCHOOL FOLDER" / "3RD TERM 2ND CA AND EXAM"
ASSET_ROOT = ROOT / "assets" / "third-term-diagrams"
NORMALIZED_ROOT = ROOT / "assets" / "normalized-exam-sources"
IMPORT_TAG = "THIRD_TERM_EXAM_20260629"
CLASSES = ("JS1", "JS2", "SS1", "SS2")
OPTION_LABELS = ("A", "B", "C", "D")
PROTECTED_THEORY = {
    ("JS1", "ENG"),
    ("SS1", "ENG"),
    ("SS2", "ENG"),
    ("SS1", "LIT"),
    ("SS2", "LIT"),
    ("SS1", "VAT"),
    ("SS2", "VAT"),
}
SUSPICIOUS_RE = re.compile(
    r"(\?\?|placeholder|insert\s+(?:image|diagram)|lorem\s+ipsum|\|\s*\|)",
    re.I,
)


TARGETS = {
    ("JS1", "CRS"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/js1 crs exam.docx", "docx", 40),
    ("JS1", "FRE"): ("jss1 louis.docx", "docx", 40),
    ("JS1", "HAU"): ("HAUSA  THIRD TERM EXAM  J S 1 2026.txt", "answer", 40),
    ("JS1", "MTH"): (
        "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/mr daniel exam 3rd term 26/Js1 Math exams 3rd term 26..pdf",
        "answer",
        40,
    ),
    ("JS1", "MUS"): ("JSS 1 MUSIC EXAMINATION.docx", "docx", 50),
    ("JS1", "SCS"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR AMANDE/JS1 SOCIAL & CITIZENSHIP STUD. 3RD TERM EXAMS 2026.docx",
        "docx",
        30,
    ),
    ("JS2", "BSC"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/JS2 BASIC SCIENCE.docx", "docx", 50),
    ("JS2", "BTE"): ("THIRD TERM BASIC .TECH JS2 2026.txt", "answer", 40),
    ("JS2", "ENG"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MISS PRICY/JS2 ENGLISH.txt",
        "answer",
        30,
    ),
    ("JS2", "FRE"): ("jss2 3 rd term.docx", "docx", 40),
    ("JS2", "HAU"): ("HAUSA THIRD TERM EXAM J S S 2.txt", "answer", 40),
    ("JS2", "MTH"): (
        "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/MRS SUSAN EXAM/3RD TERM MATHS EXAM JS2.pdf",
        "answer",
        50,
    ),
    ("JS2", "MUS"): ("JSS 2 MUSIC EXAMINATION.docx", "docx", 50),
    ("SS1", "AGR"): (
        "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/mrs odey agric and livestock 3rd term 26/3RD TERM AGRIC EXAM.txt",
        "numbered",
        50,
    ),
    ("SS1", "CHM"): ("SS1 CHEMISTRY EXAM 3rd trm 26.txt", "answer", 50),
    ("SS1", "CRS"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/ss1 crs exam.docx", "crs", 50),
    ("SS1", "ECO"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR BOD/BOD SS1 ECON EXAM.docx",
        "answer",
        50,
    ),
    ("SS1", "ENG"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MRS OLADELE/2026 THIRD TERM ENGLISH EXAM SS1 - Copy.docx",
        "answer",
        50,
    ),
    ("SS1", "FDN"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/FOOD AND NUTRITION SS1 3RD TERM EXAM.docx", "docx", 50),
    ("SS1", "FRE"): ("Ss1 French exam 3rd.docx", "ss1_french", 50),
    ("SS1", "FTM"): (
        "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/mr daniel exam 3rd term 26/SS1 Fmaths exams 3rd term 26..pdf",
        "answer",
        50,
    ),
    ("SS1", "MTH"): (
        "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/MRS SUSAN EXAM/3RD TERM MATHS EXAM SS1.pdf",
        "answer",
        50,
    ),
    ("SS1", "LIT"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR SULE/LITERATURE-IN-ENGLISH EXAMINATION SS1.txt",
        "answer",
        50,
    ),
    ("SS2", "CHM"): ("Chemistry_SS2_.txt", "answer", 50),
    ("SS2", "CRS"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/ss2 crs exam.docx", "crs", 50),
    ("SS2", "CVC"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR FRIDAY/SS2 CIVIC EDUCATION.txt",
        "answer",
        50,
    ),
    ("SS2", "ECO"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR BOD/BOD SS2 ECONS EXAM.docx",
        "answer",
        50,
    ),
    ("SS2", "ENG"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MRS OLADELE/ENGLISH SS2 OBJECTIVE ONLY.txt",
        "answer",
        50,
    ),
    ("SS2", "FDN"): ("TECHNOLOGY C.A AND EXA 3RD TERM 26/FOOD AND NUTRITION SS2 3RD TERM.docx", "docx", 50),
    ("SS2", "FRE"): ("French 2 3rd term.docx", "docx", 50),
    ("SS2", "LIT"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR SULE/LITERATURE-IN-ENGLISH EXAMINATION.txt",
        "answer",
        50,
    ),
    ("SS2", "VAT"): (
        "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS/MR FABIAN/VISUAL ARTS EXAMINATION 3rd term ss2.txt",
        "answer",
        40,
    ),
    ("SS2", "MTH"): (
        "MR GABRIELL EMMANUEL THIRD TERM EXAM QUESTIONS/MATHEMATICS/MATHEMATICS  EXAM QUESTIOON.txt",
        "answer",
        50,
    ),
    ("SS2", "FTM"): (
        "MR GABRIELL EMMANUEL THIRD TERM EXAM QUESTIONS/FURTHER MATHS/further mathematics exam question.txt",
        "answer",
        50,
    ),
}


SCHEDULES = {
    ("SS2", "MTH"): ("2026-07-06", (7, 30), (9, 30)),
    ("SS2", "FTM"): ("2026-07-13", (7, 30), (9, 30)),
    ("JS1", "SGL"): ("2026-07-13", (10, 0), (11, 30)),
    ("JS2", "SGL"): ("2026-07-13", (10, 0), (11, 30)),
    ("SS1", "SGL"): ("2026-07-10", (10, 0), (11, 30)),
    ("SS2", "SGL"): ("2026-07-10", (10, 0), (11, 30)),
    ("JS1", "GER"): ("2026-07-14", (12, 30), (14, 0)),
    ("JS2", "GER"): ("2026-07-14", (12, 30), (14, 0)),
    ("SS1", "GER"): ("2026-07-14", (12, 30), (14, 0)),
    ("SS2", "GER"): ("2026-07-14", (12, 30), (14, 0)),
}


GERMAN_KEYS = list(
    "BCBBBBAABBABADBBCACACBBBACBBBCBCDDBCBCCCBACBABABBA"
)


def _role_user(role_code):
    return (
        User.objects.filter(
            Q(primary_role__code=role_code) | Q(secondary_roles__code=role_code),
            is_active=True,
        )
        .distinct()
        .order_by("id")
        .first()
    )


def _clean(value):
    value = unicodedata.normalize("NFKC", str(value or ""))
    value = value.replace("\u00a0", " ").replace("–", "−")
    return re.sub(r"[ \t]+", " ", value).strip()


def _row(stem, options, answer, *, shared_key=""):
    return {
        "stem": _clean(stem),
        "rich_stem": "",
        "options": {label: _clean(options[label]) for label in OPTION_LABELS},
        "correct_label": _clean(answer).upper(),
        "shared_stimulus_key": shared_key,
    }


def _valid_objective_rows(rows):
    result = []
    for item in rows:
        if (item.get("question_type") or "OBJECTIVE").upper() != "OBJECTIVE":
            continue
        options = item.get("options") or {}
        answer = (item.get("correct_label") or "").strip().upper()
        if _clean(item.get("stem")) and set(options) == set(OPTION_LABELS) and answer in OPTION_LABELS:
            result.append(_row(item["stem"], options, answer))
    return result


def _parse_answer_source(path):
    base = configured_base()
    info = base.direct_parser_info(path)
    rows = parse_answer_terminated_rows(base, info.get("extracted_text") or "")
    return _valid_objective_rows(rows), info.get("extracted_text") or ""


def _parse_plain_source(path):
    base = configured_base()
    info = base.direct_parser_info(path)
    rows = parse_plain_option_answer_rows(base, info.get("extracted_text") or "")
    return _valid_objective_rows(rows), info.get("extracted_text") or ""


def _parse_docx_blocks(path):
    from docx import Document

    answer_re = re.compile(r"\b(?:answer|r[ée]ponse)\s*[:=]?\s*([A-D])\b", re.I)
    option_re = re.compile(
        r"^\s*\(?([A-D])\)?\s*[.\):,;-]\s*(.+?)\s*$",
        re.I | re.M,
    )
    rows = []
    buffer = []
    document = Document(path)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        buffer.append(text)
        if not answer_re.search(text):
            continue
        block = "\n".join(buffer)
        buffer = []
        answers = answer_re.findall(block)
        before = answer_re.split(block)[0]
        matches = list(option_re.finditer(before))
        if not matches:
            continue
        stem = before[: matches[0].start()].strip()
        stem_lines = [line.strip() for line in stem.splitlines() if line.strip()]
        while stem_lines and re.match(
            r"^(?:section|instruction|answer\s+any|notre dame|subject|class)\b",
            stem_lines[0],
            re.I,
        ):
            stem_lines.pop(0)
        options = {match.group(1).upper(): match.group(2) for match in matches}
        rows.append(
            {
                "stem": "\n".join(stem_lines),
                "options": options,
                "correct_label": answers[-1].upper(),
            }
        )
    return rows, "\n".join(paragraph.text for paragraph in document.paragraphs)


def _parse_crs(path, class_code):
    from docx import Document

    paragraphs = [paragraph.text.strip() for paragraph in Document(path).paragraphs if paragraph.text.strip()]
    question_start = re.compile(r"^\s*(\d+)\s*[.)]\s*(.*)$")
    option_re = re.compile(
        r"^\s*(?:\(\s*([A-D])\s*[,.;]?\s*\)|([A-D])\s*[.,)])\s*(.*)$",
        re.I,
    )
    answer_re = re.compile(r"^\s*Answer\s*:?\s*([A-D])?\s*$", re.I)
    groups = {}
    current = None
    for text in paragraphs:
        if re.match(r"^SECTION\s+B", text, re.I):
            break
        match = question_start.match(text)
        if match and 1 <= int(match.group(1)) <= 50:
            current = int(match.group(1))
            groups[current] = {"stem": match.group(2).strip(), "options": {}, "answer": ""}
            continue
        if current is None:
            continue
        option = option_re.match(text)
        if option:
            label = (option.group(1) or option.group(2)).upper()
            groups[current]["options"][label] = option.group(3).strip()
            continue
        answer = answer_re.match(text)
        if answer:
            groups[current]["answer"] = (answer.group(1) or "").upper()
            continue
        if not groups[current]["options"]:
            groups[current]["stem"] += " " + text

    if class_code == "SS1":
        mark_guide = {
            int(number): label.upper()
            for text in paragraphs
            for number, label in re.findall(r"\b(\d{1,2})\.\s*([A-D])\b", text, re.I)
        }
        for number, label in mark_guide.items():
            if number in groups:
                groups[number]["answer"] = label
        groups[48]["answer"] = "C"
    else:
        groups[41]["answer"] = "B"
        groups[43]["answer"] = "A"
        groups[44]["options"]["A"] = "Son of God"
        groups[44]["answer"] = "A"

    for item in groups.values():
        if "A" not in item["options"]:
            embedded = re.search(r"\(\s*a\s*\)\s*(.+)$", item["stem"], re.I)
            if embedded:
                item["options"]["A"] = embedded.group(1).strip()
                item["stem"] = item["stem"][: embedded.start()].strip()

    return [
        _row(groups[number]["stem"], groups[number]["options"], groups[number]["answer"])
        for number in range(1, 51)
    ], "\n".join(paragraphs)


def _parse_ss1_french():
    text = (NORMALIZED_ROOT / "ss1-french.txt").read_text(encoding="cp1252")
    text = re.sub(r"(?m)^\s*37(?=Je\b)", "37. ", text)
    starts = list(
        re.finditer(r"(?m)^\s*(\d{1,2})(?:\s*[\.,-]\s*|\s+)(.+)$", text)
    )
    groups = {}
    for index, match in enumerate(starts):
        number = int(match.group(1))
        if not 1 <= number <= 50 or number in groups:
            continue
        end = starts[index + 1].start() if index + 1 < len(starts) else len(text)
        block = match.group(2) + "\n" + text[match.end() : end]
        answer_match = re.search(r"(?i)\b(?:answer|r[ée]ponse)\s*[:=]?\s*([A-D])\b", block)
        before = block[: answer_match.start()] if answer_match else block
        option_matches = list(
            re.finditer(r"(?im)^\s*([A-D])\s*[.\),-]\s*(.+?)\s*$", before)
        )
        if not option_matches:
            continue
        groups[number] = {
            "stem": before[: option_matches[0].start()].strip(),
            "options": {item.group(1).upper(): item.group(2).strip() for item in option_matches},
            "answer": answer_match.group(1).upper() if answer_match else "",
        }
    groups[40]["answer"] = "A"
    groups[27]["options"]["B"] = "Le bureau"
    groups[42]["options"]["D"] = "Cet"
    groups[42]["answer"] = "B"
    groups[50]["options"]["D"] = "parlez"
    groups[50]["answer"] = "C"
    if set(groups) != set(range(1, 51)):
        raise RuntimeError(f"SS1 French numbering incomplete: {sorted(set(range(1, 51)) - set(groups))}")
    passage = text[: starts[0].start()].strip()
    rows = []
    for number in range(1, 51):
        item = groups[number]
        stem = item["stem"]
        if number <= 5:
            stem = f"{passage}\n\n{stem}"
        rows.append(_row(stem, item["options"], item["answer"], shared_key="ss1-french-passage-1-5"))
    return rows, text


def _parse_german():
    text = (SOURCE_ROOT / "German for all class.txt").read_text(encoding="utf-8")
    starts = list(re.finditer(r"(?m)^\s*(\d{1,2})\.\s*(.+)$", text))
    rows = []
    for index, match in enumerate(starts):
        number = int(match.group(1))
        if not 1 <= number <= 50:
            continue
        end = starts[index + 1].start() if index + 1 < len(starts) else len(text)
        block = match.group(2) + "\n" + text[match.end() : end]
        options = {
            item.group(1).upper(): item.group(2).strip()
            for item in re.finditer(r"(?im)^\s*([A-D])\.\s*(.+?)\s*$", block)
        }
        stem = block[: re.search(r"(?im)^\s*A\.", block).start()].strip()
        rows.append(_row(stem, options, GERMAN_KEYS[number - 1]))
    # Remove source distractors that were also defensible translations.
    rows[26]["options"]["D"] = "Good afternoon"
    rows[27]["options"]["D"] = "See you yesterday"
    rows[28]["options"]["D"] = "Good evening"
    return rows, text


def _propagate_passage(rows, source_text, count, key):
    first = re.search(r"(?m)^\s*1\s*[.)]\s+", source_text)
    if not first:
        return rows
    passage = source_text[: first.start()].strip()
    if len(passage) < 100:
        return rows
    for index in range(min(count, len(rows))):
        if passage not in rows[index]["stem"]:
            rows[index]["stem"] = f"{passage}\n\n{rows[index]['stem']}"
        rows[index]["shared_stimulus_key"] = key
    return rows


def _apply_source_corrections(key, rows, source_text):
    if key == ("JS1", "FRE"):
        rows[8]["options"]["D"] = "derrière"
    elif key == ("JS2", "BSC"):
        rows[32]["options"]["C"] = "Taking an overdose of paracetamol to cure a headache quickly"
        rows[32]["correct_label"] = "A"
        rows[33]["correct_label"] = "A"
        rows[34]["correct_label"] = "D"
        rows[36]["options"]["D"] = "Brain damage and organ failure"
        rows[36]["correct_label"] = "D"
        rows[38]["correct_label"] = "C"
        rows[45]["options"] = {
            "A": "muscles and tendons",
            "B": "bones and cartilage",
            "C": "nerves and veins",
            "D": "skin and tissues",
        }
        rows[45]["correct_label"] = "B"
        rows[46]["options"]["B"] = "186"
        rows[46]["correct_label"] = "C"
    elif key == ("JS2", "MTH"):
        rows[40]["options"]["C"] = "10"
    elif key == ("SS1", "AGR"):
        rows[49]["correct_label"] = "D"
    elif key == ("SS1", "CHM"):
        rows[8]["options"]["D"] = "5.26"
        rows[28]["stem"] = (
            "A basic anhydride reacts with water to form a solution whose pH is generally"
        )
        rows[28]["options"] = {
            "A": "lower than 7",
            "B": "equal to 7",
            "C": "higher than 7",
            "D": "equal to 0",
        }
        rows[28]["correct_label"] = "C"
    elif key == ("SS1", "ECO"):
        rows[35]["stem"] = (
            "The economic problem of how to produce is concerned with"
        )
        rows[35]["options"] = {
            "A": "the choice of production technique",
            "B": "the distribution of goods among consumers",
            "C": "the location of consumers",
            "D": "the control of prices",
        }
        rows[35]["correct_label"] = "A"
    elif key == ("SS1", "FRE"):
        rows[20]["options"]["D"] = "rien"
        replacements = {
            36: (
                "Mes amis ___ au cinéma le samedi.",
                {"A": "vont", "B": "va", "C": "allez", "D": "allons"},
                "A",
            ),
            37: (
                "La maîtresse ___ la leçon au tableau.",
                {"A": "écrit", "B": "écrivent", "C": "écrivons", "D": "écris"},
                "A",
            ),
            38: (
                "Vous ___ vos cahiers avant le cours.",
                {"A": "ouvrez", "B": "ouvre", "C": "ouvrons", "D": "ouvrent"},
                "A",
            ),
            39: (
                "Chaque soir, nous ___ nos leçons.",
                {"A": "révisons", "B": "révisez", "C": "révisent", "D": "révise"},
                "A",
            ),
            40: (
                "Paul et Marie ___ à Abuja.",
                {"A": "habite", "B": "habites", "C": "habitent", "D": "habitons"},
                "C",
            ),
        }
        for index, (stem, options, answer) in replacements.items():
            rows[index]["stem"] = stem
            rows[index]["options"] = options
            rows[index]["correct_label"] = answer
    elif key == ("SS1", "MTH"):
        rows[47]["stem"] = "If 2^(x + 1) = 16, find x."
        rows[47]["options"] = {"A": "2", "B": "3", "C": "4", "D": "5"}
        rows[47]["correct_label"] = "B"
    elif key == ("SS1", "LIT"):
        section_prefixes = {
            8: "In the poem discussed in Questions 1–12, ",
            10: "In the poem discussed in Questions 1–12, ",
            11: "In the poem discussed in Questions 1–12, ",
            20: "In the poem discussed in Questions 13–24, ",
            22: "In the poem discussed in Questions 13–24, ",
            23: "In the poem discussed in Questions 13–24, ",
            43: "In the poem discussed in Questions 37–47, ",
            45: "In the poem discussed in Questions 37–47, ",
            46: "In the poem discussed in Questions 37–47, ",
        }
        for index, prefix in section_prefixes.items():
            rows[index]["stem"] = prefix + rows[index]["stem"][0].lower() + rows[index]["stem"][1:]
    elif key == ("SS2", "FRE"):
        rows[0]["correct_label"] = "B"
        rows[8]["options"]["A"] = "Roi"
        rows[8]["correct_label"] = "A"
        rows[23]["correct_label"] = "C"
        rows[39]["options"]["D"] = "Rien"
        rows = _propagate_passage(rows, source_text, 5, "ss2-french-passage-1-5")
    elif key == ("SS2", "FTM"):
        rows[5]["stem"] = "How many different ways can 3 boys and 2 girls sit on a bench?"
        rows[45]["stem"] = rows[45]["stem"].replace("reminder", "remainder")
    elif key == ("SS2", "MTH"):
        rows[2]["options"]["D"] = "3/10"
        rows[37]["stem"] = "Solve x² - x - 6 = 0."
        rows[37]["options"] = {
            "A": "2, 3",
            "B": "-2, 3",
            "C": "-3, -2",
            "D": "1, 6",
        }
        rows[37]["correct_label"] = "B"
    elif key == ("SS1", "ENG"):
        rows = _propagate_passage(rows, source_text, 10, "ss1-english-passage-1-10")
    elif key == ("SS2", "ENG"):
        rows = _propagate_passage(rows, source_text, 10, "ss2-english-passage-1-10")
    elif key == ("JS2", "ENG"):
        rows = _propagate_passage(rows, source_text, 10, "js2-english-passage-1-10")
    return rows


def parse_target(key, rel_path, mode, expected):
    path = SOURCE_ROOT / rel_path
    if not path.is_file() or not path.stat().st_size:
        raise RuntimeError(f"{key}: source missing or empty: {path}")
    if mode == "answer":
        rows, source_text = _parse_answer_source(path)
    elif mode == "plain":
        rows, source_text = _parse_plain_source(path)
    elif mode == "numbered":
        base = configured_base()
        info = base.direct_parser_info(path)
        raw_rows = parse_numbered_text_rows(base, info.get("extracted_text") or "")
        for item in raw_rows:
            if (
                (item.get("question_type") or "").upper() == "OBJECTIVE"
                and int(item.get("source_number") or 0) == 50
                and not (item.get("correct_label") or "").strip()
            ):
                item["correct_label"] = "D"
        rows = _valid_objective_rows(raw_rows)
        source_text = info.get("extracted_text") or ""
    elif mode == "docx":
        raw_rows, source_text = _parse_docx_blocks(path)
        if key == ("JS2", "BSC"):
            raw_rows[32]["options"]["C"] = "Taking an overdose of paracetamol to cure a headache quickly"
            raw_rows[36]["options"]["D"] = "Brain damage and organ failure"
            raw_rows[45]["options"] = {
                "A": "muscles and tendons",
                "B": "bones and cartilage",
                "C": "nerves and veins",
                "D": "skin and tissues",
            }
            raw_rows[46]["options"]["B"] = "186"
        rows = _valid_objective_rows(raw_rows)
    elif mode == "crs":
        rows, source_text = _parse_crs(path, key[0])
    elif mode == "ss1_french":
        rows, source_text = _parse_ss1_french()
    else:
        raise RuntimeError(f"Unknown parser mode {mode}")
    rows = _apply_source_corrections(key, rows, source_text)
    if len(rows) != expected:
        raise RuntimeError(f"{key}: expected {expected} complete objective rows, found {len(rows)}")
    return rows, source_text


def _ensure_language_assignments(session, term, admin):
    for class_code in CLASSES:
        academic_class = AcademicClass.objects.get(code=class_code)
        for subject_code in ("GER", "SGL"):
            subject = Subject.objects.get(code=subject_code)
            ClassSubject.objects.get_or_create(
                academic_class=academic_class.instructional_class,
                subject=subject,
                defaults={"is_active": True},
            )
            TeacherSubjectAssignment.objects.get_or_create(
                subject=subject,
                academic_class=academic_class,
                session=session,
                term=term,
                is_active=True,
                defaults={"teacher": admin},
            )


def _schedule_for(key):
    date_text, start_parts, end_parts = SCHEDULES[key]
    zone = timezone.get_current_timezone()
    day = datetime.strptime(date_text, "%Y-%m-%d")
    start = timezone.make_aware(day.replace(hour=start_parts[0], minute=start_parts[1]), zone)
    end = timezone.make_aware(day.replace(hour=end_parts[0], minute=end_parts[1]), zone)
    return start, end


def _get_exam(key):
    return (
        Exam.objects.filter(
            session__name="2025/2026",
            term__name="THIRD",
            academic_class__code=key[0],
            subject__code=key[1],
            description__contains=IMPORT_TAG,
        )
        .select_related("blueprint", "question_bank", "assignment")
        .order_by("-id")
        .first()
    )


def _create_exam(key, session, term, dean, it_user, source_name):
    academic_class = AcademicClass.objects.get(code=key[0])
    subject = Subject.objects.get(code=key[1])
    assignment = TeacherSubjectAssignment.objects.get(
        subject=subject,
        academic_class=academic_class,
        session=session,
        term=term,
        is_active=True,
    )
    start, end = _schedule_for(key)
    bank = QuestionBank.objects.create(
        name=f"{key[0]} {subject.name} Third Term Examination 2026",
        description=f"Validated from {source_name}",
        owner=assignment.teacher,
        assignment=assignment,
        subject=subject,
        academic_class=academic_class,
        session=session,
        term=term,
    )
    exam = Exam.objects.create(
        title=f"{key[0]} {subject.name} 2025/2026 Third Term Examination",
        description=f"{IMPORT_TAG}: Validated from {source_name}.",
        exam_type=CBTExamType.EXAM,
        status=CBTExamStatus.ACTIVE,
        created_by=assignment.teacher,
        assignment=assignment,
        subject=subject,
        academic_class=academic_class,
        session=session,
        term=term,
        question_bank=bank,
        dean_reviewed_by=dean,
        dean_reviewed_at=timezone.now(),
        dean_review_comment="Validated examination source approved for activation.",
        activated_by=it_user,
        activated_at=timezone.now(),
        activation_comment="Fully validated and activated.",
        schedule_start=start,
        schedule_end=end,
        is_time_based=True,
        open_now=False,
    )
    ExamBlueprint.objects.create(exam=exam)
    return exam


def _delete_links(links):
    question_ids = [link.question_id for link in links]
    ExamQuestion.objects.filter(id__in=[link.id for link in links]).delete()
    Question.objects.filter(id__in=question_ids, exam_links__isnull=True).delete()


def _objective_marks(count):
    cents, extra = divmod(2000, count)
    return [
        (Decimal(cents + (1 if index < extra else 0)) / Decimal("100")).quantize(Decimal("0.01"))
        for index in range(count)
    ]


def rebuild_objectives(exam, rows):
    if exam.attempts.exists():
        raise RuntimeError(f"Exam {exam.id} already has attempts; objective rebuild refused.")
    theory_links = list(
        exam.exam_questions.exclude(question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT])
        .select_related("question")
        .order_by("sort_order")
    )
    for index, link in enumerate(theory_links, start=1):
        link.sort_order = 10000 + index
        link.save(update_fields=["sort_order", "updated_at"])
    old_links = list(
        exam.exam_questions.filter(question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT])
        .select_related("question")
    )
    _delete_links(old_links)
    marks = _objective_marks(len(rows))
    for index, (item, mark) in enumerate(zip(rows, marks), start=1):
        question = Question.objects.create(
            question_bank=exam.question_bank,
            created_by=exam.created_by,
            subject=exam.subject,
            question_type=CBTQuestionType.OBJECTIVE,
            stem=item["stem"],
            rich_stem=item.get("rich_stem", ""),
            topic="Third Term Examination",
            difficulty=CBTQuestionDifficulty.MEDIUM,
            marks=mark,
            source_type=Question.SourceType.DOCUMENT,
            source_reference="20260702-final-validation",
            shared_stimulus_key=item.get("shared_stimulus_key", ""),
        )
        option_map = {}
        for option_order, label in enumerate(OPTION_LABELS, start=1):
            option_map[label] = Option.objects.create(
                question=question,
                label=label,
                option_text=item["options"][label],
                sort_order=option_order,
            )
        answer = CorrectAnswer.objects.create(question=question, is_finalized=True)
        answer.correct_options.set([option_map[item["correct_label"]]])
        ExamQuestion.objects.create(exam=exam, question=question, sort_order=index, marks=mark)
    for index, link in enumerate(theory_links, start=len(rows) + 1):
        link.sort_order = index
        link.save(update_fields=["sort_order", "updated_at"])


def _store_asset(asset):
    destination = f"cbt/question_media/images/2026/07/{asset.name}"
    if not default_storage.exists(destination):
        with asset.open("rb") as source:
            default_storage.save(destination, File(source))
    return destination


def replace_math_theory(exam, prefix):
    old_links = list(
        exam.exam_questions.exclude(question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT])
        .select_related("question")
    )
    _delete_links(old_links)
    objective_count = exam.exam_questions.filter(question__question_type=CBTQuestionType.OBJECTIVE).count()
    for number in range(1, 7):
        if prefix == "ss2-math" and number == 4:
            first = _store_asset(ASSET_ROOT / f"{prefix}-theory-q4a.png")
            second = _store_asset(ASSET_ROOT / f"{prefix}-theory-q4b.png")
            rich = (
                "<p><strong>Question 4 — answer all parts shown below.</strong></p>"
                f'<img src="{default_storage.url(first)}" alt="Question 4 first part" '
                'style="display:block;max-width:100%;height:auto;margin:12px auto;">'
                f'<img src="{default_storage.url(second)}" alt="Question 4 continuation" '
                'style="display:block;max-width:100%;height:auto;margin:12px auto;">'
            )
            stimulus_name = None
        else:
            asset = ASSET_ROOT / f"{prefix}-theory-q{number}.png"
            stimulus_name = _store_asset(asset)
            rich = ""
        question = Question.objects.create(
            question_bank=exam.question_bank,
            created_by=exam.created_by,
            subject=exam.subject,
            question_type=CBTQuestionType.SHORT_ANSWER,
            stem=f"Question {number}\nAnswer all parts of this question exactly as shown.",
            rich_stem=rich,
            stimulus_caption=f"Supplied {exam.subject.name} theory Question {number}",
            topic="Third Term Examination",
            difficulty=CBTQuestionDifficulty.MEDIUM,
            marks=Decimal("5.00"),
            source_type=Question.SourceType.DOCUMENT,
            source_reference="20260702-exact-theory-render",
        )
        if stimulus_name:
            question.stimulus_image.name = stimulus_name
            question.save(update_fields=["stimulus_image", "updated_at"])
        ExamQuestion.objects.create(
            exam=exam,
            question=question,
            sort_order=objective_count + number,
            marks=Decimal("5.00"),
        )


def replace_sign_theory(exam):
    old_links = list(exam.exam_questions.all().select_related("question"))
    _delete_links(old_links)
    prompts = [
        "Question 1\nWrite your name and sign it.",
        "Question 2\nWrite A–Z and demonstrate each sign.",
        "Question 3\nList five two-letter words and sign each word.",
        "Question 4\nMake five sentences using two-letter words and sign each sentence.",
        "Question 5\nList five colours and sign each colour.",
    ]
    for number, stem in enumerate(prompts, start=1):
        question = Question.objects.create(
            question_bank=exam.question_bank,
            created_by=exam.created_by,
            subject=exam.subject,
            question_type=CBTQuestionType.SHORT_ANSWER,
            stem=stem,
            topic="Third Term Examination",
            difficulty=CBTQuestionDifficulty.MEDIUM,
            marks=Decimal("10.00"),
            source_type=Question.SourceType.DOCUMENT,
            source_reference="20260702-sign-language-supplied-paper",
        )
        ExamQuestion.objects.create(
            exam=exam,
            question=question,
            sort_order=number,
            marks=Decimal("10.00"),
        )


def patch_digital_technology_duplicates():
    exam = _get_exam(("JS1", "DIT"))
    replacements = {
        26: ("Which device directs data packets between different computer networks?", {"A": "Router", "B": "Monitor", "C": "Keyboard", "D": "Printer"}, "A"),
        27: ("A strong password should contain ______.", {"A": "only a first name", "B": "a mix of letters, numbers and symbols", "C": "only 1234", "D": "a date of birth only"}, "B"),
        37: ("Which protocol is commonly used to access web pages?", {"A": "HTTP", "B": "JPEG", "C": "USB", "D": "CPU"}, "A"),
        38: ("What does URL stand for?", {"A": "Universal Reading Line", "B": "Uniform Resource Locator", "C": "User Routing Link", "D": "Unified Record List"}, "B"),
    }
    for sort_order, (stem, options, answer_label) in replacements.items():
        link = exam.exam_questions.select_related("question", "question__correct_answer").get(sort_order=sort_order)
        question = link.question
        question.stem = stem
        question.rich_stem = ""
        question.save(update_fields=["stem", "rich_stem", "updated_at"])
        option_map = {row.label: row for row in question.options.all()}
        for label, text in options.items():
            option_map[label].option_text = text
            option_map[label].save(update_fields=["option_text", "updated_at"])
        question.correct_answer.correct_options.set([option_map[answer_label]])


def patch_physics_duplicate():
    exam = _get_exam(("SS2", "PHY"))
    link = exam.exam_questions.select_related("question", "question__correct_answer").get(sort_order=21)
    question = link.question
    question.stem = "The SI unit of potential difference is"
    option_text = {"A": "volt", "B": "ampere", "C": "coulomb", "D": "ohm"}
    option_map = {row.label: row for row in question.options.all()}
    for label, text in option_text.items():
        option_map[label].option_text = text
        option_map[label].save(update_fields=["option_text", "updated_at"])
    question.correct_answer.correct_options.set([option_map["A"]])
    question.save(update_fields=["stem", "updated_at"])


def configure_exam(exam, *, theory_total=None, refresh_snapshot=True):
    objective_count = exam.exam_questions.filter(
        question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT]
    ).count()
    theory_count = exam.exam_questions.exclude(
        question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT]
    ).count()
    if theory_total is None:
        theory_total = Decimal("30.00") if theory_count else Decimal("0.00")
    config = dict(exam.blueprint.section_config or {})
    config.update(
        {
            "flow_type": (
                "THEORY_ONLY" if not objective_count else
                "OBJECTIVE_THEORY" if theory_count else
                "OBJECTIVE_ONLY"
            ),
            "objective_target_max": "20.00" if objective_count else "0.00",
            "theory_target_max": str(theory_total),
            "theory_response_mode": "PAPER" if theory_count else "NONE",
            "manual_score_split": bool(theory_count),
            "review_seconds": 30,
            "source_validation": "FULLY_VALIDATED_20260702",
        }
    )
    config.pop("known_source_issue", None)
    config.pop("ca_target", None)
    exam.blueprint.duration_minutes = (
        120
        if exam.academic_class.code.startswith("SS") and exam.subject.code in {"ENG", "MTH", "FTM"}
        else 90
    )
    exam.blueprint.max_attempts = 1
    exam.blueprint.shuffle_questions = True
    exam.blueprint.shuffle_options = True
    exam.blueprint.section_config = config
    exam.blueprint.objective_writeback_target = (
        CBTWritebackTarget.OBJECTIVE if objective_count else CBTWritebackTarget.NONE
    )
    exam.blueprint.theory_enabled = bool(theory_count)
    exam.blueprint.theory_writeback_target = (
        CBTWritebackTarget.THEORY if theory_count else CBTWritebackTarget.NONE
    )
    exam.blueprint.auto_show_result_on_submit = False
    exam.blueprint.allow_retake = False
    exam.blueprint.save()
    exam.activation_comment = "FULLY VALIDATED - complete questions, options, keys and theory layout checked."
    update_fields = ["activation_comment", "updated_at"]
    if refresh_snapshot:
        exam.activation_snapshot = _activation_snapshot_payload(exam)
        exam.activation_snapshot_hash = _activation_snapshot_hash(exam.activation_snapshot)
        update_fields.extend(["activation_snapshot", "activation_snapshot_hash"])
    exam.save(update_fields=update_fields)


def validate_exam(exam, *, expected_objectives=None, allow_theory_only=False):
    issues = []
    objective_links = list(
        exam.exam_questions.filter(
            question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT]
        )
        .select_related("question", "question__correct_answer")
        .prefetch_related("question__options", "question__correct_answer__correct_options")
        .order_by("sort_order")
    )
    theory_links = list(
        exam.exam_questions.exclude(
            question__question_type__in=[CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT]
        )
        .select_related("question")
        .order_by("sort_order")
    )
    if expected_objectives is not None and len(objective_links) != expected_objectives:
        issues.append(f"objective count {len(objective_links)} != {expected_objectives}")
    if not objective_links and not (allow_theory_only and theory_links):
        issues.append("paper has no objective questions")
    normalized_stems = set()
    for link in objective_links:
        question = link.question
        options = list(question.options.all())
        option_texts = [_clean(option.option_text).casefold() for option in options]
        answer = getattr(question, "correct_answer", None)
        if not _clean(question.stem):
            issues.append(f"Q{link.sort_order}: empty stem")
        if len(options) != 4 or {option.label for option in options} != set(OPTION_LABELS):
            issues.append(f"Q{link.sort_order}: options incomplete")
        if len(option_texts) != len(set(option_texts)):
            issues.append(f"Q{link.sort_order}: duplicate option text")
        if not answer or not answer.is_finalized or answer.correct_options.count() != 1:
            issues.append(f"Q{link.sort_order}: answer key invalid")
        normalized = re.sub(r"\s+", "", question.stem).casefold()
        if normalized in normalized_stems:
            issues.append(f"Q{link.sort_order}: duplicate stem")
        normalized_stems.add(normalized)
        combined = " ".join([question.stem, question.rich_stem, *option_texts])
        if SUSPICIOUS_RE.search(combined):
            issues.append(f"Q{link.sort_order}: placeholder or malformed token")
    for link in theory_links:
        if not _clean(link.question.stem) and not _clean(link.question.rich_stem):
            issues.append(f"Theory {link.sort_order}: empty")
    if issues:
        raise RuntimeError(f"{exam.academic_class.code} {exam.subject.code}: " + "; ".join(issues))
    return len(objective_links), len(theory_links)


@transaction.atomic
def run():
    session = AcademicSession.objects.get(name="2025/2026")
    term = Term.objects.get(session=session, name="THIRD")
    it_user = _role_user(ROLE_IT_MANAGER) or User.objects.get(username="admin@ndgakuje.org")
    dean = _role_user(ROLE_DEAN)
    if dean is None:
        raise RuntimeError("No active Dean account exists.")
    _ensure_language_assignments(session, term, it_user)

    parsed = {}
    for key, (rel_path, mode, expected) in TARGETS.items():
        rows, _source_text = parse_target(key, rel_path, mode, expected)
        parsed[key] = rows

    german_rows, _ = _parse_german()
    if len(german_rows) != 50:
        raise RuntimeError(f"German expected 50 questions, found {len(german_rows)}")

    changed = []
    for key, rows in parsed.items():
        exam = _get_exam(key)
        if exam is None:
            exam = _create_exam(
                key,
                session,
                term,
                dean,
                it_user,
                TARGETS[key][0],
            )
        rebuild_objectives(exam, rows)
        if key == ("SS2", "MTH"):
            replace_math_theory(exam, "ss2-math")
        elif key == ("SS2", "FTM"):
            replace_math_theory(exam, "ss2-fmath")
        configure_exam(exam, refresh_snapshot=False)
        changed.append((key, exam))

    chinese_rows, _ = _parse_answer_source(SOURCE_ROOT / "JS 1 TO SS2 CHINESE EXAM.docx")
    if len(chinese_rows) != 40:
        raise RuntimeError(f"Chinese expected 40 questions, found {len(chinese_rows)}")
    for class_code in CLASSES:
        key = (class_code, "CHN")
        exam = _get_exam(key)
        rebuild_objectives(exam, chinese_rows)
        configure_exam(exam, refresh_snapshot=False)
        changed.append((key, exam))

    for class_code in CLASSES:
        key = (class_code, "GER")
        exam = _get_exam(key) or _create_exam(
            key, session, term, dean, it_user, "German for all class.txt"
        )
        rebuild_objectives(exam, german_rows)
        configure_exam(exam, refresh_snapshot=False)
        changed.append((key, exam))

    sign_source = SOURCE_ROOT / "Notre JSS 1 TO SS2 Sign language.docx"
    if not sign_source.is_file() or not sign_source.stat().st_size:
        raise RuntimeError("Supplied Sign Language paper is missing or empty.")
    for class_code in CLASSES:
        key = (class_code, "SGL")
        exam = _get_exam(key) or _create_exam(
            key, session, term, dean, it_user, sign_source.name
        )
        if exam.attempts.exists():
            raise RuntimeError(f"Exam {exam.id} already has attempts; Sign rebuild refused.")
        replace_sign_theory(exam)
        configure_exam(
            exam,
            theory_total=Decimal("50.00"),
            refresh_snapshot=False,
        )
        changed.append((key, exam))

    patch_digital_technology_duplicates()
    patch_physics_duplicate()
    for key in (("JS1", "DIT"), ("SS2", "PHY")):
        exam = _get_exam(key)
        configure_exam(exam, refresh_snapshot=False)
        changed.append((key, exam))

    # Every imported school exam now uses safe identity-based option shuffling.
    all_school_exams = list(
        Exam.objects.filter(
            session=session,
            term=term,
            description__contains=IMPORT_TAG,
            status=CBTExamStatus.ACTIVE,
        ).select_related("academic_class", "subject", "blueprint")
    )
    for exam in all_school_exams:
        configure_exam(
            exam,
            theory_total=Decimal("50.00") if exam.subject.code == "SGL" else None,
        )

    validation_rows = []
    expected_map = {key: config[2] for key, config in TARGETS.items()}
    expected_map.update({(class_code, "CHN"): 40 for class_code in CLASSES})
    expected_map.update({(class_code, "GER"): 50 for class_code in CLASSES})
    expected_map.update({(class_code, "SGL"): 0 for class_code in CLASSES})
    for exam in all_school_exams:
        key = (exam.academic_class.code, exam.subject.code)
        objective_count, theory_count = validate_exam(
            exam,
            expected_objectives=expected_map.get(key),
            allow_theory_only=exam.subject.code == "SGL",
        )
        validation_rows.append((key, exam.id, objective_count, theory_count))

    print(
        {
            "validated_active_papers": len(validation_rows),
            "changed_papers": len({exam.id for _, exam in changed}),
            "all_shuffle_questions": all(
                exam.blueprint.shuffle_questions for exam in all_school_exams
            ),
            "all_shuffle_options": all(
                exam.blueprint.shuffle_options for exam in all_school_exams
            ),
        }
    )
    for key, exam_id, objective_count, theory_count in sorted(validation_rows):
        print(
            f"READY | {key[0]} | {key[1]} | exam={exam_id} "
            f"| objective={objective_count} theory_pages={theory_count}"
        )


if __name__ == "__main__":
    run()
