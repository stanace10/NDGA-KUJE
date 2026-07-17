"""Audit and import the supplied 2025/2026 Third Term examination papers.

The manifest is deliberately explicit: a missing paper is reported and is never
replaced with a CA document or a paper for another class. Exams are scheduled
from the approved junior/senior timetable and remain closed until their slot.

Run without ``--apply`` for the readiness audit.
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", os.getenv("DJANGO_SETTINGS_MODULE", "core.settings"))

import django

django.setup()

from docx import Document
from django.core.files import File

from apps.academics.models import AcademicSession, TeacherSubjectAssignment, Term
from apps.cbt.models import CBTExamType, CBTWritebackTarget, Exam
from apps.cbt.services import _apply_exam_row_marks as apply_exam_row_marks
from scripts.import_ca23_third_term_20260624 import load_base_importer
from scripts.import_thursday_ca import parse_structured_docx_rows


SOURCE_ROOT = ROOT / "SCHOOL FOLDER" / "3RD TERM 2ND CA AND EXAM"
IMPORT_TAG = "THIRD_TERM_EXAM_20260629"

SLOTS = {
    "J1": ("2026-07-06", "7:30-9:00"),
    "J2": ("2026-07-06", "9:30-11:00"),
    "J3": ("2026-07-06", "12:10-1:40"),
    "J4": ("2026-07-07", "7:30-9:00"),
    "J5": ("2026-07-07", "9:30-11:00"),
    "J6": ("2026-07-07", "12:10-1:40"),
    "J7": ("2026-07-08", "7:30-9:00"),
    "J8": ("2026-07-08", "9:30-11:00"),
    "J9": ("2026-07-08", "12:10-1:40"),
    "J10": ("2026-07-09", "7:30-9:00"),
    "J11": ("2026-07-09", "9:30-11:00"),
    "J12": ("2026-07-09", "12:10-1:40"),
    "J13": ("2026-07-10", "7:30-9:00"),
    "J14": ("2026-07-10", "9:30-11:00"),
    "J15": ("2026-07-10", "12:10-1:40"),
    "J16": ("2026-07-13", "7:30-9:00"),
    "J17": ("2026-07-13", "9:30-11:00"),
    "J18": ("2026-07-13", "12:10-1:40"),
    "J19": ("2026-07-14", "7:30-9:00"),
    "J20": ("2026-07-14", "9:30-11:00"),
    "J21": ("2026-07-14", "12:10-1:40"),
    "S1": ("2026-07-06", "7:30-9:00"),
    "S2": ("2026-07-06", "9:30-11:00"),
    "S3": ("2026-07-06", "12:10-1:40"),
    "S4": ("2026-07-07", "7:30-9:00"),
    "S5": ("2026-07-07", "9:30-11:00"),
    "S6": ("2026-07-07", "12:10-1:40"),
    "S7": ("2026-07-08", "7:30-9:00"),
    "S8": ("2026-07-08", "9:30-11:00"),
    "S9": ("2026-07-08", "12:10-1:40"),
    "S10": ("2026-07-09", "7:30-9:00"),
    "S11": ("2026-07-09", "9:30-11:00"),
    "S12": ("2026-07-09", "12:10-1:40"),
    "S13": ("2026-07-10", "7:30-9:00"),
    "S14": ("2026-07-10", "9:30-11:00"),
    "S15": ("2026-07-10", "12:10-1:40"),
    "S16": ("2026-07-13", "7:30-9:00"),
    "S17": ("2026-07-13", "9:30-11:00"),
    "S18": ("2026-07-13", "12:10-1:40"),
}

# (class, subject) -> timetable slot. This is the readiness checklist.
EXPECTED = {
    ("JS1", "MTH"): "J1", ("JS1", "SCS"): "J2", ("JS1", "MUS"): "J3",
    ("JS1", "ENG"): "J4", ("JS1", "DIT"): "J5", ("JS1", "HIS"): "J6",
    ("JS1", "INS"): "J7", ("JS1", "CRS"): "J8", ("JS1", "LIV"): "J9",
    ("JS1", "CCA"): "J11", ("JS1", "FRE"): "J12", ("JS1", "BST"): "J13",
    ("JS1", "CHN"): "J14", ("JS1", "FAS"): "J15", ("JS1", "PHE"): "J16",
    ("JS1", "SGL"): "J17", ("JS1", "IGB"): "J19", ("JS1", "YOR"): "J19",
    ("JS1", "HAU"): "J20", ("JS1", "GER"): "J21",
    ("JS2", "MTH"): "J1", ("JS2", "SST"): "J2", ("JS2", "MUS"): "J3",
    ("JS2", "ENG"): "J4", ("JS2", "CSC"): "J5", ("JS2", "HIS"): "J6",
    ("JS2", "BTE"): "J7", ("JS2", "CRS"): "J8", ("JS2", "AGR"): "J9",
    ("JS2", "BSC"): "J10", ("JS2", "CCA"): "J11", ("JS2", "FRE"): "J12",
    ("JS2", "BST"): "J13", ("JS2", "CHN"): "J14", ("JS2", "HEC"): "J15",
    ("JS2", "PHE"): "J16", ("JS2", "SGL"): "J17", ("JS2", "CVC"): "J18",
    ("JS2", "IGB"): "J19", ("JS2", "YOR"): "J19", ("JS2", "HAU"): "J20",
    ("JS2", "GER"): "J21",
    ("SS1", "MTH"): "S1", ("SS1", "VAT"): "S2", ("SS1", "AGR"): "S2",
    ("SS1", "CHS"): "S3", ("SS1", "CHM"): "S4", ("SS1", "GOV"): "S4",
    ("SS1", "COM"): "S4", ("SS1", "ENG"): "S5", ("SS1", "BIO"): "S7",
    ("SS1", "FDN"): "S8", ("SS1", "GMT"): "S9", ("SS1", "PHY"): "S10",
    ("SS1", "LIT"): "S10", ("SS1", "ACC"): "S10", ("SS1", "DIT"): "S11",
    ("SS1", "CRS"): "S12", ("SS1", "CHN"): "S13", ("SS1", "SGL"): "S14",
    ("SS1", "ECO"): "S15", ("SS1", "FTM"): "S16", ("SS1", "GEO"): "S17",
    ("SS1", "FRE"): "S17", ("SS1", "LIV"): "S18",
    ("SS2", "MTH"): "S1", ("SS2", "VAT"): "S2", ("SS2", "AGR"): "S2",
    ("SS2", "CVC"): "S3", ("SS2", "CHM"): "S4", ("SS2", "GOV"): "S4",
    ("SS2", "COM"): "S4", ("SS2", "ENG"): "S5", ("SS2", "FSH"): "S6",
    ("SS2", "BIO"): "S7", ("SS2", "FDN"): "S8", ("SS2", "DAP"): "S8",
    ("SS2", "GMT"): "S9", ("SS2", "PHY"): "S10", ("SS2", "LIT"): "S10",
    ("SS2", "ACC"): "S10", ("SS2", "CPS"): "S11", ("SS2", "CRS"): "S12",
    ("SS2", "CHN"): "S13", ("SS2", "SGL"): "S14", ("SS2", "ECO"): "S15",
    ("SS2", "FTM"): "S16", ("SS2", "TDR"): "S17", ("SS2", "GEO"): "S17",
    ("SS2", "FRE"): "S17",
}

H = "HUMANITIES AND LANGUAGE 3RD TERM 2026/EXAMS"
SC = "SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM"
TE = "TECHNOLOGY C.A AND EXA 3RD TERM 26/TECHNOLOGY EXAM 3RD TERM 26"

# Values may contain an objective file followed by its separate theory file.
SOURCES = {
    ("JS1", "MTH"): ["SCIENCE DEPT2ND CA FOR 3RD TERM 26/EXAM/mr daniel exam 3rd term 26/Js1 Math exams 3rd term 26..pdf"],
    ("JS1", "SCS"): [f"{H}/MR AMANDE/JS1 SOCIAL & CITIZENSHIP STUD. 3RD TERM EXAMS 2026.docx"],
    ("JS1", "MUS"): [f"{H}/MR MUSIC/MUSIC JSS 1 EXAM.docx"],
    ("JS1", "ENG"): [f"{H}/MISS PRICY/JSS1 THIRD TERM ENGLISH STUDIES EXAMINATION 2026.pdf"],
    ("JS1", "DIT"): [f"{TE}/ICT JUNIOR EXAM 3RD TERM 26/JS1 DIGITAL TECHNOLOGY  3RD TERM EXAM  2026.txt"],
    ("JS1", "HIS"): [f"{H}/MR FRIDAY/JSS1 HISTORY 2026 EXAM.docx"],
    ("JS1", "INS"): ["intermediate EXAM JS1 3RD TERM 2025-2026.txt"],
    ("JS1", "CRS"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/js1 crs exam.docx"],
    ("JS1", "LIV"): [f"{SC}/mrs odey agric and livestock 3rd term 26/3rd term jss1 livestoc exam.txt"],
    ("JS1", "CCA"): ["jss1 cca examination 3rd term.txt"],
    ("JS1", "FRE"): [f"{H}/MR LOUIS BABEM/french jss1 examen 3rd.docx"],
    ("JS1", "BST"): [f"{H}/MRS ARINZE/j.s.1. 3rd term business studies exam 2026.txt"],
    ("JS1", "FAS"): [f"{TE}/GARMENT MAKING 3RD TERM 26/JSS 1 FASHION DESIGN AND GARMENT MAKING 3RD TERM EXAM.pdf"],
    ("JS1", "YOR"): [f"{H}/MRS ESTHER AMONA/2026 YORUBA THIRD TERM EXAM.docx"],
    ("JS1", "HAU"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/HAUSA JSS1.docx"],
    ("JS1", "PHE"): [f"{TE}/ABIGAIL PHE EXAM 3RD TER 26/JS1 EXAM 3RD TERM  26.txt"],
    ("JS1", "CHN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/NOTRE DAME CHINESE EXAM.docx"],
    ("JS1", "IGB"): ["Igbo civic for js1 and js2/JS1 IGBO.docx"],
    ("JS2", "MTH"): [f"{SC}/MRS SUSAN EXAM/3RD TERM MATHS EXAM JS2.pdf"],
    ("JS2", "SST"): [f"{H}/MR AMANDE/JS2 SOCIAL STUDIES EXAM 3RD TERM 2026.txt"],
    ("JS2", "MUS"): [f"{H}/MR MUSIC/MUSIC JSS 2 EXAM.docx"],
    ("JS2", "ENG"): [f"{H}/MISS PRICY/PRISCY'S 3RD TERM ENG EXAM2026.pdf"],
    ("JS2", "CSC"): [f"{TE}/ICT JUNIOR EXAM 3RD TERM 26/JS2  COMPUTER STUDIES 3RD TERM EXAM 2026.txt"],
    ("JS2", "HIS"): [f"{H}/MR FRIDAY/JSS2 HISTORY 2026 EXAM.docx"],
    ("JS2", "BTE"): ["SCIENCE DEPT2ND CA FOR 3RD TERM 26/THIRD TERM B.TEH JS2 2026.txt"],
    ("JS2", "CRS"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/js2 exam crs.docx"],
    ("JS2", "AGR"): [f"{SC}/mrs odey agric and livestock 3rd term 26/3RD TERM EXAMINATION JSS2   AGRICULTURAL SC.2026.pdf"],
    ("JS2", "BSC"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/JS2 BASIC SCIENCE.docx"],
    ("JS2", "CCA"): [f"{H}/MR AMANDE/Js2 CCA third term Exam 2026.txt"],
    ("JS2", "FRE"): [f"{H}/MR LOUIS BABEM/Jss 2 french examen 3rd.docx"],
    ("JS2", "BST"): [f"{H}/MRS ARINZE/j.s2 3rd  term  business studies  exam 2026.txt"],
    ("JS2", "HEC"): [f"{SC}/JSS 2 HOME ECONOMIC 3RD TERM EXAM.pdf"],
    ("JS2", "CVC"): ["Igbo civic for js1 and js2/JS2 CIVIC.docx"],
    ("JS2", "IGB"): ["Igbo civic for js1 and js2/JS2 IGBO.docx"],
    ("JS2", "YOR"): [f"{H}/MRS ESTHER AMONA/2026 YORUBA THIRD TERM EXAM.docx"],
    ("JS2", "HAU"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/HAUSA  JSS 2 THIRD TERM EXAM 2026.docx"],
    ("JS2", "PHE"): [f"{TE}/ABIGAIL PHE EXAM 3RD TER 26/JS2 EXAM 3RD TERM 26.txt"],
    ("JS2", "CHN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/NOTRE DAME CHINESE EXAM.docx"],
    ("SS1", "MTH"): [f"{SC}/MRS SUSAN EXAM/3RD TERM MATHS EXAM SS1.pdf"],
    ("SS1", "VAT"): [f"{H}/MR FABIAN/ss1 VA. EXAMS 3RD TERM.docx"],
    ("SS1", "AGR"): [f"{SC}/mrs odey agric and livestock 3rd term 26/3RD TERM AGRIC EXAM.txt"],
    ("SS1", "CHS"): [f"{H}/MR FRIDAY/SS1 EXAM 3RD TERM.docx"],
    ("SS1", "CHM"): ["SCIENCE DEPT2ND CA FOR 3RD TERM 26/SS1 CHEMISTRY EXAM 3rd trm 26.txt"],
    ("SS1", "GOV"): [f"{H}/MR ITODO/SS 1 GOVERNMENT EXAM.docx"],
    ("SS1", "COM"): [f"{H}/MRS ARINZE/s.s.1.3rd term commerce exam.txt"],
    ("SS1", "ENG"): [f"{H}/MRS OLADELE/2026 THIRD TERM ENGLISH EXAM SS1 - Copy.docx"],
    ("SS1", "BIO"): [f"{SC}/MRS UKAIRO EXAM 3RD TERM 26/SS1 3RD TERM EXAM BIOLOGY26.txt"],
    ("SS1", "FDN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/FOOD AND NUTRITION SS1 3RD TERM EXAM.docx"],
    ("SS1", "GMT"): [f"{TE}/GARMENT MAKING 3RD TERM 26/SS1 GARMENT MAKING 3RD TERM EXAM.pdf"],
    ("SS1", "PHY"): [f"{TE}/MR C HIBUKE EXAM 3RD TERM 26/SS 1 EXAM THIRD TERM PHYSICS.docx"],
    ("SS1", "LIT"): [f"{H}/MR SULE/MR SULE THIRD EXAM LIT SS 1.docx"],
    ("SS1", "ACC"): [f"{H}/MR BOD/FINANCIAL ACCOUNT SS1 EXAM.docx"],
    ("SS1", "DIT"): [f"{TE}/MR MARCEL EXAM 26 3RD TERM/SS1 DIGITAL TECHNOLOGY/SS1_Digital_Tech_Objectives.txt"],
    ("SS1", "CRS"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/ss1 crs exam.docx"],
    ("SS1", "CHN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/NOTRE DAME CHINESE EXAM.docx"],
    ("SS1", "ECO"): [f"{H}/MR BOD/ECONOMICS SS1 EXAM.docx"],
    ("SS1", "FTM"): [f"{SC}/mr daniel exam 3rd term 26/SS1 Fmaths exams 3rd term 26..pdf"],
    ("SS1", "GEO"): [f"{H}/MS MARY/GEOGRAPHY EXAMINATION SS 1 THIRD TERM 2026.docx"],
    ("SS1", "FRE"): [f"{H}/MR LOUIS BABEM/Ss1 French exam 3rd.docx"],
    ("SS1", "LIV"): [f"{SC}/mrs odey agric and livestock 3rd term 26/3rd exam ss1 livestock farming.txt"],
    ("SS2", "VAT"): [f"{H}/MR FABIAN/VISUAL ARTS EXAMINATION 3rd term ss2.docx"],
    ("SS2", "AGR"): [f"{SC}/mrs odey agric and livestock 3rd term 26/3RD TERMA  AGRIC SS2 EXAM.txt"],
    ("SS2", "CVC"): [f"{H}/MR FRIDAY/SS2 EXAM 3RD TERM.docx"],
    ("SS2", "CHM"): ["SCIENCE DEPT2ND CA FOR 3RD TERM 26/Chemistry_SS2_.txt"],
    ("SS2", "GOV"): [f"{H}/MR ITODO/SS 2 GOVERNMENT EXAM.docx"],
    ("SS2", "COM"): [f"{H}/MRS ARINZE/s.s.2 3rd term exam 2026.txt3.txt"],
    ("SS2", "ENG"): [f"{H}/MRS OLADELE/2026 THIRD TERM ENGLISH EXAM SS2 - Copy.docx"],
    ("SS2", "FSH"): [f"{SC}/MRS UKAIRO EXAM 3RD TERM 26/FISHERIES EXAM SS2 3RD TERM 26.txt"],
    ("SS2", "BIO"): ["SCIENCE DEPT2ND CA FOR 3RD TERM 26/SS2 BIOLOGY THIRD TERM EXAM 26.txt"],
    ("SS2", "FDN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/FOOD AND NUTRITION SS2 3RD TERM.docx"],
    ("SS2", "GMT"): [f"{TE}/GARMENT MAKING 3RD TERM 26/SS2 GARMENT MAKING 3RD EXAM.pdf"],
    ("SS2", "DAP"): [
        f"{TE}/MR MARCEL EXAM 26 3RD TERM/SS2 DATA PROCESSING/SS2_DataProcessing_Objectives.txt",
        f"{TE}/MR MARCEL EXAM 26 3RD TERM/SS2 DATA PROCESSING/SS2_DataProcessing_Theory.txt",
    ],
    ("SS2", "PHY"): [f"{TE}/MR C HIBUKE EXAM 3RD TERM 26/SS 2 EXAM THIRD TERM PHYSICS.docx"],
    ("SS2", "LIT"): [f"{H}/MR SULE/MR SULE THIRD EXAM LIT SS 2.docx"],
    ("SS2", "ACC"): [f"{H}/MR BOD/FINANCIAL ACCOUNT SS2 EXAM.docx"],
    ("SS2", "CPS"): [
        f"{TE}/MR MARCEL EXAM 26 3RD TERM/SS2 ICT (COMPUTER)/SS2_ICT_Objectives.txt",
        f"{TE}/MR MARCEL EXAM 26 3RD TERM/SS2 ICT (COMPUTER)/SS2_ICT_Theory.txt",
    ],
    ("SS2", "CRS"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/ss2 crs exam.docx"],
    ("SS2", "CHN"): ["TECHNOLOGY C.A AND EXA 3RD TERM 26/NOTRE DAME CHINESE EXAM.docx"],
    ("SS2", "ECO"): [f"{H}/MR BOD/ECONOMICS SS2 EXAM.docx"],
    ("SS2", "TDR"): [f"{TE}/MR C HIBUKE EXAM 3RD TERM 26/SS 2 TECH. DRAWING EXAM.docx"],
    ("SS2", "GEO"): [f"{H}/MS MARY/GEOGRAPHY EXAM SS 2 FOR THIRD TERM 2026.docx"],
    ("SS2", "FRE"): [f"{H}/MR LOUIS BABEM/Ss2 french exam 3rd.docx"],
}


def configured_base():
    base = load_base_importer()
    original_normalize_rows = base.normalize_rows

    def normalize_exam_rows(parsed_rows, extracted_text):
        rows, objective_count, theory_count = original_normalize_rows(parsed_rows, extracted_text)
        objective_rows = [
            row for row in rows
            if (row.get("question_type") or "OBJECTIVE").upper() == "OBJECTIVE"
        ]
        theory_rows = [
            row for row in rows
            if (row.get("question_type") or "OBJECTIVE").upper() != "OBJECTIVE"
        ]
        if theory_rows:
            theory_text = "\n\n".join(
                str(row.get("stem") or "").strip()
                for row in theory_rows
                if str(row.get("stem") or "").strip()
            )
            theory_rows = [
                {
                    "question_type": "THEORY",
                    "stem": theory_text,
                    "model_answer": "",
                }
            ] if theory_text else []
        return objective_rows + theory_rows, objective_count, len(theory_rows)

    base.normalize_rows = normalize_exam_rows
    base.IMPORT_TAG = IMPORT_TAG
    base.CBTExamType = SimpleNamespace(CA=CBTExamType.EXAM)
    base.CBTWritebackTarget = SimpleNamespace(
        CA1=CBTWritebackTarget.OBJECTIVE,
        NONE=CBTWritebackTarget.THEORY,
    )
    base.title_for = lambda assignment, _slot: (
        f"{assignment.academic_class.code} {assignment.subject.name} 2025/2026 Third Term Examination"
    )
    base._apply_exam_row_marks = lambda *, exam, objective_total, theory_total: apply_exam_row_marks(
        exam=exam,
        objective_total="20.00",
        theory_total="30.00",
    )
    def parse_igbo_docx(path):
        paragraphs = [
            paragraph.text.strip()
            for paragraph in Document(path).paragraphs
            if paragraph.text.strip()
        ]
        theory_index = next(
            (
                index
                for index, text in enumerate(paragraphs)
                if "THEORY" in text.upper()
                or re.search(r"(?i)\b(?:AKWỤKWỌ|AGBA)\s+B\b", text)
            ),
            len(paragraphs),
        )
        objective_paragraphs = paragraphs[:theory_index]
        theory_paragraphs = paragraphs[theory_index:]
        answer_re = re.compile(r"(?i)^\s*Az[iị]za\s*:\s*([A-D])(?:\b|[\s\.:;\-)])")
        option_re = re.compile(r"(?i)^\s*([A-D])\s*[\).:\-]\s*(.+?)\s*$")
        header_re = re.compile(
            r"(?i)^(?:ULE\s+TAAM|ASỤSỤ\s+IGBO|3RD\s+TERM|THIRD\s+TERM)"
        )
        chunks = []
        current = []
        for paragraph in objective_paragraphs:
            answer_match = answer_re.match(paragraph)
            if answer_match:
                chunks.append((current, answer_match.group(1).upper()))
                current = []
            else:
                current.append(paragraph)

        objective_rows = []
        for chunk, correct_label in chunks:
            lines = []
            for paragraph in chunk:
                lines.extend(
                    line.strip()
                    for line in paragraph.splitlines()
                    if line.strip()
                )
            lines = [line for line in lines if not header_re.match(line)]
            labelled = []
            first_option_index = None
            for index, line in enumerate(lines):
                match = option_re.match(line)
                if match:
                    if first_option_index is None:
                        first_option_index = index
                    labelled.append((match.group(1).upper(), match.group(2).strip()))
            if len(labelled) >= 4 and {label for label, _text in labelled[:4]} == set("ABCD"):
                stem_lines = lines[:first_option_index]
                options = dict(labelled[:4])
            elif len(lines) >= 5:
                stem_lines = lines[:-4]
                options = dict(zip("ABCD", lines[-4:]))
            else:
                continue
            stem = re.sub(r"\s+", " ", " ".join(stem_lines)).strip()
            if not stem or not all(str(options.get(label) or "").strip() for label in "ABCD"):
                continue
            objective_rows.append(
                {
                    "question_type": "OBJECTIVE",
                    "stem": stem,
                    "options": {
                        label: re.sub(r"\s+", " ", str(options[label])).strip()
                        for label in "ABCD"
                    },
                    "correct_label": correct_label,
                }
            )

        theory_text = "\n\n".join(theory_paragraphs).strip()
        theory_rows = (
            [{"question_type": "THEORY", "stem": theory_text, "model_answer": ""}]
            if theory_text
            else []
        )
        extracted_lines = []
        for paragraph in paragraphs:
            answer_match = answer_re.match(paragraph)
            extracted_lines.append(
                f"Answer: {answer_match.group(1).upper()}"
                if answer_match
                else paragraph
            )
        extracted_text = "\n".join(extracted_lines)
        return {
            "parsed_rows": objective_rows + theory_rows,
            "extracted_text": extracted_text,
            "parser_used": "deterministic_igbo_docx_blocks",
            "instructions": [],
        }

    def parse_yoruba_docx(path, class_code):
        paragraphs = [
            paragraph.text.strip()
            for paragraph in Document(path).paragraphs
            if paragraph.text.strip()
        ]
        class_re = re.compile(r"(?i)^CLASS\s*:\s*JS\s*([12])\b")
        class_matches = [
            (index, class_re.match(text))
            for index, text in enumerate(paragraphs)
            if class_re.match(text)
        ]
        wanted_level = class_code[-1]
        start = next(
            (index for index, match in class_matches if match.group(1) == wanted_level),
            None,
        )
        if start is None:
            raise RuntimeError(f"{class_code} Yoruba section was not found.")
        end = next(
            (index for index, _match in class_matches if index > start),
            len(paragraphs),
        )
        section = paragraphs[max(0, start - 2):end]
        theory_index = next(
            (
                index
                for index, text in enumerate(section)
                if "THEORY" in text.upper() or re.search(r"(?i)^IPIN\s+KEJI\b", text)
            ),
            len(section),
        )
        objective_paragraphs = section[:theory_index]
        theory_paragraphs = section[theory_index:]
        answer_re = re.compile(r"(?i)^\s*Answer\s*:\s*([A-D])(?:\b|[\s\.:;\-)])")
        option_re = re.compile(
            r"(?i)^\s*(?:\(([A-D])\)|([A-D])[\).:\-]|([A-D])\s+)\s*(.+?)\s*$"
        )
        header_re = re.compile(
            r"(?i)^(?:THIRD\s+TERM|3RD\s+TERM|SUBJECT\s*:|CLASS\s*:|INSTRUCTION\s*:)"
        )
        chunks = []
        current = []
        for paragraph in objective_paragraphs:
            answer_match = answer_re.match(paragraph)
            if answer_match:
                chunks.append((current, answer_match.group(1).upper()))
                current = []
            else:
                current.append(paragraph)

        objective_rows = []
        for source_number, (chunk, correct_label) in enumerate(chunks, start=1):
            lines = []
            for paragraph in chunk:
                lines.extend(
                    line.strip()
                    for line in paragraph.splitlines()
                    if line.strip()
                )
            lines = [line for line in lines if not header_re.match(line)]
            labelled = []
            first_option_index = None
            for index, line in enumerate(lines):
                match = option_re.match(line)
                if not match:
                    continue
                label = (match.group(1) or match.group(2) or match.group(3) or "").upper()
                if label not in "ABCD":
                    continue
                if first_option_index is None:
                    first_option_index = index
                labelled.append((label, match.group(4).strip()))
            if len(labelled) >= 4 and {label for label, _text in labelled[-4:]} == set("ABCD"):
                first_option_index = next(
                    index
                    for index, line in enumerate(lines)
                    if option_re.match(line)
                )
                stem_lines = lines[:first_option_index]
                options = dict(labelled[-4:])
            elif len(lines) >= 5:
                stem_lines = lines[:-4]
                options = dict(zip("ABCD", lines[-4:]))
            else:
                continue
            stem = re.sub(r"\s+", " ", " ".join(stem_lines)).strip()
            if not stem or not all(str(options.get(label) or "").strip() for label in "ABCD"):
                continue
            objective_rows.append(
                {
                    "question_type": "OBJECTIVE",
                    "stem": stem,
                    "options": {
                        label: re.sub(r"\s+", " ", str(options[label])).strip()
                        for label in "ABCD"
                    },
                    "correct_label": correct_label,
                    "source_number": source_number,
                }
            )

        theory_text = "\n\n".join(theory_paragraphs).strip()
        theory_rows = (
            [{"question_type": "THEORY", "stem": theory_text, "model_answer": ""}]
            if theory_text
            else []
        )
        canonical_lines = []
        for number, row in enumerate(objective_rows, start=1):
            canonical_lines.append(f"{number}. {row['stem']}")
            canonical_lines.extend(
                f"{label}. {row['options'][label]}" for label in "ABCD"
            )
            canonical_lines.append(f"Answer: {row['correct_label']}")
        if theory_text:
            canonical_lines.extend(["SECTION B: THEORY", theory_text])
        return {
            "parsed_rows": objective_rows + theory_rows,
            "extracted_text": "\n".join(canonical_lines),
            "parser_used": "deterministic_yoruba_class_blocks",
            "instructions": [],
        }

    def local_rows(path, subject_name):
        if path.suffix.lower() == ".docx" and "IGBO" in subject_name.upper():
            return parse_igbo_docx(path)
        if path.suffix.lower() == ".docx" and "YORUBA" in subject_name.upper():
            class_code = path.name.split("-", 1)[0].upper()
            return parse_yoruba_docx(path, class_code)
        direct = {**base.direct_parser_info(path), "parser_used": "deterministic_local"}
        if path.suffix.lower() != ".docx":
            return direct
        try:
            structured = parse_structured_docx_rows(path)
        except Exception:
            return direct

        def quality(payload):
            rows, objective_count, theory_count = base.normalize_rows(
                payload.get("parsed_rows") or [],
                payload.get("extracted_text") or "",
            )
            answered = sum(
                1 for row in rows
                if (row.get("question_type") or "").upper() == "OBJECTIVE"
                and (row.get("correct_label") or "").strip()
            )
            return answered == objective_count and objective_count > 0, answered, objective_count, min(theory_count, 10)

        if quality(structured) > quality(direct):
            structured["parser_used"] = "deterministic_docx_structural"
            return structured
        return direct

    base.choose_rows_for_path = local_rows
    return base


def combined_source(base, paths, temp_dir: Path, key):
    if key[1] == "YOR" and len(paths) == 1:
        target = temp_dir / f"{key[0]}-YOR-third-term-exam.docx"
        shutil.copy2(paths[0], target)
        return target
    if len(paths) == 1:
        return paths[0]
    text = []
    for path in paths:
        text.append(base.direct_parser_info(path).get("extracted_text") or "")
    target = temp_dir / f"{key[0]}-{key[1]}-third-term-exam.txt"
    target.write_text("\n\n".join(text), encoding="utf-8")
    return target


def attach_supplied_diagrams(exam, key):
    """Attach diagrams that are embedded/drawn in the supplied Yoruba paper."""
    if key != ("JS1", "YOR"):
        return 0
    diagram_dir = ROOT / "assets" / "third-term-diagrams"
    assets = [
        diagram_dir / f"js1-yoruba-q{number}.png"
        for number in range(37, 41)
    ]
    if not all(path.is_file() and path.stat().st_size for path in assets):
        missing = [path.name for path in assets if not path.is_file() or not path.stat().st_size]
        raise RuntimeError(f"Yoruba diagram export missing: {', '.join(missing)}")
    links = list(
        exam.exam_questions.select_related("question")
        .filter(question__question_type="OBJECTIVE")
        .order_by("sort_order")
    )
    if len(links) < 40:
        raise RuntimeError(f"JS1 Yoruba expected 40 objective rows, found {len(links)}")
    for number, asset in zip(range(37, 41), assets):
        question = links[number - 1].question
        with asset.open("rb") as handle:
            question.stimulus_image.save(asset.name, File(handle), save=False)
        question.stimulus_caption = f"Diagram supplied for question {number}"
        question.save(update_fields=["stimulus_image", "stimulus_caption", "updated_at"])
    return len(assets)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--source-root", default=str(SOURCE_ROOT))
    parser.add_argument(
        "--only",
        action="append",
        default=[],
        help="Limit to CLASS:SUBJECT, for example --only JS1:PHE. Repeat as needed.",
    )
    args = parser.parse_args()
    only_keys = set()
    for value in args.only:
        parts = [part.strip().upper() for part in str(value).split(":", 1)]
        if len(parts) != 2 or not all(parts):
            raise SystemExit(f"Invalid --only value: {value!r}")
        only_keys.add(tuple(parts))

    source_root = Path(args.source_root)
    session = AcademicSession.objects.get(name="2025/2026")
    term = Term.objects.get(session=session, name="THIRD")
    assignments = {
        (row.academic_class.code, row.subject.code): row
        for row in TeacherSubjectAssignment.objects.filter(session=session, term=term, is_active=True)
        .select_related("teacher", "subject", "academic_class")
    }
    base = configured_base()
    ready = []
    blocked = []
    imported = []

    with tempfile.TemporaryDirectory(prefix="ndga-third-exam-") as folder:
        temp_dir = Path(folder)
        for key, slot_code in EXPECTED.items():
            if only_keys and key not in only_keys:
                continue
            assignment = assignments.get(key)
            rel_paths = SOURCES.get(key)
            if not assignment:
                blocked.append((*key, "teacher assignment missing"))
                continue
            if not rel_paths:
                blocked.append((*key, "paper not supplied"))
                continue
            paths = [source_root / value for value in rel_paths]
            missing = [str(path.relative_to(source_root)) for path in paths if not path.is_file()]
            if missing:
                blocked.append((*key, f"source missing: {', '.join(missing)}"))
                continue
            empty = [path.name for path in paths if path.stat().st_size == 0]
            if empty:
                blocked.append((*key, f"empty source file: {', '.join(empty)}"))
                continue
            try:
                source_path = combined_source(base, paths, temp_dir, key)
                info = base.choose_rows_for_path(source_path, assignment.subject.name)
                rows, objective_count, theory_count = base.normalize_rows(
                    info.get("parsed_rows") or [],
                    info.get("extracted_text") or "",
                )
                answered = sum(
                    1 for row in rows
                    if (row.get("question_type") or "").upper() == "OBJECTIVE"
                    and (row.get("correct_label") or "").strip()
                )
                if objective_count <= 0:
                    blocked.append((*key, "no objective questions parsed"))
                    continue
                if answered != objective_count:
                    blocked.append((*key, f"answer key incomplete ({answered}/{objective_count})"))
                    continue
                answer_markers = len(list(base.ANSWER_LABEL_RE.finditer(info.get("extracted_text") or "")))
                if answer_markers and answer_markers != objective_count:
                    blocked.append(
                        (*key, f"source/parser question mismatch ({objective_count} parsed, {answer_markers} answer markers)")
                    )
                    continue
                if theory_count <= 0 and key[1] != "CHN":
                    blocked.append((*key, "no theory questions parsed"))
                    continue
                ready.append((*key, objective_count, theory_count, info.get("parser_used") or ""))
                if args.apply:
                    date_text, slot_label = SLOTS[slot_code]
                    result = base.import_exam(
                        source_path=source_path,
                        assignment=assignment,
                        it_user=base.User.objects.get(username="admin@ndgakuje.org"),
                        dean_user=base.User.objects.filter(username="emmanuel@ndgakuje.org").first()
                        or base.User.objects.filter(primary_role__code="DEAN").first(),
                        date_text=date_text,
                        slot_label=slot_label,
                        dry_run=False,
                    )
                    exam = Exam.objects.get(pk=result["exam_id"])
                    exam.description = f"{IMPORT_TAG}: Imported from {', '.join(rel_paths)}."
                    exam.save(update_fields=["description", "updated_at"])
                    base.Question.objects.filter(exam_links__exam=exam).update(topic="Third Term Examination")
                    attach_supplied_diagrams(exam, key)
                    blueprint = exam.blueprint
                    config = dict(blueprint.section_config or {})
                    config.update(
                        {
                            "flow_type": "OBJECTIVE_THEORY" if theory_count else "OBJECTIVE_ONLY",
                            "objective_target_max": "20.00",
                            "theory_target_max": "30.00" if theory_count else "0.00",
                            "theory_response_mode": "PAPER" if theory_count else "NONE",
                            "manual_score_split": bool(theory_count),
                        }
                    )
                    config.pop("ca_target", None)
                    blueprint.duration_minutes = 90
                    blueprint.theory_enabled = bool(theory_count)
                    blueprint.section_config = config
                    blueprint.objective_writeback_target = CBTWritebackTarget.OBJECTIVE
                    blueprint.theory_writeback_target = CBTWritebackTarget.THEORY
                    blueprint.auto_show_result_on_submit = False
                    blueprint.save(
                        update_fields=[
                            "duration_minutes",
                            "theory_enabled",
                            "section_config",
                            "objective_writeback_target",
                            "theory_writeback_target",
                            "auto_show_result_on_submit",
                            "updated_at",
                        ]
                    )
                    imported.append((key, exam.id))
            except Exception as exc:
                blocked.append((*key, f"parse/import error: {exc}"))

    print(f"MODE={'APPLY' if args.apply else 'AUDIT'} READY={len(ready)} BLOCKED={len(blocked)} IMPORTED={len(imported)}")
    for class_code, subject_code, objective_count, theory_count, parser_name in ready:
        print(f"READY | {class_code} | {subject_code} | objective={objective_count} theory={theory_count} | {parser_name}")
    for class_code, subject_code, reason in blocked:
        print(f"BLOCKED | {class_code} | {subject_code} | {reason}")
    for (class_code, subject_code), exam_id in imported:
        print(f"IMPORTED | {class_code} | {subject_code} | exam={exam_id}")


if __name__ == "__main__":
    main()
