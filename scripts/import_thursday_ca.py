import argparse
import os
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from html import escape
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

os.environ.setdefault("DJANGO_SETTINGS_MODULE", os.getenv("DJANGO_SETTINGS_MODULE", "core.settings.local"))

import django

django.setup()

from django.core.files import File
from django.db import transaction
from django.utils import timezone
from docx import Document

from apps.accounts.models import User
from apps.academics.models import AcademicSession, TeacherSubjectAssignment, Term
from apps.cbt.models import (
    CBTDocumentStatus,
    CBTExamType,
    CBTQuestionDifficulty,
    CBTQuestionType,
    CBTWritebackTarget,
    CorrectAnswer,
    Exam,
    ExamDocumentImport,
    ExamQuestion,
    Option,
    Question,
    QuestionBank,
)
from apps.cbt.services import (
    _apply_exam_row_marks,
    _objective_marker_count,
    _parsed_rows_quality,
    _parse_questions_with_openai,
    ensure_default_blueprint,
    extract_text_from_document,
    parse_question_document,
)
from apps.cbt.workflow import dean_approve_exam, submit_exam_to_dean

DEFAULT_SOURCE_DIR = "Thursday CA"
DEFAULT_SESSION = "2025/2026"
DEFAULT_TERM = "SECOND"
DEFAULT_ACTUAL_DATE = "2026-03-12"

EXAM_DAY_PREFIX = {
    "2026-03-12": "THU",
    "2026-03-13": "FRI",
    "2026-03-14": "SAT",
}

SUBJECT_ALIAS_MAP = {
    "MATHS": "MATHEMATICS",
    "MATHEMATICS": "MATHEMATICS",
    "ENGLISH": "ENGLISH LANGUAGE",
    "ENGLISHSTUDIES": "ENGLISH LANGUAGE",
    "ICT": "DIGITAL TECHNOLOGY",
    "DIGITALTECHNOLOGY": "DIGITAL TECHNOLOGY",
    "COMPUTER": "COMPUTER SCIENCE",
    "COMPUTERSCIENCE": "COMPUTER SCIENCE",
    "COMPUTERSTUDIES": "COMPUTER STUDIES",
    "CONPUTER": "COMPUTER SCIENCE",
    "BASICTECH": "BASIC TECHNOLOGY",
    "BASICTECHNOLOGY": "BASIC TECHNOLOGY",
    "GOVT": "GOVERNMENT",
    "GOVERNMENT": "GOVERNMENT",
    "FOODNUT": "FOOD AND NUTRITION",
    "FOODANDNUTRITION": "FOOD AND NUTRITION",
    "FINANCIALACCOUNT": "ACCOUNTING",
    "ACCOUNT": "ACCOUNTING",
    "ACCOUNTING": "ACCOUNTING",
    "VISUALART": "VISUAL ART",
    "SOCIALSTUDIES": "SOCIAL STUDIES",
    "SOCIALCITIZENSHIPSTUDIES": "SOCIAL AND CITIZENSHIP STUDIES",
    "LIVESTOCKFARMING": "LIVESTOCK",
    "LIVESTOCK": "LIVESTOCK",
    "AGRIC": "AGRICULTURAL SCIENCE",
    "AGRICULTURALSCIENCE": "AGRICULTURAL SCIENCE",
    "CIVICEDUCATION": "CIVIC EDUCATION",
    "CIVICEDUC": "CIVIC EDUCATION",
    "CCA": "CCA",
    "YORUBA": "YORUBA LANGUAGE",
    "HAUSA": "HAUSA LANGUAGE",
    "FRENCH": "FRENCH",
    "MUSIC": "MUSIC",
    "PHYSICS": "PHYSICS",
    "CHEMISTRY": "CHEMISTRY",
    "COMMERCE": "COMMERCE",
    "FASHION": "FASHION",
    "GARMENT": "GARMENT MAKING THEORY",
}

CLASS_SPECIFIC_ALIASES = {
    ("JS1", "GARMENT"): "FASHION",
    ("JS2", "COMPUTER"): "COMPUTER SCIENCE",
    ("JS2", "COMPUTERSTUDIES"): "COMPUTER SCIENCE",
    ("JS2", "CONPUTER"): "COMPUTER SCIENCE",
    ("JS3", "COMPUTER"): "COMPUTER SCIENCE",
    ("JS3", "COMPUTERSTUDIES"): "COMPUTER SCIENCE",
    ("JS3", "CONPUTER"): "COMPUTER SCIENCE",
    ("JS3", "BASICTECH"): "BASIC TECHNOLOGY",
    ("JS3", "BASICTECHNOLOGY"): "BASIC TECHNOLOGY",
    ("SS2", "COMPUTER"): "COMPUTER STUDIES",
    ("SS2", "COMPUTERSTUDIES"): "COMPUTER STUDIES",
}

SAFE_TEXT_REPLACEMENTS = {
    ord("\u2080"): "0",
    ord("\u2081"): "1",
    ord("\u2082"): "2",
    ord("\u2083"): "3",
    ord("\u2084"): "4",
    ord("\u2085"): "5",
    ord("\u2086"): "6",
    ord("\u2087"): "7",
    ord("\u2088"): "8",
    ord("\u2089"): "9",
    ord("\u2070"): "0",
    ord("\u00b9"): "1",
    ord("\u00b2"): "2",
    ord("\u00b3"): "3",
    ord("\u2074"): "4",
    ord("\u2075"): "5",
    ord("\u2076"): "6",
    ord("\u2077"): "7",
    ord("\u2078"): "8",
    ord("\u2079"): "9",
    ord("\u03c0"): "pi",
    ord("\u2212"): "-",
    ord("\u2010"): "-",
    ord("\u2011"): "-",
    ord("\u2012"): "-",
    ord("\u2013"): "-",
    ord("\u2014"): "-",
    ord("\u2022"): "-",
    ord("\uf0b7"): "-",
    ord("\u20a6"): "N",
    ord("\u00a0"): " ",
}

ANSWER_LINE_RE = re.compile(
    r"^\s*(?:answer|ans)\s*[:\-]?\s*([A-Da-d])(?:\b|[\).\:-])",
    re.IGNORECASE,
)
OPTION_LINE_RE = re.compile(r"^\s*([A-Da-d])[\)\.\:-]\s*(.+)$")
OPTION_ONLY_RE = re.compile(r"^\s*(.+)$")
SINGLE_ANSWER_KEY_RE = re.compile(r"^\s*([A-Da-d])\s*$")
SECTION_RE = re.compile(r"^\s*SECTION\s+([A-Z0-9]+)", re.IGNORECASE)
THEORY_HEADER_RE = re.compile(
    r"^\s*(?:SECTION\s+[A-Z0-9]+\s*[:\-]?\s*THEORY|THEORY(?:\s+QUESTIONS?)?)\b",
    re.IGNORECASE,
)
OBJECTIVE_HEADER_RE = re.compile(
    r"^\s*(?:OBJECTIVE(?:S|\s+QUESTIONS?)?|COMPREHENSION\s+QUESTIONS)\b",
    re.IGNORECASE,
)
INSTRUCTION_RE = re.compile(
    r"^\s*(?:INSTRUCTION|INSTRUCTIONS|ANSWER ALL THE QUESTIONS|READ THE PASSAGE|READ THE TEXTE|OBJECTIVES?)\b",
    re.IGNORECASE,
)
PASSAGE_RE = re.compile(r"^\s*(?:TEXTE|PASSAGE)\b", re.IGNORECASE)
THEORY_NUMBER_RE = re.compile(r"^\s*\d+[A-Za-z]?[\.\)]")
INLINE_OPTION_MARKER_RE = re.compile(r"(?i)\b([A-D])\s*[\)\.\:-]\s*")


@dataclass
class RichPart:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False

    def render(self, value=None):
        html = escape(str(self.text if value is None else value))
        html = html.replace("\n", "<br>")
        if self.bold:
            html = f"<strong>{html}</strong>"
        if self.italic:
            html = f"<em>{html}</em>"
        if self.underline:
            html = f"<u>{html}</u>"
        return html


@dataclass
class ParagraphRow:
    text: str
    parts: list[RichPart]

    @property
    def html(self):
        return "".join(part.render() for part in self.parts)

    def slice_html(self, start=0, end=None):
        if end is None:
            end = len(self.text)
        if start < 0:
            start = 0
        if end < start:
            return ""
        cursor = 0
        chunks = []
        for part in self.parts:
            part_text = part.text
            next_cursor = cursor + len(part_text)
            if next_cursor <= start:
                cursor = next_cursor
                continue
            if cursor >= end:
                break
            local_start = max(start - cursor, 0)
            local_end = min(end - cursor, len(part_text))
            if local_end > local_start:
                chunks.append(part.render(part_text[local_start:local_end]))
            cursor = next_cursor
        return "".join(chunks).strip()


def collapse(value):
    return " ".join(str(value or "").split())


def normalize(value):
    return re.sub(r"[^A-Z0-9]+", "", collapse(value).upper().replace("&", " AND "))


def safe_text(value):
    return str(value or "").translate(SAFE_TEXT_REPLACEMENTS)


def rich_text(value):
    return escape(safe_text(value)).replace("\n", "<br>")


def paragraph_rows_from_doc(doc):
    rows = []
    for paragraph in doc.paragraphs:
        parts = []
        for run in paragraph.runs:
            if not run.text:
                continue
            parts.append(
                RichPart(
                    text=safe_text(run.text),
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                    underline=bool(run.underline),
                )
            )
        if not parts:
            text = safe_text(paragraph.text)
            if text.strip():
                parts.append(RichPart(text=text))
        text = safe_text("".join(part.text for part in parts))
        if text.strip():
            rows.append(ParagraphRow(text=text, parts=parts))
    return rows


def _normalize_option_text(value):
    return safe_text(collapse(value)).strip(" -")


def _normalize_stem_text(value):
    return safe_text(collapse(value)).strip()


def _append_instruction(instructions, value):
    cleaned = _normalize_stem_text(value)
    if not cleaned:
        return
    if cleaned not in instructions:
        instructions.append(cleaned)


def _resolve_correct_label(options, answer_raw):
    answer_text = _normalize_option_text(answer_raw)
    if not answer_text:
        return ""
    match = re.match(r"^\s*([A-Da-d])(?:\b|[\).\:-])?", answer_text)
    if match:
        return match.group(1).upper()
    normalized_answer = collapse(answer_text).lower()
    for label, option_text in options.items():
        candidate = collapse(option_text).lower()
        if candidate and normalized_answer == candidate:
            return label
    for label, option_text in options.items():
        candidate = collapse(option_text).lower()
        if candidate and normalized_answer in candidate:
            return label
    return ""


def _extract_inline_objective_row(paragraph):
    value = safe_text(paragraph.text)
    if not value.strip():
        return None

    answer_raw = ""
    answer_match = re.search(
        r"(?is)\b(?:answer|ans)\s*[:\-]?\s*(.+)$",
        value,
    )
    if answer_match:
        answer_raw = safe_text(answer_match.group(1))
        value = value[: answer_match.start()].strip()

    markers = list(INLINE_OPTION_MARKER_RE.finditer(value))
    if len(markers) < 4:
        return None

    first_seen = {}
    for marker in markers:
        label = marker.group(1).upper()
        if label not in first_seen:
            first_seen[label] = marker
    if not all(label in first_seen for label in ("A", "B", "C", "D")):
        return None

    ordered = [first_seen[label] for label in ("A", "B", "C", "D")]
    if not (ordered[0].start() < ordered[1].start() < ordered[2].start() < ordered[3].start()):
        return None

    stem = _normalize_stem_text(value[: ordered[0].start()].strip(" :-"))
    if not stem:
        return None

    options = {}
    for index, marker in enumerate(ordered):
        label = marker.group(1).upper()
        start = marker.end()
        end = ordered[index + 1].start() if index + 1 < len(ordered) else len(value)
        option_text = _normalize_option_text(value[start:end])
        if not option_text:
            return None
        options[label] = option_text

    return {
        "question_type": "OBJECTIVE",
        "stem": stem,
        "rich_stem": paragraph.slice_html(0, ordered[0].start()) or rich_text(stem),
        "options": options,
        "correct_label": _resolve_correct_label(options, answer_raw),
    }


def parse_structured_docx_rows(path):
    doc = Document(path)
    rows = paragraph_rows_from_doc(doc)
    parsed_rows = []
    instructions = []
    current_section = ""
    mode = "OBJECTIVE"
    current = None
    answer_key_mode = False
    answer_key_values = []

    def finalize_current():
        nonlocal current
        if not current:
            return
        stem = _normalize_stem_text(current.get("stem"))
        rich_stem_value = (current.get("rich_stem") or "").strip() or rich_text(stem)
        if current.get("question_type") == "OBJECTIVE":
            options = current.get("options") or {}
            if len(options) >= 4 and stem:
                ordered = {label: _normalize_option_text(options.get(label, "")) for label in ("A", "B", "C", "D")}
                parsed_rows.append(
                    {
                        "question_type": "OBJECTIVE",
                        "stem": stem,
                        "rich_stem": rich_stem_value,
                        "options": ordered,
                        "correct_label": (current.get("correct_label") or "").strip().upper(),
                        "section_label": current.get("section_label") or current_section,
                    }
                )
        elif stem:
            note = "\n".join(_normalize_stem_text(value) for value in current.get("model_answer_lines") or [] if _normalize_stem_text(value))
            parsed_rows.append(
                {
                    "question_type": "THEORY",
                    "stem": stem,
                    "rich_stem": rich_stem_value,
                    "model_answer": safe_text(note),
                    "section_label": current.get("section_label") or current_section,
                }
            )
        current = None

    for paragraph in rows:
        line = _normalize_stem_text(paragraph.text)
        if not line:
            continue

        if answer_key_mode:
            answer_key_match = SINGLE_ANSWER_KEY_RE.match(line)
            if answer_key_match:
                answer_key_values.append(answer_key_match.group(1).upper())
                continue
            answer_key_mode = False

        if line.upper() == "ANSWER":
            finalize_current()
            answer_key_mode = True
            continue

        section_match = SECTION_RE.match(line)
        if section_match:
            finalize_current()
            current_section = safe_text(line)
            if THEORY_HEADER_RE.match(line):
                mode = "THEORY"
            elif "SECTION" in line.upper():
                mode = "OBJECTIVE"
            _append_instruction(instructions, line)
            continue

        if THEORY_HEADER_RE.match(line):
            finalize_current()
            mode = "THEORY"
            current_section = safe_text(line)
            _append_instruction(instructions, line)
            continue

        if OBJECTIVE_HEADER_RE.match(line):
            finalize_current()
            mode = "OBJECTIVE"
            current_section = safe_text(line)
            _append_instruction(instructions, line)
            continue

        if "QUESTION" in line.upper() and len(line.split()) <= 4:
            finalize_current()
            mode = "OBJECTIVE"
            current_section = safe_text(line)
            _append_instruction(instructions, line)
            continue

        if INSTRUCTION_RE.match(line) or PASSAGE_RE.match(line):
            finalize_current()
            _append_instruction(instructions, line)
            continue

        if len(line) > 120 and not line.endswith("?") and current is None and not parsed_rows:
            _append_instruction(instructions, line)
            continue

        answer_match = ANSWER_LINE_RE.match(line)
        if answer_match:
            if current and current.get("question_type") == "OBJECTIVE":
                current["correct_label"] = answer_match.group(1).upper()
                finalize_current()
            else:
                answer_key_values.append(answer_match.group(1).upper())
            continue

        if mode == "THEORY":
            if current and THEORY_NUMBER_RE.match(line):
                finalize_current()
            if current is None:
                current = {
                    "question_type": "THEORY",
                    "stem": line,
                    "rich_stem": paragraph.html or rich_text(line),
                    "model_answer_lines": [],
                    "section_label": current_section,
                }
            else:
                current["stem"] = f"{current['stem']}\n{line}".strip()
                current["rich_stem"] = f"{current['rich_stem']}<br>{paragraph.html or rich_text(line)}".strip()
            continue

        inline_row = _extract_inline_objective_row(paragraph)
        if inline_row is not None:
            finalize_current()
            inline_row["section_label"] = current_section
            parsed_rows.append(inline_row)
            continue

        option_match = OPTION_LINE_RE.match(line)
        if option_match and current and current.get("question_type") == "OBJECTIVE":
            current["options"][option_match.group(1).upper()] = _normalize_option_text(option_match.group(2))
            if len(current["options"]) >= 4 and current.get("correct_label"):
                finalize_current()
            continue

        if current and current.get("question_type") == "OBJECTIVE" and len(current.get("options") or {}) < 4:
            next_label = ("A", "B", "C", "D")[len(current["options"])]
            current["options"][next_label] = _normalize_option_text(line)
            continue

        finalize_current()
        current = {
            "question_type": "OBJECTIVE",
            "stem": line,
            "rich_stem": paragraph.html or rich_text(line),
            "options": {},
            "correct_label": "",
            "section_label": current_section,
        }

    finalize_current()

    unanswered_rows = [
        row for row in parsed_rows
        if row.get("question_type") == "OBJECTIVE" and not row.get("correct_label")
    ]
    for row, correct_label in zip(unanswered_rows, answer_key_values):
        row["correct_label"] = correct_label

    return {
        "parsed_rows": parsed_rows,
        "instructions": instructions,
        "parser_used": "docx_structural",
    }


def instruction_text_for_exam(*, instructions, flow_type):
    rows = []
    seen = set()
    for entry in instructions or []:
        cleaned = safe_text(collapse(entry))
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        rows.append(cleaned)
    if flow_type in {"OBJECTIVE_THEORY", "THEORY_ONLY"}:
        theory_note = "Write theory answers in the theory response format shown on the exam screen."
        if theory_note.lower() not in seen:
            rows.append(theory_note)
    if not rows:
        return "Answer all questions carefully."
    return "\n".join(rows)


def display_path(value):
    path = Path(value).resolve()
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(value)


def class_code_from_value(value):
    cleaned = collapse(value).upper().replace("JSS", "JS").replace("SSS", "SS")
    match = re.search(r"(JS|SS)\s*([123])", cleaned)
    if not match:
        return normalize(cleaned)
    return f"{match.group(1)}{match.group(2)}"


def normalize_time_text(value):
    text = safe_text(collapse(value)).replace(";", ":")
    text = re.sub(r"\s*:\s*", ":", text)
    text = re.sub(r"\s*-\s*", "-", text)
    text = re.sub(r"(?<=\d)\s+(?=\d)", ":", text)
    return text


def parse_time_window(label):
    text = normalize_time_text(label)
    match = re.search(r"(\d{1,2}:\d{2})-(\d{1,2}:\d{2})", text)
    if not match:
        return None, None

    def _parse_school_time(value):
        parsed = datetime.strptime(value, "%H:%M")
        hour = parsed.hour
        # Timetable afternoon slots are written as 1:30 / 2:10 without PM.
        if hour < 7:
            hour += 12
        return parsed.replace(hour=hour).time()

    start = _parse_school_time(match.group(1))
    end = _parse_school_time(match.group(2))
    return start, end


def timetable_slot_map(source_dir):
    slot_map = {}
    for filename in ("SECOND CA TIMETABLE  JUNIOR.docx", "SECOND CA TIMETABLE  SENIOR.docx"):
        path = Path(source_dir) / filename
        if not path.exists():
            continue
        doc = Document(path)
        if not doc.tables:
            continue
        table = doc.tables[0]
        headers = [collapse(cell.text) for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [collapse(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue
            if not cells[0].upper().startswith("WEDNESDAY"):
                continue
            class_code = class_code_from_value(cells[1])
            per_class = slot_map.setdefault(class_code, [])
            for index in range(2, min(len(headers), len(cells))):
                cell_text = safe_text(cells[index])
                header_text = headers[index]
                if not cell_text:
                    continue
                if cell_text.upper() in {"BREAK", "ANGELUS"}:
                    continue
                per_class.append({
                    "slot_label": normalize_time_text(header_text),
                    "raw_cell": cell_text,
                })
    return slot_map


def resolve_subject_name(class_code, doc_text, file_name, assignments):
    assignment_subjects = {normalize(item.subject.name): item for item in assignments}
    search_space = []
    subject_match = re.search(r"SUBJECT\s*[:\-]\s*(.+)", doc_text, re.IGNORECASE)
    if subject_match:
        search_space.append(subject_match.group(1).splitlines()[0].strip())
    search_space.append(Path(file_name).stem)
    combined_norm = normalize(" ".join(search_space))

    for (alias_class, alias_key), canonical in CLASS_SPECIFIC_ALIASES.items():
        if alias_class == class_code and alias_key in combined_norm:
            target = normalize(canonical)
            if target in assignment_subjects:
                return assignment_subjects[target]

    for alias, canonical in SUBJECT_ALIAS_MAP.items():
        if alias in combined_norm:
            target = normalize(canonical)
            if target in assignment_subjects:
                return assignment_subjects[target]

    for subject_key, assignment in assignment_subjects.items():
        if subject_key and subject_key in combined_norm:
            return assignment

    lowered = combined_norm.lower()
    best = None
    best_score = -1
    for subject_key, assignment in assignment_subjects.items():
        score = 0
        chunks = [chunk for chunk in re.split(r"[^A-Z0-9]+", assignment.subject.name.upper()) if chunk]
        for chunk in chunks:
            token = normalize(chunk)
            if token and token.lower() in lowered:
                score += len(token)
        if score > best_score:
            best = assignment
            best_score = score
    if best_score > 0:
        return best
    return None


def choose_parser_rows(path, subject_name):
    extracted_text = safe_text(extract_text_from_document(str(path)))
    payload = parse_question_document(extracted_text)
    parsed_rows = list(payload.get("parsed_questions") or [])
    flagged_blocks = list(payload.get("flagged_blocks") or [])
    objective_count = sum(1 for row in parsed_rows if (row.get("question_type") or "").upper() == "OBJECTIVE")
    theory_count = sum(1 for row in parsed_rows if (row.get("question_type") or "").upper() != "OBJECTIVE")
    marker_count = _objective_marker_count(extracted_text)
    deterministic_quality = _parsed_rows_quality(parsed_rows)
    parser_used = "deterministic"
    instructions = []

    if str(path).lower().endswith(".docx"):
        try:
            structured = parse_structured_docx_rows(path)
        except Exception:
            structured = {}
        structured_rows = list(structured.get("parsed_rows") or [])
        structured_quality = _parsed_rows_quality(structured_rows)
        if structured_rows and structured_quality >= deterministic_quality:
            parsed_rows = structured_rows
            objective_count = sum(
                1 for row in parsed_rows if (row.get("question_type") or "").upper() == "OBJECTIVE"
            )
            theory_count = sum(
                1 for row in parsed_rows if (row.get("question_type") or "").upper() != "OBJECTIVE"
            )
            parser_used = structured.get("parser_used") or "docx_structural"
            instructions = list(structured.get("instructions") or [])
            flagged_blocks = []
        elif structured.get("instructions"):
            instructions = list(structured.get("instructions") or [])

    suspicious = (
        not parsed_rows
        or (marker_count >= 8 and objective_count < 4)
        or ("SECTION B" in extracted_text.upper() and theory_count < 1)
        or (normalize(subject_name) in {"MATHEMATICS", "PHYSICS"} and len(parsed_rows) < 8)
        or bool(flagged_blocks)
    )
    if suspicious:
        ai_rows = _parse_questions_with_openai(
            extracted_text=safe_text(payload.get("normalized_text") or extracted_text),
            subject_name=subject_name,
            expected_count=max(len(parsed_rows), 10),
        ) or []
        if ai_rows and _parsed_rows_quality(ai_rows) >= deterministic_quality:
            parsed_rows = ai_rows
            objective_count = sum(1 for row in parsed_rows if (row.get("question_type") or "").upper() == "OBJECTIVE")
            theory_count = sum(1 for row in parsed_rows if (row.get("question_type") or "").upper() != "OBJECTIVE")
            parser_used = "ai_fallback_fulltext"

    return {
        "extracted_text": extracted_text,
        "parsed_rows": parsed_rows,
        "objective_count": objective_count,
        "theory_count": theory_count,
        "parser_used": parser_used,
        "marker_count": marker_count,
        "flagged_block_count": len(flagged_blocks),
        "instructions": instructions,
    }


def slot_for_subject(class_code, subject_name, slot_map):
    entries = slot_map.get(class_code, [])
    target = normalize(subject_name)
    alias_targets = {target}
    for alias, canonical in SUBJECT_ALIAS_MAP.items():
        if normalize(canonical) == target:
            alias_targets.add(alias)
    for entry in entries:
        haystack = normalize(entry["raw_cell"])
        if any(item and item in haystack for item in alias_targets):
            return entry
    return None


def build_exam_title(actual_date, class_code, subject_name, slot_label):
    prefix = EXAM_DAY_PREFIX.get(actual_date, actual_date)
    if slot_label:
        return f"{prefix} {slot_label} {class_code} {subject_name} Second CA"
    return f"{prefix} {class_code} {subject_name} Second CA"


def ensure_datetime_window(actual_date, slot_label):
    if not slot_label:
        return None, None, 40
    start_time, end_time = parse_time_window(slot_label)
    if not start_time or not end_time:
        return None, None, 40
    day = datetime.strptime(actual_date, "%Y-%m-%d").date()
    start_at = timezone.make_aware(datetime.combine(day, start_time))
    end_at = timezone.make_aware(datetime.combine(day, end_time))
    minutes = max(int((end_at - start_at).total_seconds() // 60), 40)
    return start_at, end_at, minutes


def replace_existing_imports(source_name, assignment):
    matches = ExamDocumentImport.objects.filter(source_filename=source_name, assignment=assignment)
    for row in matches.select_related("exam"):
        exam = row.exam
        row.delete()
        if exam is not None:
            exam.delete()


@transaction.atomic
def import_one(path, it_user, dean_user, term, slot_map, actual_date):
    class_code = class_code_from_value(path.parent.name)
    assignments = list(
        TeacherSubjectAssignment.objects.select_related("teacher", "subject", "academic_class")
        .filter(academic_class__code=class_code, session=term.session, term=term, is_active=True)
    )
    if not assignments:
        raise RuntimeError(f"No active subject assignments found for class {class_code}.")

    doc = Document(path)
    doc_text = safe_text("\n".join(collapse(p.text) for p in doc.paragraphs if collapse(p.text)))
    assignment = resolve_subject_name(class_code, doc_text, path.name, assignments)
    if assignment is None:
        raise RuntimeError(f"Could not match subject/teacher assignment for {path.name} ({class_code}).")

    parser_info = choose_parser_rows(path, assignment.subject.name)
    parsed_rows = parser_info["parsed_rows"]
    if not parsed_rows:
        raise RuntimeError(f"No questions parsed from {path.name}.")
    exam_instructions = list(parser_info.get("instructions") or [])

    flow_type = "OBJECTIVE_ONLY"
    if parser_info["objective_count"] > 0 and parser_info["theory_count"] > 0:
        flow_type = "OBJECTIVE_THEORY"
    elif parser_info["objective_count"] <= 0:
        flow_type = "THEORY_ONLY"

    slot = slot_for_subject(class_code, assignment.subject.name, slot_map)
    slot_label = slot["slot_label"] if slot else ""
    schedule_start, schedule_end, duration_minutes = ensure_datetime_window(actual_date, slot_label)
    title = build_exam_title(actual_date, class_code, assignment.subject.name, slot_label)

    replace_existing_imports(path.name, assignment)

    question_bank, _ = QuestionBank.objects.get_or_create(
        owner=assignment.teacher,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        name=title + " Bank",
        defaults={"description": "Imported from Thursday CA folder."},
    )

    exam = Exam.objects.create(
        title=title,
        description=f"Imported from {path.parent.name}/{path.name}. Tomorrow paper shifted from Wednesday timetable row.",
        exam_type=CBTExamType.CA,
        status="DRAFT",
        created_by=assignment.teacher,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        question_bank=question_bank,
        schedule_start=schedule_start,
        schedule_end=schedule_end,
        is_time_based=True,
        open_now=False,
        is_free_test=False,
    )

    with path.open("rb") as handle:
        import_row = ExamDocumentImport.objects.create(
            uploaded_by=it_user,
            assignment=assignment,
            exam=exam,
            source_file=File(handle, name=path.name),
            source_filename=path.name,
            extraction_status=CBTDocumentStatus.PENDING,
        )

    created_objective = 0
    created_theory = 0
    sort_order = 0
    for row in parsed_rows:
        question_type = (row.get("question_type") or "OBJECTIVE").upper()
        if flow_type == "OBJECTIVE_ONLY" and question_type != "OBJECTIVE":
            continue
        if flow_type == "THEORY_ONLY" and question_type == "OBJECTIVE":
            continue
        sort_order += 1
        is_objective = question_type == "OBJECTIVE"
        question = Question.objects.create(
            question_bank=question_bank,
            created_by=assignment.teacher,
            subject=assignment.subject,
            question_type=CBTQuestionType.OBJECTIVE if is_objective else CBTQuestionType.SHORT_ANSWER,
            stem=safe_text(row.get("stem") or f"Question {sort_order}"),
            rich_stem=(row.get("rich_stem") or "").strip(),
            topic="Imported CA",
            difficulty=CBTQuestionDifficulty.MEDIUM,
            marks=1,
            source_type=Question.SourceType.DOCUMENT,
            source_reference=str(import_row.id),
        )
        if is_objective:
            for option_sort, label in enumerate(("A", "B", "C", "D"), start=1):
                Option.objects.create(
                    question=question,
                    label=label,
                    option_text=safe_text((row.get("options") or {}).get(label, f"Option {label}")),
                    sort_order=option_sort,
                )
            answer = CorrectAnswer.objects.create(question=question, is_finalized=False)
            correct_label = (row.get("correct_label") or "").strip().upper()
            if correct_label:
                option_qs = question.options.filter(label=correct_label)
                if option_qs.exists():
                    answer.correct_options.set(option_qs)
                    answer.is_finalized = True
                    answer.save(update_fields=["is_finalized", "updated_at"])
            created_objective += 1
        else:
            CorrectAnswer.objects.create(
                question=question,
                is_finalized=False,
                note=safe_text(row.get("model_answer") or ""),
            )
            created_theory += 1
        ExamQuestion.objects.create(exam=exam, question=question, sort_order=sort_order, marks=1)

    blueprint = ensure_default_blueprint(exam)
    blueprint.duration_minutes = duration_minutes
    blueprint.max_attempts = 1
    blueprint.shuffle_questions = True
    blueprint.shuffle_options = True
    blueprint.instructions = instruction_text_for_exam(
        instructions=exam_instructions,
        flow_type=flow_type,
    )
    blueprint.section_config = {
        "flow_type": flow_type,
        "objective_count": created_objective,
        "theory_count": created_theory,
        "theory_response_mode": "PAPER",
        "ca_target": CBTWritebackTarget.CA2,
        "is_free_test": False,
        "manual_score_split": flow_type == "OBJECTIVE_THEORY",
        "objective_target_max": "10.00" if flow_type in {"OBJECTIVE_ONLY", "OBJECTIVE_THEORY"} else "0.00",
        "theory_target_max": "10.00" if flow_type in {"THEORY_ONLY", "OBJECTIVE_THEORY"} else "0.00",
        "calculator_mode": "SCIENTIFIC" if normalize(assignment.subject.name) in {"MATHEMATICS", "PHYSICS", "CHEMISTRY", "FURTHERMATHEMATICS"} else "NONE",
    }
    if flow_type == "OBJECTIVE_ONLY":
        blueprint.objective_writeback_target = CBTWritebackTarget.CA2
        blueprint.theory_enabled = False
        blueprint.theory_writeback_target = CBTWritebackTarget.NONE
    elif flow_type == "THEORY_ONLY":
        blueprint.objective_writeback_target = CBTWritebackTarget.NONE
        blueprint.theory_enabled = True
        blueprint.theory_writeback_target = CBTWritebackTarget.CA3
    else:
        blueprint.objective_writeback_target = CBTWritebackTarget.CA2
        blueprint.theory_enabled = True
        blueprint.theory_writeback_target = CBTWritebackTarget.CA3
    blueprint.save(
        update_fields=[
            "duration_minutes",
            "max_attempts",
            "shuffle_questions",
            "shuffle_options",
            "instructions",
            "section_config",
            "objective_writeback_target",
            "theory_enabled",
            "theory_writeback_target",
            "updated_at",
        ]
    )
    _apply_exam_row_marks(
        exam=exam,
        objective_total=blueprint.section_config["objective_target_max"],
        theory_total=blueprint.section_config["theory_target_max"],
    )

    import_row.extraction_status = CBTDocumentStatus.SUCCESS
    import_row.extracted_text = parser_info["extracted_text"]
    import_row.parse_summary = {
        "parser_used": parser_info["parser_used"],
        "question_count": created_objective + created_theory,
        "objective_count": created_objective,
        "theory_count": created_theory,
        "instruction_count": len(exam_instructions),
        "flagged_block_count": parser_info["flagged_block_count"],
        "objective_marker_count": parser_info["marker_count"],
        "schedule_slot": slot_label,
        "actual_exam_date": actual_date,
        "generated_at": timezone.now().isoformat(),
    }
    import_row.error_message = ""
    import_row.save(update_fields=["extraction_status", "extracted_text", "parse_summary", "error_message", "updated_at"])

    submit_exam_to_dean(exam=exam, actor=assignment.teacher, comment="Imported by IT from teacher-supplied CA document.")
    dean_approve_exam(exam=exam, actor=dean_user, comment="Imported and approved for IT activation.")

    return {
        "file": display_path(path),
        "title": exam.title,
        "exam_id": exam.id,
        "teacher": assignment.teacher.username,
        "subject": assignment.subject.name,
        "class_code": class_code,
        "flow_type": flow_type,
        "objective_count": created_objective,
        "theory_count": created_theory,
        "slot_label": slot_label,
        "status": exam.status,
        "parser_used": parser_info["parser_used"],
    }


def main():
    parser = argparse.ArgumentParser(description="Import Thursday CA papers into NDGA CBT as approved Second CA exams.")
    parser.add_argument("--source-dir", default=DEFAULT_SOURCE_DIR)
    parser.add_argument("--session", default=DEFAULT_SESSION)
    parser.add_argument("--term", default=DEFAULT_TERM)
    parser.add_argument("--actual-date", default=DEFAULT_ACTUAL_DATE)
    args = parser.parse_args()

    session = AcademicSession.objects.get(name=args.session)
    term = Term.objects.get(session=session, name=args.term)
    it_user = User.objects.get(username="admin@ndgakuje.org")
    dean_user = User.objects.filter(username="emmanuel@ndgakuje.org").first() or User.objects.filter(primary_role__code="DEAN").first()
    if dean_user is None:
        raise RuntimeError("No dean user found for approval step.")

    slot_map = timetable_slot_map(args.source_dir)
    rows = []
    failures = []
    for path in sorted(Path(args.source_dir).glob("*/*.docx")):
        try:
            rows.append(import_one(path, it_user, dean_user, term, slot_map, args.actual_date))
        except Exception as exc:
            failures.append((display_path(path), str(exc)))

    print(f"Imported: {len(rows)} | Failed: {len(failures)}")
    for row in rows:
        print(f"OK | {row['exam_id']} | {row['status']} | {row['class_code']} | {row['subject']} | {row['flow_type']} | obj={row['objective_count']} | theory={row['theory_count']} | {row['parser_used']} | {row['slot_label']} | {row['title']}")
    for path, error in failures:
        print(f"FAIL | {path} | {error}")
    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
