from __future__ import annotations

import re
from pathlib import Path

from django.core.management.base import BaseCommand, CommandError

from apps.accounts.models import StudentProfile
from apps.notifications.models import BirthdayContact, BirthdayContactType

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


MONTHS = {
    "JANUARY": 1,
    "FEBRUARY": 2,
    "FEBRUAURY": 2,
    "MARCH": 3,
    "APRIL": 4,
    "MAY": 5,
    "JUNE": 6,
    "JULY": 7,
    "AUGUST": 8,
    "SEPTEMBER": 9,
    "OCTOBER": 10,
    "NOVEMBER": 11,
    "DECEMBER": 12,
}
EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
PHONE_RE = re.compile(r"(?<!\d)(?:\+?234|0)\d{9,10}(?!\d)")
DAY_RE = re.compile(r"\((\d{1,2})(?:st|nd|rd|th)?\)", re.IGNORECASE)
TITLE_WORDS = {"mr", "mrs", "miss", "ms", "dr", "chief", "prof", "rev", "sir", "lady"}


def _clean_name(value):
    value = re.sub(EMAIL_RE, "", value or "")
    value = re.sub(PHONE_RE, "", value)
    value = re.sub(r"\(\s*\d{1,2}\s*(?:st|nd|rd|th)?\s*\)", "", value, flags=re.IGNORECASE)
    value = value.replace("–", "-")
    value = value.split("-", 1)[0]
    return " ".join(value.strip(" -").split())


def _name_tokens(value):
    cleaned = re.sub(r"[^a-z0-9 ]+", " ", (value or "").lower())
    tokens = [item for item in cleaned.split() if item and item not in TITLE_WORDS]
    return tokens


def _names_match(imported_name, guardian_name):
    imported_tokens = _name_tokens(imported_name)
    guardian_tokens = _name_tokens(guardian_name)
    if not imported_tokens or not guardian_tokens:
        return False
    imported_set = set(imported_tokens)
    guardian_set = set(guardian_tokens)
    if imported_set == guardian_set:
        return True
    overlap = imported_set & guardian_set
    return len(overlap) >= 2 and (
        overlap == imported_set
        or overlap == guardian_set
        or len(overlap) >= min(len(imported_set), len(guardian_set)) - 1
    )


def _student_display_name(profile):
    user = profile.user
    return user.get_full_name() or user.display_name or user.username


def _candidate_student_profiles_for_contact(contact):
    profiles = (
        StudentProfile.objects.select_related("user")
        .filter(user__is_active=True)
        .order_by("user__last_name", "user__first_name")
    )
    contact_tokens = _name_tokens(contact.full_name)
    rows = []
    for profile in profiles.exclude(guardian_name=""):
        if _names_match(contact.full_name, profile.guardian_name):
            return [profile]
    if not contact_tokens:
        return []
    surname = contact_tokens[-1]
    for profile in profiles:
        user = profile.user
        student_tokens = _name_tokens(
            " ".join(
                [
                    user.first_name,
                    user.last_name,
                    getattr(profile, "middle_name", ""),
                    getattr(profile, "student_number", ""),
                    user.username,
                ]
            )
        )
        if surname in student_tokens:
            rows.append(profile)
    return rows if len(rows) == 1 else []


def _match_student_guardian(contact):
    if contact.contact_type != BirthdayContactType.PARENT:
        return False
    for profile in _candidate_student_profiles_for_contact(contact):
        changed = []
        if profile.guardian_email and contact.email != profile.guardian_email.strip().lower():
            contact.email = profile.guardian_email.strip().lower()
            changed.append("email")
        if profile.guardian_phone and contact.phone != profile.guardian_phone.strip():
            contact.phone = profile.guardian_phone.strip()
            changed.append("phone")
        student_name = _student_display_name(profile)
        if contact.student_name != student_name:
            contact.student_name = student_name
            changed.append("student_name")
        admission_no = (profile.student_number or profile.user.username or "").strip().upper()
        if contact.student_admission_no != admission_no:
            contact.student_admission_no = admission_no
            changed.append("student_admission_no")
        if contact.linked_user_id != profile.user_id:
            contact.linked_user = profile.user
            changed.append("linked_user")
        if changed:
            changed.append("updated_at")
            contact.save(update_fields=changed)
        return True
    return False


def _iter_doc_lines(path):
    document = Document(str(path))
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                yield " | ".join(cells)


def _import_file(path, *, contact_type, default_month=None, match_students=True):
    current_month = default_month
    pending = None
    imported = 0
    updated = 0
    skipped = 0

    def flush_pending():
        nonlocal pending, imported, updated, skipped
        if not pending:
            return
        if not (pending["name"] and pending["month"] and pending["day"]):
            skipped += 1
            pending = None
            return
        obj, created = BirthdayContact.objects.update_or_create(
            contact_type=contact_type,
            full_name=pending["name"],
            birth_month=pending["month"],
            birth_day=pending["day"],
            defaults={
                "email": pending.get("email", ""),
                "phone": pending.get("phone", ""),
                "source_label": path.name,
                "is_active": True,
            },
        )
        if match_students:
            _match_student_guardian(obj)
        imported += int(created)
        updated += int(not created)
        pending = None

    for raw_line in _iter_doc_lines(path):
        line = " ".join(raw_line.split())
        upper = line.strip().upper()
        if upper in MONTHS:
            flush_pending()
            current_month = MONTHS[upper]
            continue

        emails = EMAIL_RE.findall(line)
        phones = PHONE_RE.findall(line)
        day_match = DAY_RE.search(line)

        if emails and pending and not day_match:
            pending["email"] = pending.get("email") or emails[0].strip().lower()
            continue

        if phones and pending and not day_match and not _clean_name(line):
            pending["phone"] = pending.get("phone") or phones[0].strip()
            continue

        if day_match and current_month:
            flush_pending()
            name = _clean_name(line)
            if not name:
                skipped += 1
                continue
            pending = {
                "name": name,
                "month": current_month,
                "day": int(day_match.group(1)),
                "email": emails[0].strip().lower() if emails else "",
                "phone": phones[0].strip() if phones else "",
            }
            continue

    flush_pending()
    return {"imported": imported, "updated": updated, "skipped": skipped}


class Command(BaseCommand):
    help = "Import parent and staff birthday contacts from NDGA birthday Word documents."

    def add_arguments(self, parser):
        parser.add_argument(
            "--documents-dir",
            default=r"C:\Users\NDGA ADMIN\Documents",
            help="Folder containing birthday .docx files.",
        )
        parser.add_argument(
            "--no-student-match",
            action="store_true",
            help="Import birthday documents without matching parent contacts to student guardian records.",
        )

    def handle(self, *args, **options):
        if Document is None:
            raise CommandError("python-docx is required to import birthday Word documents.")

        documents_dir = Path(options["documents_dir"])
        if not documents_dir.exists():
            raise CommandError(f"Documents folder not found: {documents_dir}")

        totals = {"imported": 0, "updated": 0, "skipped": 0, "files": 0}
        for path in sorted(documents_dir.glob("*.docx")):
            upper_name = path.name.upper()
            if "STAFF BIRTHDAY" in upper_name or "STAFF BIRTHDAYS" in upper_name:
                result = _import_file(
                    path,
                    contact_type=BirthdayContactType.STAFF,
                    match_students=not options["no_student_match"],
                )
            elif "BORN" in upper_name and "PARENT" in upper_name:
                default_month = next((number for label, number in MONTHS.items() if label in upper_name), None)
                result = _import_file(
                    path,
                    contact_type=BirthdayContactType.PARENT,
                    default_month=default_month,
                    match_students=not options["no_student_match"],
                )
            else:
                continue
            totals["files"] += 1
            for key in ("imported", "updated", "skipped"):
                totals[key] += result[key]

        self.stdout.write(
            self.style.SUCCESS(
                "Birthday import complete: "
                f"files={totals['files']} imported={totals['imported']} "
                f"updated={totals['updated']} skipped={totals['skipped']}"
            )
        )
