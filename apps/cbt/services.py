from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from pathlib import Path
import random
import shutil
import tempfile
import zipfile
from copy import deepcopy

from django.conf import settings
from django.core.exceptions import ValidationError
from django.db import transaction
from django.db.models import Max, Q
from django.utils.text import slugify
from django.utils import timezone

from apps.accounts.constants import (
    ROLE_IT_MANAGER,
    ROLE_PRINCIPAL,
    ROLE_STUDENT,
    ROLE_VP,
)
from apps.accounts.permissions import has_any_role
from apps.academics.models import (
    StudentClassEnrollment,
    StudentSubjectEnrollment,
    TeacherSubjectAssignment,
)
from apps.cbt.models import (
    CBTSimulationCallbackType,
    CBTAttemptStatus,
    CBTDocumentStatus,
    CBTExamType,
    CBTExamStatus,
    CBTQuestionType,
    CBTQuestionDifficulty,
    CBTSimulationAttemptStatus,
    CBTSimulationScoreMode,
    CBTSimulationWrapperStatus,
    CBTWritebackTarget,
    CorrectAnswer,
    Exam,
    ExamAttempt,
    ExamAttemptAnswer,
    ExamBlueprint,
    ExamDocumentImport,
    ExamQuestion,
    ExamReviewAction,
    ExamSimulation,
    Option,
    Question,
    QuestionBank,
    SimulationAttemptRecord,
    SimulationWrapper,
)
from apps.cbt.simulation_catalog import CURATED_SIMULATION_CATALOG, grouped_catalog_labels
from apps.cbt.subject_matrix import resolve_subject_simulation_profile
from apps.results.models import ResultSheet, ResultSheetStatus, StudentSubjectScore
from apps.audit.services import log_event, log_lockdown_violation
from apps.audit.models import AuditCategory, AuditStatus
from apps.setup_wizard.services import get_setup_state
from apps.sync.services import queue_exam_attempt_sync, queue_simulation_attempt_sync

QUESTION_START_PATTERN = re.compile(r"^\s*(\d+)[\).\:-]\s*(.+)$")
QUESTION_PREFIX_PATTERN = re.compile(r"^\s*(?:question|q)\s*(\d+)\s*[\).\:-]?\s*(.*)$", re.IGNORECASE)
QUESTION_BLOCK_SPLIT_PATTERN = re.compile(r"(?im)^(?:\s*(?:question|q)\s*)?(\d{1,3})\s*[\)\.\:-]\s*")
OPTION_PATTERN = re.compile(r"^\s*([A-Da-d])(?:\s*[\)\.\:-]|\s+)\s*(.+)$")
ANSWER_PATTERN = re.compile(
    r"^\s*(?:answer|ans|correct(?:\s*answer|\s*option)?)\s*[\.:\-]?\s*(.+)$",
    re.IGNORECASE,
)
DOUBLE_NEWLINE_SPLIT_PATTERN = re.compile(r"\n\s*\n+")
INLINE_OPTION_BOUNDARY_PATTERN = re.compile(r"(?<=\S)\s+(?=[A-Da-d][\)\.\:-]\s*)")
INLINE_ANSWER_BOUNDARY_PATTERN = re.compile(
    r"(?<=\S)\s+(?=(?:answer|ans|correct(?:\s*answer|\s*option)?)\s*[\.:\-])",
    re.IGNORECASE,
)
INLINE_QUESTION_BOUNDARY_PATTERN = re.compile(r"(?<=\S)\s+(?=\d+\s*[\)\.\:-]\s*)")
INLINE_OPTION_GLUE_PATTERN = re.compile(r"([^\s])([A-Da-d][\)\.\:-]\s*)")
INLINE_ANSWER_GLUE_PATTERN = re.compile(
    r"([^\s])((?:answer|ans|correct(?:\s*answer|\s*option)?)\s*[\.:\-]\s*)",
    re.IGNORECASE,
)
INLINE_QUESTION_GLUE_PATTERN = re.compile(r"([^\s])(\d+\s*[\)\.\:-]\s*)")
INLINE_QUESTION_PREFIX_GLUE_PATTERN = re.compile(
    r"([^\s])((?:question|q)\s*\d+\s*[\)\.\:-]\s*)",
    re.IGNORECASE,
)
INLINE_OPTION_MARKER_PATTERN = re.compile(r"(?i)\b([A-D])\s*[\)\.\:-]\s*")
HEADER_NOISE_PATTERN = re.compile(
    r"^\s*(?:topic|subject|section|instruction|instructions|time|duration|class|name|date)\b",
    re.IGNORECASE,
)


@dataclass
class RawQuestionBlock:
    number: str
    body: str


def can_manage_all_cbt(user):
    return has_any_role(user, {ROLE_IT_MANAGER, ROLE_PRINCIPAL, ROLE_VP})


def authoring_assignment_queryset(
    user,
    *,
    include_all_periods=False,
    selected_session=None,
    selected_term=None,
):
    setup_state = get_setup_state()
    qs = TeacherSubjectAssignment.objects.select_related(
        "teacher",
        "subject",
        "academic_class",
        "session",
        "term",
    ).filter(is_active=True)
    if selected_session:
        qs = qs.filter(session=selected_session)
    if selected_term:
        qs = qs.filter(term=selected_term)
    if (
        not include_all_periods
        and selected_session is None
        and selected_term is None
        and setup_state.current_session_id
        and setup_state.current_term_id
    ):
        qs = qs.filter(
            session=setup_state.current_session,
            term=setup_state.current_term,
        )
    if can_manage_all_cbt(user):
        return qs.order_by("academic_class__code", "subject__name", "teacher__username")
    return qs.filter(teacher=user).order_by("academic_class__code", "subject__name")


def authoring_question_bank_queryset(user):
    qs = QuestionBank.objects.select_related(
        "owner",
        "subject",
        "academic_class",
        "session",
        "term",
    )
    if can_manage_all_cbt(user):
        return qs.order_by("-updated_at")
    return qs.filter(owner=user).order_by("-updated_at")


def authoring_exam_queryset(user):
    qs = Exam.objects.select_related(
        "created_by",
        "subject",
        "academic_class",
        "session",
        "term",
        "question_bank",
    )
    if can_manage_all_cbt(user):
        return qs.order_by("-updated_at")
    return qs.filter(created_by=user).order_by("-updated_at")


def simulation_registry_queryset():
    return SimulationWrapper.objects.select_related(
        "created_by",
        "dean_reviewed_by",
    ).order_by("tool_name", "-updated_at")


def simulation_catalog_grouped_labels():
    return grouped_catalog_labels()


def _is_safe_zip_member(member_name):
    member_path = Path(member_name)
    if member_path.is_absolute():
        return False
    if ".." in member_path.parts:
        return False
    return True


def store_simulation_bundle(*, wrapper, uploaded_bundle):
    if uploaded_bundle is None:
        raise ValidationError("Simulation bundle file is required.")

    media_root = Path(settings.MEDIA_ROOT)
    media_root.mkdir(parents=True, exist_ok=True)
    slug = slugify(wrapper.tool_name) or "simulation"
    bundle_root = media_root / "sims" / f"{wrapper.id}-{slug}"
    if bundle_root.exists():
        shutil.rmtree(bundle_root)
    bundle_root.mkdir(parents=True, exist_ok=True)

    tmp_bundle_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as temp_zip:
            tmp_bundle_path = Path(temp_zip.name)
            for chunk in uploaded_bundle.chunks():
                temp_zip.write(chunk)

        try:
            with zipfile.ZipFile(tmp_bundle_path) as zip_file:
                for member in zip_file.infolist():
                    member_name = member.filename or ""
                    if not member_name.strip():
                        continue
                    if not _is_safe_zip_member(member_name):
                        raise ValidationError("Invalid ZIP structure detected.")
                    zip_file.extract(member, bundle_root)
        except zipfile.BadZipFile as exc:
            raise ValidationError("Upload a valid ZIP simulation bundle.") from exc

        html_candidates = list(bundle_root.rglob("index.html"))
        if not html_candidates:
            html_candidates = list(bundle_root.rglob("*.html"))
        if not html_candidates:
            raise ValidationError("No HTML entry file found in uploaded ZIP.")

        entry_file = sorted(
            html_candidates,
            key=lambda row: (len(row.parts), row.name.lower()),
        )[0]
        relative = entry_file.relative_to(media_root).as_posix()
        media_url = settings.MEDIA_URL.rstrip("/")
        return f"{media_url}/{relative}"
    finally:
        if tmp_bundle_path and tmp_bundle_path.exists():
            tmp_bundle_path.unlink(missing_ok=True)


@transaction.atomic
def seed_simulation_library_rows(*, actor, rows):
    created = 0
    updated = 0
    skipped = 0
    seeded_ids = []
    for row in rows:
        defaults = {
            "tool_type": row.get("tool_type", "").strip(),
            "source_provider": row.get("source_provider", "OTHER"),
            "source_reference_url": row.get("source_reference_url", "").strip(),
            "description": row.get("description", "").strip(),
            "online_url": row.get("online_url", "").strip(),
            "offline_asset_path": row.get("offline_asset_path", "").strip(),
            "score_mode": row.get("score_mode", CBTSimulationScoreMode.AUTO),
            "max_score": row.get("max_score", "10.00"),
            "scoring_callback_type": row.get(
                "scoring_callback_type",
                CBTSimulationCallbackType.POST_MESSAGE,
            ),
            "evidence_required": bool(row.get("evidence_required", False)),
            "is_active": True,
        }
        wrapper, was_created = SimulationWrapper.objects.get_or_create(
            tool_name=row["tool_name"],
            tool_category=row["tool_category"],
            defaults={**defaults, "created_by": actor},
        )
        if was_created:
            created += 1
            seeded_ids.append(wrapper.id)
            continue

        changed_fields = []
        for field_name, field_value in defaults.items():
            if getattr(wrapper, field_name) != field_value:
                setattr(wrapper, field_name, field_value)
                changed_fields.append(field_name)
        if changed_fields:
            wrapper.save(update_fields=[*changed_fields, "updated_at"])
            updated += 1
        else:
            skipped += 1
        seeded_ids.append(wrapper.id)

    return {
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "seeded_ids": seeded_ids,
        "total_seed_rows": len(rows),
    }


def seed_curated_simulation_library(*, actor):
    return seed_simulation_library_rows(
        actor=actor,
        rows=CURATED_SIMULATION_CATALOG,
    )


def simulation_recommendation_for_subject(subject):
    return resolve_subject_simulation_profile(subject)


def approved_simulation_queryset(subject=None):
    return SimulationWrapper.objects.filter(
        status=CBTSimulationWrapperStatus.APPROVED,
        is_active=True,
    ).order_by("tool_name")


def recommended_simulation_queryset(subject=None):
    approved = approved_simulation_queryset()
    profile = simulation_recommendation_for_subject(subject)
    preferred_categories = profile.get("tool_categories") or []
    preferred_providers = profile.get("preferred_providers") or []
    if not preferred_categories and not preferred_providers:
        return approved

    filter_q = Q()
    if preferred_categories:
        filter_q |= Q(tool_category__in=preferred_categories)
    if preferred_providers:
        filter_q |= Q(source_provider__in=preferred_providers)

    recommended = approved.filter(filter_q)

    subject_name = (getattr(subject, "name", "") or "").strip().lower()
    subject_tokens = [token for token in re.findall(r"[a-z0-9]+", subject_name) if len(token) >= 4]
    if subject_tokens:
        keyword_q = Q()
        for token in subject_tokens:
            keyword_q |= Q(tool_name__icontains=token)
            keyword_q |= Q(description__icontains=token)
        keyword_hits = approved.filter(keyword_q)
        if keyword_hits.exists():
            recommended = approved.filter(Q(id__in=recommended.values("id")) | keyword_q)

    if recommended.exists():
        return recommended
    return approved


def _extract_docx_text_with_xml_order(path):
    import xml.etree.ElementTree as ET

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    def _clean_join(parts):
        value = re.sub(r"\s+", " ", " ".join(part for part in parts if part)).strip()
        return value

    with zipfile.ZipFile(path) as zf:
        xml_raw = zf.read("word/document.xml")
    root = ET.fromstring(xml_raw)

    extracted_rows = []
    for paragraph in root.findall(".//w:body/w:p", namespaces):
        fragments = []
        has_equation = False
        has_embedded_object = False

        for node in paragraph.iter():
            tag = node.tag
            if not isinstance(tag, str):
                continue
            if tag.endswith("}t"):
                text_value = (node.text or "").strip()
                if text_value:
                    fragments.append(text_value)
                continue
            if tag.endswith("}oMath") or tag.endswith("}oMathPara"):
                has_equation = True
                equation_tokens = []
                for eq_node in node.iter():
                    eq_tag = eq_node.tag if isinstance(eq_node.tag, str) else ""
                    if eq_tag.endswith("}t"):
                        token = (eq_node.text or "").strip()
                        if token:
                            equation_tokens.append(token)
                    elif eq_tag.endswith("}chr"):
                        token = (eq_node.attrib.get(f"{{{namespaces['m']}}}val", "") or "").strip()
                        if token:
                            equation_tokens.append(token)
                    elif eq_tag.endswith("}sym"):
                        token = (eq_node.attrib.get(f"{{{namespaces['m']}}}char", "") or "").strip()
                        if token:
                            equation_tokens.append(token)
                eq_text = _clean_join(equation_tokens)
                fragments.append(f"[EQ: {eq_text}]" if eq_text else "[EQUATION]")
                continue
            if tag.endswith("}drawing") or tag.endswith("}pict") or tag.endswith("}shape"):
                has_embedded_object = True

        row_text = _clean_join(fragments)
        if row_text:
            extracted_rows.append(row_text)
            continue
        if has_equation:
            extracted_rows.append("[EQUATION]")
        elif has_embedded_object:
            extracted_rows.append("[EMBEDDED_OBJECT]")

    if extracted_rows:
        return "\n".join(extracted_rows)

    # Fallback generic extraction of any text runs if paragraph order extraction is empty.
    texts = [node.text for node in root.findall(".//w:t", namespaces) if node.text]
    math_texts = [node.text for node in root.findall(".//m:t", namespaces) if node.text]
    merged = texts + math_texts
    return "\n".join(merged)


def _extract_docx_text(path):
    try:
        from docx import Document

        document = Document(str(path))
        paragraph_rows = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                cells = [re.sub(r"\s+", " ", (cell.text or "").strip()) for cell in row.cells]
                cell_text = " | ".join(cell for cell in cells if cell)
                if cell_text:
                    table_rows.append(cell_text)
        direct_text = "\n".join(paragraph_rows + table_rows)
    except Exception:
        direct_text = ""

    xml_order_text = ""
    try:
        xml_order_text = _extract_docx_text_with_xml_order(path)
    except Exception:
        xml_order_text = ""

    def _density(value):
        return len(re.sub(r"\s+", "", str(value or "")))

    if _density(xml_order_text) >= _density(direct_text):
        return xml_order_text or direct_text
    return direct_text or xml_order_text


def extract_text_from_document(path_or_file):
    path = Path(str(path_or_file))
    extension = path.suffix.lower()
    if extension == ".pdf":
        extraction_candidates = []

        def _remember_candidate(raw_value):
            value = str(raw_value or "").strip()
            if value:
                extraction_candidates.append(value)

        try:
            import pdfplumber

            plumber_rows = []
            with pdfplumber.open(str(path)) as document:
                for page in document.pages:
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                    if not page_text.strip():
                        page_text = page.extract_text_simple() or ""
                    if page_text.strip():
                        plumber_rows.append(page_text)
            _remember_candidate("\n\n".join(plumber_rows))
        except Exception:
            pass

        try:
            from pdfminer.high_level import extract_text as pdf_extract_text

            _remember_candidate(pdf_extract_text(str(path)) or "")
        except Exception:
            pass

        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            _remember_candidate("\n".join((page.extract_text() or "") for page in reader.pages))
        except Exception:
            pass

        try:
            from PyPDF2 import PdfReader as LegacyPdfReader

            reader = LegacyPdfReader(str(path))
            _remember_candidate("\n".join((page.extract_text() or "") for page in reader.pages))
        except Exception:
            pass

        def _text_density(raw_value):
            compact = re.sub(r"\s+", "", str(raw_value or ""))
            return len(compact)

        text = ""
        if extraction_candidates:
            text = max(
                extraction_candidates,
                key=lambda candidate: (
                    _question_structure_score(candidate),
                    _text_density(candidate),
                ),
            )
        if _text_density(text) < 60:
            try:
                import fitz
                import pytesseract
                from PIL import Image

                ocr_chunks = []
                with fitz.open(str(path)) as document:
                    for page in document:
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        ocr_chunks.append(pytesseract.image_to_string(image) or "")
                ocr_text = "\n".join(chunk for chunk in ocr_chunks if chunk.strip())
                if _text_density(ocr_text) > _text_density(text):
                    text = ocr_text
            except Exception:
                pass
        if _text_density(text) < 60:
            try:
                import pytesseract
                from pdf2image import convert_from_path

                pages = convert_from_path(str(path), dpi=220)
                ocr_text = "\n".join(pytesseract.image_to_string(image) or "" for image in pages)
                if _text_density(ocr_text) > _text_density(text):
                    text = ocr_text
            except Exception:
                pass
        if not str(text or "").strip():
            raise ValidationError(
                "Could not extract readable text from PDF. Upload a clearer PDF, image, or paste text."
            )
        return str(text or "")
    if extension == ".docx":
        try:
            text = _extract_docx_text(path)
            if text.strip():
                return text
            raise ValidationError("DOCX file is empty or unreadable.")
        except Exception:
            try:
                text = _extract_docx_text_with_xml_order(path)
                if text.strip():
                    return text
                raise ValidationError("DOCX file is empty or unreadable.")
            except Exception as exc:
                raise ValidationError("Could not parse DOCX content.") from exc
    if extension == ".doc":
        # Legacy binary .doc files are unreliable with naive UTF decoding.
        # Try common external converters first, then fail clearly if unreadable.
        import subprocess

        converter_commands = [
            ["antiword", str(path)],
            ["catdoc", str(path)],
        ]
        for command in converter_commands:
            try:
                result = subprocess.run(
                    command,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="ignore",
                    timeout=20,
                    check=False,
                )
            except Exception:
                continue
            candidate = (result.stdout or "").strip()
            if result.returncode == 0 and len(re.sub(r"\s+", "", candidate)) >= 40:
                return candidate
        raise ValidationError(
            "Legacy .doc extraction is not reliable on this server. Save as .docx or PDF and upload again."
        )

    if extension in {".txt", ".rtf"}:
        raw = path.read_bytes()
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            pass
        try:
            return raw.decode("utf-16")
        except UnicodeDecodeError:
            pass
        return raw.decode("latin-1", errors="ignore")
    if extension in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}:
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(path)
            return pytesseract.image_to_string(image) or ""
        except Exception as exc:
            raise ValidationError(
                "Image OCR not available. Install pytesseract + pillow and configure tesseract binary."
            ) from exc
    raise ValidationError(
        "Unsupported file type. Upload PDF, DOC, DOCX, TXT, or image file (PNG/JPG/JPEG/BMP/TIFF/WEBP)."
    )


def _normalize_extracted_text(raw_text):
    text = str(raw_text or "")
    if not text:
        return ""
    text = (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\u00A0", " ")
        .replace("\u2007", " ")
        .replace("\u202F", " ")
    )
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
    text = INLINE_QUESTION_PREFIX_GLUE_PATTERN.sub(r"\1\n\2", text)
    text = INLINE_OPTION_GLUE_PATTERN.sub(r"\1\n\2", text)
    text = INLINE_ANSWER_GLUE_PATTERN.sub(r"\1\n\2", text)
    text = INLINE_QUESTION_GLUE_PATTERN.sub(r"\1\n\2", text)
    text = INLINE_OPTION_BOUNDARY_PATTERN.sub("\n", text)
    text = INLINE_ANSWER_BOUNDARY_PATTERN.sub("\n", text)
    text = INLINE_QUESTION_BOUNDARY_PATTERN.sub("\n", text)
    text = re.sub(r"(?im)(?<!\n)\(([A-Da-d])\)\s*", r"\n\1) ", text)
    text = re.sub(r"(?im)(?<!\n)([A-Da-d])\s*[\)\.\:-]\s*", r"\n\1) ", text)
    text = re.sub(r"(?im)^(?:question|q)\s*(\d+)\s*[\)\.\:-]?\s*", r"\1. ", text)
    text = re.sub(r"(?im)^(\d{1,3})\s*\)\s*$", r"\1. ", text)
    text = re.sub(r"(?im)^(\d{1,3})\s*\)\s*(\S)", r"\1. \2", text)

    cleaned_rows = []
    seen_question_like = False
    for row in text.splitlines():
        line = row.strip()
        if not line:
            cleaned_rows.append("")
            continue
        looks_like_question = bool(
            QUESTION_BLOCK_SPLIT_PATTERN.match(line)
            or OPTION_PATTERN.match(line)
            or ANSWER_PATTERN.match(line)
        )
        if not seen_question_like and HEADER_NOISE_PATTERN.match(line) and not looks_like_question:
            continue
        if looks_like_question:
            seen_question_like = True
        cleaned_rows.append(line)

    text = "\n".join(cleaned_rows)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _question_structure_score(raw_value):
    text = str(raw_value or "")
    compact_len = len(re.sub(r"\s+", "", text))
    question_markers = len(
        re.findall(r"(?im)^(?:\s*(?:question|q)\s*)?\d{1,3}\s*[\)\.\:-]\s*", text)
    )
    option_markers = len(
        re.findall(r"(?im)(?:^|\s)[A-Da-d]\s*[\)\.\:-]\s*", text)
    )
    answer_markers = len(
        re.findall(
            r"(?im)(?:^|\s)(?:answer|ans|correct(?:\s*answer|\s*option)?)\s*[\.:\-]?\s*",
            text,
        )
    )
    # Bias selection toward candidate text that preserves exam structure,
    # not just raw character volume.
    return (
        (question_markers * 60)
        + (option_markers * 30)
        + (answer_markers * 20)
        + min(compact_len, 12000) / 120
    )


def _split_into_question_blocks(normalized_text):
    content = (normalized_text or "").strip()
    if not content:
        return []
    matches = list(QUESTION_BLOCK_SPLIT_PATTERN.finditer(content))
    blocks = []
    if matches:
        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
            body = content[start:end].strip()
            if body:
                blocks.append(RawQuestionBlock(number=match.group(1), body=body))
        if blocks:
            return blocks

    chunks = [row.strip() for row in DOUBLE_NEWLINE_SPLIT_PATTERN.split(content) if row.strip()]
    if len(chunks) > 1:
        return [RawQuestionBlock(number=str(index), body=row) for index, row in enumerate(chunks, start=1)]
    return [RawQuestionBlock(number="?", body=content)]


def _resolve_correct_label_from_answer(*, options, answer_raw):
    answer_text = (answer_raw or "").strip()
    if not answer_text:
        return ""
    label_match = re.match(r"^\s*([A-Da-d])(?:\b|[\).\:-])?", answer_text)
    if label_match:
        label = label_match.group(1).upper()
        if label in options:
            return label
    normalized_answer = re.sub(r"\s+", " ", answer_text).strip().lower()
    if not normalized_answer:
        return ""
    for label, value in options.items():
        normalized_option = re.sub(r"\s+", " ", str(value or "")).strip().lower()
        if normalized_option and normalized_answer == normalized_option:
            return label
    for label, value in options.items():
        normalized_option = re.sub(r"\s+", " ", str(value or "")).strip().lower()
        if normalized_option and normalized_answer in normalized_option:
            return label
    return ""


def _extract_inline_objective_from_text(text):
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    if not value:
        return None

    answer_raw = ""
    answer_match = re.search(
        r"(?is)\b(?:answer|ans|correct(?:\s*answer|\s*option)?)\s*[\.:\-]?\s*(.+)$",
        value,
    )
    if answer_match:
        answer_raw = (answer_match.group(1) or "").strip()
        value = value[: answer_match.start()].strip()

    markers = list(INLINE_OPTION_MARKER_PATTERN.finditer(value))
    if len(markers) < 4:
        return None

    first_seen = {}
    for marker in markers:
        label = marker.group(1).upper()
        if label not in first_seen:
            first_seen[label] = marker
    if not all(label in first_seen for label in ("A", "B", "C", "D")):
        return None

    ordered_markers = [first_seen[label] for label in ("A", "B", "C", "D")]
    if not (ordered_markers[0].start() < ordered_markers[1].start() < ordered_markers[2].start() < ordered_markers[3].start()):
        return None

    stem = value[: ordered_markers[0].start()].strip(" :-")
    if not stem:
        return None

    options = {}
    for index, marker in enumerate(ordered_markers):
        label = marker.group(1).upper()
        start = marker.end()
        end = ordered_markers[index + 1].start() if index + 1 < len(ordered_markers) else len(value)
        option_text = value[start:end].strip(" :-")
        option_text = re.sub(r"\s+", " ", option_text).strip()
        if not option_text:
            return None
        options[label] = option_text

    return {"stem": stem, "options": options, "answer_raw": answer_raw}


def _parse_question_block(block_text):
    text = (block_text or "").strip()
    if not text:
        return None
    lines = [re.sub(r"\s+", " ", row.strip()) for row in text.splitlines() if row.strip()]
    if not lines:
        return None

    question_match = QUESTION_START_PATTERN.match(lines[0]) or QUESTION_PREFIX_PATTERN.match(lines[0])
    if question_match:
        lines[0] = (question_match.group(2) or "").strip()
    lines = [row for row in lines if row]
    if not lines:
        return None

    stem_chunks = []
    options = {}
    last_option_label = ""
    answer_raw = ""
    in_answer_block = False

    for line in lines:
        answer_match = ANSWER_PATTERN.match(line)
        if answer_match:
            answer_piece = (answer_match.group(1) or "").strip()
            if answer_piece:
                answer_raw = f"{answer_raw} {answer_piece}".strip() if answer_raw else answer_piece
            in_answer_block = True
            last_option_label = ""
            continue

        option_match = OPTION_PATTERN.match(line)
        if option_match and not in_answer_block:
            label = option_match.group(1).upper()
            option_text = (option_match.group(2) or "").strip()
            if option_text:
                existing = options.get(label, "")
                options[label] = f"{existing} {option_text}".strip() if existing else option_text
                last_option_label = label
            continue

        if in_answer_block:
            answer_raw = f"{answer_raw} {line}".strip() if answer_raw else line
            continue

        if last_option_label:
            options[last_option_label] = f"{options[last_option_label]} {line}".strip()
            continue

        stem_chunks.append(line)

    stem = re.sub(r"\s+", " ", " ".join(stem_chunks)).strip()
    ordered_options = {
        label: options[label]
        for label in ("A", "B", "C", "D")
        if options.get(label)
    }

    if stem and len(ordered_options) < 4:
        inline_payload = _extract_inline_objective_from_text(text)
        if inline_payload:
            stem = inline_payload["stem"]
            ordered_options = inline_payload["options"]
            if not answer_raw:
                answer_raw = inline_payload["answer_raw"]

    if stem and len(ordered_options) >= 4:
        return {
            "question_type": "OBJECTIVE",
            "stem": stem,
            "options": ordered_options,
            "correct_label": _resolve_correct_label_from_answer(
                options=ordered_options,
                answer_raw=answer_raw,
            ),
        }
    if stem:
        return {
            "question_type": "THEORY",
            "stem": stem,
            "model_answer": re.sub(r"\s+", " ", answer_raw).strip(),
        }
    return None


def _extract_json_payload(raw_text):
    text = (raw_text or "").strip()
    if not text:
        return ""
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s*```$", "", text)
    start_brace = text.find("{")
    start_list = text.find("[")
    starts = [value for value in [start_brace, start_list] if value >= 0]
    if starts:
        text = text[min(starts):]
    last_brace = text.rfind("}")
    last_list = text.rfind("]")
    end_candidates = [value for value in [last_brace, last_list] if value >= 0]
    if end_candidates:
        text = text[: max(end_candidates) + 1]
    return text.strip()


def _openai_json_response(*, system_prompt, user_prompt):
    api_key = (
        getattr(settings, "OPENAI_API_KEY", "") or os.getenv("OPENAI_API_KEY", "")
    ).strip()
    if not api_key:
        return None
    model = (
        getattr(settings, "OPENAI_CBT_MODEL", "") or os.getenv("OPENAI_CBT_MODEL", "")
    ).strip() or "gpt-4.1-mini"
    try:
        from openai import OpenAI
    except Exception:
        return None

    client = OpenAI(api_key=api_key)
    raw_text = ""
    try:
        response = client.responses.create(
            model=model,
            temperature=0.1,
            input=[
                {"role": "system", "content": [{"type": "text", "text": system_prompt}]},
                {"role": "user", "content": [{"type": "text", "text": user_prompt}]},
            ],
        )
        raw_text = getattr(response, "output_text", "") or ""
    except Exception:
        try:
            response = client.chat.completions.create(
                model=model,
                temperature=0.1,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
            raw_text = (response.choices[0].message.content or "").strip()
        except Exception:
            return None

    payload_text = _extract_json_payload(raw_text)
    if not payload_text:
        return None
    try:
        return json.loads(payload_text)
    except json.JSONDecodeError:
        return None


def _normalize_objective_ai_rows(rows):
    normalized = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        stem = re.sub(r"\s+", " ", str(row.get("stem") or "")).strip()
        options = row.get("options") or {}
        if not stem or not isinstance(options, dict):
            continue
        option_map = {}
        for label in ["A", "B", "C", "D"]:
            value = re.sub(r"\s+", " ", str(options.get(label, "") or "")).strip()
            if value:
                option_map[label] = value
        if len(option_map) < 4:
            continue
        correct_label = str(row.get("correct_label") or "").strip().upper()
        if correct_label not in option_map:
            correct_label = "A"
        normalized.append(
            {
                "stem": stem,
                "options": {label: option_map[label] for label in ["A", "B", "C", "D"]},
                "correct_label": correct_label,
            }
        )
    return normalized


def _normalize_parsed_question_rows(rows):
    normalized = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        question_type = str(row.get("question_type") or "OBJECTIVE").strip().upper()
        stem = re.sub(r"\s+", " ", str(row.get("stem") or "")).strip()
        if not stem:
            continue
        if question_type == "OBJECTIVE":
            options = row.get("options") or {}
            if not isinstance(options, dict):
                continue
            option_map = {}
            for label in ["A", "B", "C", "D"]:
                value = re.sub(r"\s+", " ", str(options.get(label, "") or "")).strip()
                if value:
                    option_map[label] = value
            if len(option_map) < 4:
                continue
            correct_label = str(row.get("correct_label") or "").strip().upper()
            if correct_label and correct_label not in option_map:
                correct_label = ""
            normalized.append(
                {
                    "question_type": "OBJECTIVE",
                    "stem": stem,
                    "options": option_map,
                    "correct_label": correct_label,
                }
            )
            continue
        model_answer = re.sub(r"\s+", " ", str(row.get("model_answer") or "")).strip()
        normalized.append(
            {
                "question_type": "THEORY",
                "stem": stem,
                "model_answer": model_answer,
            }
        )
    return normalized


def _generate_ai_question_blocks_with_openai(*, subject_name, topic, question_count, lesson_note_text=""):
    lesson_note_excerpt = _normalize_extracted_text(lesson_note_text)[:12000]
    user_prompt = (
        "Create objective CBT questions in strict JSON format.\n"
        "Return ONLY JSON in this format:\n"
        '{"questions":[{"stem":"...","options":{"A":"...","B":"...","C":"...","D":"..."},"correct_label":"A"}]}\n'
        f"Subject: {subject_name}\n"
        f"Topic: {topic}\n"
        f"Question count: {max(int(question_count or 0), 1)}\n"
        "Constraints:\n"
        "- Questions must be classroom-level and solvable.\n"
        "- If subject is math, use calculation/problem-solving style where possible.\n"
        "- Options must be realistic distractors.\n"
        "- correct_label must be one of A,B,C,D.\n"
    )
    if lesson_note_excerpt.strip():
        user_prompt += f"\nLesson note/context:\n{lesson_note_excerpt}"

    data = _openai_json_response(
        system_prompt=(
            "You are an NDGA exam drafting assistant. "
            "Return strict JSON only. Do not include markdown."
        ),
        user_prompt=user_prompt,
    )
    if not isinstance(data, dict):
        return []
    rows = _normalize_objective_ai_rows(data.get("questions") or [])
    if not rows:
        return []
    target_count = max(int(question_count or 0), 1)
    return rows[:target_count]


def _parse_questions_with_openai(*, extracted_text, subject_name, expected_count=0):
    text_excerpt = _normalize_extracted_text(extracted_text)[:18000]
    if not text_excerpt.strip():
        return []
    user_prompt = (
        "Parse this question document text into JSON.\n"
        "Return ONLY valid JSON in this exact shape:\n"
        '{"questions":[{"question_type":"OBJECTIVE","stem":"...","options":{"A":"...","B":"...","C":"...","D":"..."},"correct_label":"A"},'
        '{"question_type":"THEORY","stem":"...","model_answer":"..."}]}\n'
        "Rules:\n"
        "- Keep original question meaning.\n"
        "- If and only if options A-D exist, output OBJECTIVE.\n"
        "- If no options, output THEORY.\n"
        "- For OBJECTIVE rows, include all options A, B, C, D.\n"
        "- If answer key exists, set correct_label to A/B/C/D.\n"
        "- Do not invent answers if none is provided.\n"
        "- Do not add markdown, comments, or prose.\n"
        f"Subject: {subject_name}\n"
    )
    if expected_count:
        user_prompt += f"Expected around {expected_count} questions.\n"
    user_prompt += f"\nSource text:\n{text_excerpt}"

    data = _openai_json_response(
        system_prompt=(
            "You are an NDGA document-to-CBT parser. "
            "Return strict JSON only. Do not include markdown."
        ),
        user_prompt=user_prompt,
    )
    if not isinstance(data, dict):
        return []
    return _normalize_parsed_question_rows(data.get("questions") or [])


def _parse_questions_linewise(normalized_text):
    source_rows = [str(row or "") for row in (normalized_text or "").splitlines()]
    lines = []
    for row in source_rows:
        expanded = row
        expanded = INLINE_OPTION_BOUNDARY_PATTERN.sub("\n", expanded)
        expanded = INLINE_ANSWER_BOUNDARY_PATTERN.sub("\n", expanded)
        expanded = INLINE_QUESTION_BOUNDARY_PATTERN.sub("\n", expanded)
        expanded = re.sub(r"(?im)(?<!\n)\(([A-Da-d])\)\s*", r"\n\1) ", expanded)
        expanded = re.sub(r"(?im)(?<!\n)([A-Da-d])\s*[\)\.\:-]\s*", r"\n\1) ", expanded)
        for piece in expanded.splitlines():
            cleaned = re.sub(r"\s+", " ", piece.strip())
            if cleaned:
                lines.append(cleaned)

    parsed = []
    current = None
    pending_stem_lines = []
    last_option_label = None
    in_answer_block = False

    def _new_question(stem_text=""):
        return {
            "stem": (stem_text or "").strip(),
            "options": {},
            "answer_raw": "",
        }

    def _append_stem(target, text):
        cleaned = (text or "").strip()
        if not cleaned:
            return
        if target.get("stem"):
            target["stem"] = f"{target['stem']} {cleaned}".strip()
        else:
            target["stem"] = cleaned

    def flush_current():
        nonlocal current, last_option_label, in_answer_block
        if not current:
            return
        options = current.get("options", {})
        normalized_options = {
            label: options[label]
            for label in ["A", "B", "C", "D"]
            if options.get(label)
        }
        stem = current.get("stem", "").strip()
        answer_raw = (current.get("answer_raw") or "").strip()
        if stem and len(normalized_options) >= 4:
            parsed.append(
                {
                    "question_type": "OBJECTIVE",
                    "stem": stem,
                    "options": normalized_options,
                    "correct_label": _resolve_correct_label_from_answer(
                        options=normalized_options,
                        answer_raw=answer_raw,
                    ),
                }
            )
        elif stem:
            parsed.append(
                {
                    "question_type": "THEORY",
                    "stem": stem,
                    "model_answer": answer_raw,
                }
            )
        current = None
        last_option_label = None
        in_answer_block = False

    for line in lines:
        if not line:
            continue
        line = re.sub(r"^[\-\u2022\u25AA\u25CF\u00B7]+\s*", "", line).strip()
        if not line:
            continue

        numbered_match = QUESTION_START_PATTERN.match(line)
        prefixed_match = QUESTION_PREFIX_PATTERN.match(line)
        question_match = numbered_match or prefixed_match
        if question_match:
            flush_current()
            stem_from_line = (question_match.group(2) or "").strip()
            if not stem_from_line and pending_stem_lines:
                stem_from_line = " ".join(pending_stem_lines).strip()
                pending_stem_lines = []
            current = _new_question(stem_from_line)
            continue

        option_match = OPTION_PATTERN.match(line)
        if option_match is None:
            option_match = re.match(r"^\s*([A-Da-d])\s+(.+)$", line)
        if current is None:
            if option_match:
                stem = " ".join(pending_stem_lines).strip()
                pending_stem_lines = []
                current = _new_question(stem)
            else:
                pending_stem_lines.append(line)
                continue

        answer_match = ANSWER_PATTERN.match(line)
        if answer_match:
            answer_text = (answer_match.group(1) or "").strip()
            if answer_text:
                existing_answer = current.get("answer_raw", "")
                current["answer_raw"] = (
                    f"{existing_answer} {answer_text}".strip() if existing_answer else answer_text
                )
            in_answer_block = True
            last_option_label = None
            continue

        if option_match and not in_answer_block:
            label = option_match.group(1).upper()
            option_text = option_match.group(2).strip()
            existing = current["options"].get(label, "")
            current["options"][label] = (
                f"{existing} {option_text}".strip() if existing else option_text
            )
            last_option_label = label
            continue

        if in_answer_block:
            if current.get("options") and re.match(r"^[A-Za-z0-9(]", line):
                flush_current()
                pending_stem_lines = [line]
                continue
            existing_answer = current.get("answer_raw", "")
            current["answer_raw"] = (
                f"{existing_answer} {line}".strip() if existing_answer else line
            )
            continue

        if last_option_label:
            current["options"][last_option_label] = (
                f"{current['options'][last_option_label]} {line}"
            ).strip()
            continue

        if current.get("options"):
            flush_current()
            pending_stem_lines = [line]
            continue

        _append_stem(current, line)

    flush_current()
    if pending_stem_lines:
        stem = " ".join(pending_stem_lines).strip()
        if stem:
            parsed.append(
                {
                    "question_type": "THEORY",
                    "stem": stem,
                    "model_answer": "",
                }
            )
    return parsed


def _is_low_confidence_parse(rows, *, normalized_text):
    if not rows:
        return True
    objective_count = sum(1 for row in rows if row.get("question_type") == "OBJECTIVE")
    normalized_compact_len = len(re.sub(r"\s+", "", normalized_text or ""))
    if len(rows) == 1 and objective_count == 0:
        only_stem = (rows[0].get("stem") or "").strip()
        if len(only_stem) > 200:
            return True
        if normalized_compact_len > 260:
            return True
    if objective_count == 0 and len(rows) <= 2 and normalized_compact_len > 700:
        return True
    return False


def _objective_marker_count(text):
    return len(
        re.findall(r"(?im)(?:^|\s)[A-Da-d]\s*[\)\.\:-]\s*", str(text or ""))
    )


def _parsed_rows_quality(rows):
    objective_count = sum(1 for row in rows if row.get("question_type") == "OBJECTIVE")
    theory_count = sum(1 for row in rows if row.get("question_type") != "OBJECTIVE")
    stem_chars = sum(len((row.get("stem") or "").strip()) for row in rows)
    return (objective_count * 100) + (len(rows) * 12) + min(stem_chars, 4000) / 80 - (theory_count * 2)


def parse_question_document(extracted_text):
    normalized_text = _normalize_extracted_text(extracted_text)
    if not normalized_text:
        return {
            "normalized_text": "",
            "parsed_questions": [],
            "flagged_blocks": [],
            "used_line_fallback": False,
            "low_confidence": True,
        }

    blocks = _split_into_question_blocks(normalized_text)
    block_rows = []
    flagged_blocks = []
    used_line_fallback = False

    if len(blocks) > 1:
        for block in blocks:
            parsed = _parse_question_block(block.body)
            if parsed:
                block_rows.append(parsed)
            else:
                flagged_blocks.append(
                    {
                        "number": block.number,
                        "body": block.body,
                        "reason": "deterministic_block_parse_failed",
                    }
                )
    else:
        used_line_fallback = True
        block_rows = _parse_questions_linewise(normalized_text)
        if not block_rows:
            flagged_blocks.append(
                {
                    "number": "?",
                    "body": normalized_text,
                    "reason": "linewise_parse_failed",
                }
            )

    line_rows = _parse_questions_linewise(normalized_text)
    parsed_rows = block_rows
    if line_rows:
        block_quality = _parsed_rows_quality(block_rows)
        line_quality = _parsed_rows_quality(line_rows)
        block_objective = sum(1 for row in block_rows if row.get("question_type") == "OBJECTIVE")
        line_objective = sum(1 for row in line_rows if row.get("question_type") == "OBJECTIVE")
        option_markers = _objective_marker_count(normalized_text)

        should_use_line_rows = (
            line_quality > block_quality
            or (
                option_markers >= 8
                and block_objective == 0
                and line_objective > 0
            )
            or (
                len(line_rows) >= max(len(block_rows) + 2, 4)
                and line_objective >= block_objective
            )
        )
        if should_use_line_rows:
            parsed_rows = line_rows
            used_line_fallback = True
            # If linewise parser produced strong structure, ignore block-parse flags.
            if line_quality >= block_quality:
                flagged_blocks = []

    low_confidence = _is_low_confidence_parse(parsed_rows, normalized_text=normalized_text)
    if low_confidence and normalized_text:
        flagged_blocks.append(
            {
                "number": "?",
                "body": normalized_text[:2500],
                "reason": "low_confidence_parse",
            }
        )
        parsed_rows = []

    return {
        "normalized_text": normalized_text,
        "parsed_questions": parsed_rows,
        "flagged_blocks": flagged_blocks,
        "used_line_fallback": used_line_fallback,
        "low_confidence": low_confidence,
    }


def _repair_flagged_blocks_with_openai(*, flagged_blocks, subject_name):
    if not flagged_blocks:
        return []
    repair_text_parts = []
    for row in flagged_blocks:
        body = (row.get("body") or "").strip()
        if not body:
            continue
        block_no = row.get("number") or "?"
        repair_text_parts.append(f"BLOCK {block_no}:\n{body}")
    repair_text = "\n\n".join(repair_text_parts).strip()
    if not repair_text:
        return []
    return _parse_questions_with_openai(
        extracted_text=repair_text,
        subject_name=subject_name,
        expected_count=len(flagged_blocks),
    )


def parse_objective_questions(extracted_text):
    return parse_question_document(extracted_text).get("parsed_questions") or []


def _truncate_text(value, *, max_length=180):
    text = re.sub(r"\s+", " ", (value or "").strip())
    if len(text) <= max_length:
        return text
    return f"{text[: max_length - 3].rstrip()}..."


def _extract_lesson_note_points(note_text, *, limit=120):
    raw = (note_text or "").strip()
    if not raw:
        return []
    candidates = []
    for row in re.split(r"[\r\n]+", raw):
        chunk = row.strip()
        if not chunk:
            continue
        for sentence in re.split(r"(?<=[\.\?\!;:])\s+", chunk):
            cleaned = re.sub(r"\s+", " ", sentence).strip(" -\t")
            if len(cleaned) >= 14:
                candidates.append(cleaned)
    deduped = []
    seen = set()
    for row in candidates:
        key = row.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
        if len(deduped) >= limit:
            break
    return deduped


def _build_ai_question_options(*, subject_name, topic, rng, lesson_point=""):
    subject = (subject_name or "").lower()
    topic_text = topic.strip()
    lesson_point = _truncate_text(lesson_point, max_length=170)
    if lesson_point:
        correct = lesson_point
        distractors = [
            f"A statement that is not supported by the class note on {topic_text}.",
            f"An incorrect interpretation of {topic_text} from the note.",
            f"A claim unrelated to the taught point on {topic_text}.",
        ]
    elif "math" in subject:
        correct = f"The correct worked outcome for {topic_text}."
        distractors = [
            f"A wrong estimate for {topic_text}.",
            f"A value that ignores the correct method for {topic_text}.",
            f"An answer with incorrect operation order for {topic_text}.",
        ]
    elif "physics" in subject or "chem" in subject or "bio" in subject or "science" in subject:
        correct = f"The scientifically valid explanation for {topic_text}."
        distractors = [
            f"A claim that confuses core principles of {topic_text}.",
            f"A statement that ignores lab evidence for {topic_text}.",
            f"An interpretation not supported by observed results in {topic_text}.",
        ]
    else:
        correct = f"The best supported answer about {topic_text}."
        distractors = [
            f"A partially correct statement about {topic_text}.",
            f"An unrelated interpretation of {topic_text}.",
            f"A misleading summary for {topic_text}.",
        ]
    normalized_correct = _truncate_text(correct, max_length=180)
    option_texts = [normalized_correct] + [
        _truncate_text(item, max_length=180) for item in distractors
    ]
    rng.shuffle(option_texts)
    labels = ["A", "B", "C", "D"]
    options = {label: option_texts[idx] for idx, label in enumerate(labels)}
    correct_label = labels[option_texts.index(normalized_correct)]
    return options, correct_label


def generate_ai_question_blocks(*, subject_name, topic, question_count, lesson_note_text=""):
    templates = [
        "Which option best matches what was taught about {topic}?",
        "Based on class teaching, identify the accurate statement for {topic}.",
        "Which choice reflects the correct understanding of {topic}?",
        "Select the best supported statement about {topic}.",
        "From your class note, which option is correct for {topic}?",
        "Choose the option most consistent with the taught point on {topic}.",
    ]
    seed = f"{subject_name}|{topic}|{question_count}".lower()
    rng = random.Random(seed)
    lesson_points = _extract_lesson_note_points(lesson_note_text)
    rows = []
    for index in range(max(int(question_count or 0), 1)):
        stem_template = templates[index % len(templates)]
        lesson_point = lesson_points[index % len(lesson_points)] if lesson_points else ""
        if lesson_point:
            point_preview = _truncate_text(lesson_point, max_length=120)
            stem = (
                f"{stem_template.format(topic=topic.strip(), subject=subject_name)} "
                f"Reference point: {point_preview}"
            )
        else:
            stem = stem_template.format(topic=topic.strip(), subject=subject_name)
        options, correct_label = _build_ai_question_options(
            subject_name=subject_name,
            topic=topic,
            rng=rng,
            lesson_point=lesson_point,
        )
        rows.append(
            {
                "stem": f"{index + 1}. {stem}",
                "options": options,
                "correct_label": correct_label,
            }
        )
    return rows


def _default_section_totals(*, exam_type, flow_type):
    if exam_type == CBTExamType.FREE_TEST:
        return Decimal("100.00"), Decimal("0.00")
    if exam_type in {CBTExamType.CA, CBTExamType.PRACTICAL}:
        if flow_type == "OBJECTIVE_THEORY":
            return Decimal("5.00"), Decimal("5.00")
        if flow_type == "THEORY_ONLY":
            return Decimal("0.00"), Decimal("10.00")
        return Decimal("10.00"), Decimal("0.00")
    if exam_type == CBTExamType.EXAM:
        if flow_type == "OBJECTIVE_THEORY":
            return Decimal("40.00"), Decimal("20.00")
        if flow_type == "THEORY_ONLY":
            return Decimal("0.00"), Decimal("20.00")
        return Decimal("40.00"), Decimal("0.00")
    return Decimal("10.00"), Decimal("0.00")


def is_free_test_exam(exam):
    return bool(
        getattr(exam, "is_free_test", False)
        or getattr(exam, "exam_type", "") == CBTExamType.FREE_TEST
    )


def _normalized_ca_target_for_flow(*, ca_target, flow_type):
    target = (ca_target or "").strip()
    if flow_type == "OBJECTIVE_THEORY" and target == CBTWritebackTarget.CA3:
        target = CBTWritebackTarget.CA2
    return target


def has_completed_ca_target_exam(*, assignment, ca_target, flow_type="OBJECTIVE_THEORY", exclude_exam_id=None):
    target = _normalized_ca_target_for_flow(ca_target=ca_target, flow_type=flow_type)
    if target not in {
        CBTWritebackTarget.CA1,
        CBTWritebackTarget.CA2,
        CBTWritebackTarget.CA3,
        CBTWritebackTarget.CA4,
    }:
        return False
    candidate_exams = (
        Exam.objects.filter(
            assignment=assignment,
            exam_type__in=[CBTExamType.CA, CBTExamType.PRACTICAL, CBTExamType.SIM],
        )
        .exclude(id=exclude_exam_id)
        .select_related("blueprint")
    )
    for exam in candidate_exams:
        blueprint = getattr(exam, "blueprint", None)
        section_config = getattr(blueprint, "section_config", {}) if blueprint else {}
        if not isinstance(section_config, dict):
            section_config = {}
        exam_target = (section_config.get("ca_target") or "").strip()
        if not exam_target:
            exam_target = (
                CBTWritebackTarget.CA4
                if exam.exam_type == CBTExamType.PRACTICAL
                else CBTWritebackTarget.CA1
            )
        exam_flow_type = (
            section_config.get("flow_type")
            or ("SIMULATION" if exam.exam_type == CBTExamType.SIM else "OBJECTIVE_THEORY")
        )
        exam_target = _normalized_ca_target_for_flow(
            ca_target=exam_target,
            flow_type=exam_flow_type,
        )
        if exam_target != target:
            continue
        if exam.status == CBTExamStatus.CLOSED:
            return True
        if exam.attempts.filter(
            status__in=[CBTAttemptStatus.SUBMITTED, CBTAttemptStatus.FINALIZED]
        ).exists():
            return True
    return False


def _distributed_section_marks(*, target_total, question_count):
    if question_count <= 0:
        return []
    total = Decimal(str(target_total or 0)).quantize(Decimal("0.01"))
    if total <= 0:
        return [Decimal("1.00")] * question_count
    per_mark = (total / Decimal(question_count)).quantize(Decimal("0.01"))
    if per_mark <= 0:
        per_mark = Decimal("0.01")
    distributed = [per_mark for _ in range(question_count)]
    remainder = (total - (per_mark * Decimal(question_count - 1))).quantize(Decimal("0.01"))
    if remainder > 0:
        distributed[-1] = remainder
    return distributed


def _apply_exam_row_marks(*, exam, objective_total, theory_total):
    objective_rows = list(
        exam.exam_questions.select_related("question")
        .filter(question__question_type__in=OBJECTIVE_TYPES)
        .order_by("sort_order", "id")
    )
    theory_rows = list(
        exam.exam_questions.select_related("question")
        .exclude(question__question_type__in=OBJECTIVE_TYPES)
        .order_by("sort_order", "id")
    )
    objective_marks = _distributed_section_marks(
        target_total=objective_total,
        question_count=len(objective_rows),
    )
    theory_marks = _distributed_section_marks(
        target_total=theory_total,
        question_count=len(theory_rows),
    )
    for row, mark in zip(objective_rows, objective_marks):
        row.marks = mark
        row.save(update_fields=["marks", "updated_at"])
        row.question.marks = mark
        row.question.save(update_fields=["marks", "updated_at"])
    for row, mark in zip(theory_rows, theory_marks):
        row.marks = mark
        row.save(update_fields=["marks", "updated_at"])
        row.question.marks = mark
        row.question.save(update_fields=["marks", "updated_at"])


def _extract_text_from_uploaded_file(uploaded_file):
    if not uploaded_file:
        return ""
    suffix = Path(getattr(uploaded_file, "name", "")).suffix.lower() or ".txt"
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
            temp_path = handle.name
            for chunk in uploaded_file.chunks():
                handle.write(chunk)
        return extract_text_from_document(temp_path)
    finally:
        if temp_path:
            Path(temp_path).unlink(missing_ok=True)


def _should_write_import_debug_artifacts():
    value = (
        getattr(settings, "CBT_IMPORT_DEBUG", "")
        or os.getenv("CBT_IMPORT_DEBUG", "")
    )
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def _write_import_debug_artifacts(
    *,
    import_row,
    raw_text,
    normalized_text,
    parsed_rows,
    flagged_blocks=None,
):
    if not _should_write_import_debug_artifacts():
        return {}
    media_root = Path(settings.MEDIA_ROOT)
    debug_root = media_root / "cbt" / "import_debug" / str(import_row.id)
    debug_root.mkdir(parents=True, exist_ok=True)

    raw_path = debug_root / "raw_text.txt"
    normalized_path = debug_root / "normalized_text.txt"
    blocks_path = debug_root / "blocks.json"

    raw_path.write_text(str(raw_text or ""), encoding="utf-8", errors="ignore")
    normalized_path.write_text(str(normalized_text or ""), encoding="utf-8", errors="ignore")
    blocks_payload = _split_into_question_blocks(normalized_text or "")
    blocks_rows = [
        {
            "number": row.number,
            "body": row.body,
        }
        for row in blocks_payload
    ]
    blocks_path.write_text(
        json.dumps(
            {
                "generated_at": timezone.now().isoformat(),
                "block_count": len(blocks_rows),
                "parsed_question_count": len(parsed_rows or []),
                "flagged_block_count": len(flagged_blocks or []),
                "blocks": blocks_rows,
                "flagged_blocks": flagged_blocks or [],
                "parsed_rows": parsed_rows or [],
            },
            indent=2,
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    return {
        "debug_enabled": True,
        "debug_dir": str(debug_root.relative_to(media_root)).replace("\\", "/"),
        "raw_text_file": str(raw_path.relative_to(media_root)).replace("\\", "/"),
        "normalized_text_file": str(normalized_path.relative_to(media_root)).replace("\\", "/"),
        "blocks_file": str(blocks_path.relative_to(media_root)).replace("\\", "/"),
    }


def ensure_default_blueprint(exam):
    blueprint, _ = ExamBlueprint.objects.get_or_create(
        exam=exam,
        defaults={
            "duration_minutes": 60,
            "max_attempts": 1,
            "shuffle_questions": True,
            "shuffle_options": True,
            "instructions": "",
            "section_config": [],
            "passing_score": 0,
        },
    )
    return blueprint


@transaction.atomic
def build_exam_with_ai_draft(
    *,
    actor,
    assignment,
    title,
    topic,
    question_count,
    exam_type,
    ca_target="",
    difficulty=CBTQuestionDifficulty.MEDIUM,
    lesson_note_text="",
    lesson_note_file=None,
):
    note_from_file = _extract_text_from_uploaded_file(lesson_note_file)
    combined_note = "\n".join(
        part.strip()
        for part in [lesson_note_text or "", note_from_file or ""]
        if part and part.strip()
    )
    bank_name = f"AI Draft Bank - {title}"
    question_bank, _ = QuestionBank.objects.get_or_create(
        owner=actor,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        name=bank_name,
        defaults={"description": f"AI-assisted draft set for {topic}."},
    )

    exam = Exam.objects.create(
        title=title,
        description=(
            f"AI-assisted draft generated for topic: {topic}."
            if not combined_note
            else f"AI-assisted draft generated for topic: {topic} using uploaded/pasted lesson note context."
        ),
        exam_type=exam_type,
        status=CBTExamStatus.DRAFT,
        created_by=actor,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        question_bank=question_bank,
        is_free_test=(exam_type == CBTExamType.FREE_TEST),
    )
    ensure_default_blueprint(exam)

    target_count = max(int(question_count or 0), 1)
    if exam_type != CBTExamType.FREE_TEST:
        target_count = max(target_count, 2)
    planned_theory_count = (
        0
        if exam_type == CBTExamType.FREE_TEST
        else max(1, min(5, target_count // 4 or 1))
    )
    planned_objective_count = (
        target_count if exam_type == CBTExamType.FREE_TEST else max(1, target_count - planned_theory_count)
    )
    generated = []
    seen_stems = set()

    def _append_generated_rows(rows):
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            stem = re.sub(r"\s+", " ", str(row.get("stem") or "")).strip()
            options = row.get("options") or {}
            if not stem or not isinstance(options, dict):
                continue
            option_map = {}
            for label in ("A", "B", "C", "D"):
                value = re.sub(r"\s+", " ", str(options.get(label) or "")).strip()
                if value:
                    option_map[label] = value
            if len(option_map) < 4:
                continue
            correct_label = str(row.get("correct_label") or "").strip().upper()
            if correct_label not in option_map:
                correct_label = "A"
            key = stem.lower()
            if key in seen_stems:
                continue
            seen_stems.add(key)
            generated.append(
                {
                    "stem": stem,
                    "options": {label: option_map[label] for label in ("A", "B", "C", "D")},
                    "correct_label": correct_label,
                }
            )
            if len(generated) >= target_count:
                break

    # First: if teacher materials already contain objective questions, use them directly.
    if combined_note.strip():
        note_payload = parse_question_document(combined_note)
        objective_rows = [
            {
                "stem": row.get("stem") or "",
                "options": row.get("options") or {},
                "correct_label": row.get("correct_label") or "",
            }
            for row in (note_payload.get("parsed_questions") or [])
            if str(row.get("question_type") or "").upper() == "OBJECTIVE"
        ]
        _append_generated_rows(objective_rows)

    remaining = target_count - len(generated)
    if remaining > 0:
        ai_rows = _generate_ai_question_blocks_with_openai(
            subject_name=assignment.subject.name,
            topic=topic,
            question_count=remaining,
            lesson_note_text=combined_note,
        )
        _append_generated_rows(ai_rows)

    remaining = target_count - len(generated)
    if remaining > 0:
        deterministic_rows = generate_ai_question_blocks(
            subject_name=assignment.subject.name,
            topic=topic,
            question_count=remaining,
            lesson_note_text=combined_note,
        )
        _append_generated_rows(deterministic_rows)

    if not generated:
        raise ValidationError(
            "AI draft could not generate valid objective questions from the provided materials."
        )

    created_objective_count = 0
    created_theory_count = 0
    for index, row in enumerate(generated[:target_count], start=1):
        if index <= planned_objective_count:
            question = Question.objects.create(
                question_bank=question_bank,
                created_by=actor,
                subject=assignment.subject,
                question_type=CBTQuestionType.OBJECTIVE,
                stem=row["stem"],
                topic=topic,
                difficulty=difficulty,
                marks=Decimal("1.00"),
                source_type=Question.SourceType.MANUAL,
                source_reference="AI_DRAFT_NOTE" if combined_note else "AI_DRAFT",
            )
            ordered_labels = ["A", "B", "C", "D"]
            for sort_order, label in enumerate(ordered_labels, start=1):
                Option.objects.create(
                    question=question,
                    label=label,
                    option_text=row["options"][label],
                    sort_order=sort_order,
                )
            correct = CorrectAnswer.objects.create(question=question, is_finalized=True)
            correct.correct_options.set(question.options.filter(label=row["correct_label"]))
            ExamQuestion.objects.create(
                exam=exam,
                question=question,
                sort_order=index,
                marks=Decimal("1.00"),
            )
            created_objective_count += 1
            continue

        theory_question = Question.objects.create(
            question_bank=question_bank,
            created_by=actor,
            subject=assignment.subject,
            question_type=CBTQuestionType.SHORT_ANSWER,
            stem=row["stem"],
            topic=topic,
            difficulty=difficulty,
            marks=Decimal("1.00"),
            source_type=Question.SourceType.MANUAL,
            source_reference="THEORY_MODE:PAPER",
        )
        CorrectAnswer.objects.create(question=theory_question, is_finalized=False)
        ExamQuestion.objects.create(
            exam=exam,
            question=theory_question,
            sort_order=index,
            marks=Decimal("1.00"),
        )
        created_theory_count += 1

    flow_type = "OBJECTIVE_ONLY" if exam_type == CBTExamType.FREE_TEST else "OBJECTIVE_THEORY"
    objective_total, theory_total = _default_section_totals(
        exam_type=exam_type,
        flow_type=flow_type,
    )
    blueprint = ensure_default_blueprint(exam)
    if exam_type == CBTExamType.FREE_TEST:
        objective_target = CBTWritebackTarget.NONE
        theory_target = CBTWritebackTarget.NONE
        effective_ca_target = ""
    elif exam_type in {CBTExamType.CA, CBTExamType.PRACTICAL}:
        default_ca_target = (
            CBTWritebackTarget.CA4
            if exam_type == CBTExamType.PRACTICAL
            else CBTWritebackTarget.CA1
        )
        effective_ca_target = (ca_target or default_ca_target).strip()
        if effective_ca_target == CBTWritebackTarget.CA3:
            effective_ca_target = CBTWritebackTarget.CA2
        if effective_ca_target == CBTWritebackTarget.CA2:
            objective_target = CBTWritebackTarget.CA2
            theory_target = CBTWritebackTarget.CA3
            objective_total = Decimal("10.00")
            theory_total = Decimal("10.00")
        else:
            objective_target = effective_ca_target
            theory_target = effective_ca_target
    else:
        objective_target = CBTWritebackTarget.OBJECTIVE
        theory_target = CBTWritebackTarget.THEORY
        effective_ca_target = ""
    blueprint.section_config = {
        "flow_type": flow_type,
        "objective_count": created_objective_count,
        "theory_count": created_theory_count,
        "theory_response_mode": "PAPER",
        "ca_target": effective_ca_target,
        "is_free_test": exam_type == CBTExamType.FREE_TEST,
        "manual_score_split": False,
        "objective_target_max": str(objective_total),
        "theory_target_max": str(theory_total),
    }
    blueprint.objective_writeback_target = objective_target
    blueprint.theory_enabled = created_theory_count > 0
    blueprint.theory_writeback_target = theory_target
    blueprint.save(
        update_fields=[
            "section_config",
            "objective_writeback_target",
            "theory_enabled",
            "theory_writeback_target",
            "updated_at",
        ]
    )
    _apply_exam_row_marks(
        exam=exam,
        objective_total=objective_total,
        theory_total=theory_total,
    )

    return exam, (created_objective_count + created_theory_count)


@transaction.atomic
def build_exam_from_uploaded_document(
    *,
    actor,
    assignment,
    title,
    exam_type,
    ca_target="",
    source_file,
):
    import_row = ExamDocumentImport.objects.create(
        uploaded_by=actor,
        assignment=assignment,
        source_file=source_file,
        source_filename=source_file.name,
        extraction_status=CBTDocumentStatus.PENDING,
    )

    bank_name = f"Imported Bank - {title}"
    question_bank, _ = QuestionBank.objects.get_or_create(
        owner=actor,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        name=bank_name,
        defaults={"description": "Generated from uploaded document."},
    )

    exam = Exam.objects.create(
        title=title,
        description="Generated from uploaded document.",
        exam_type=exam_type,
        status=CBTExamStatus.DRAFT,
        created_by=actor,
        assignment=assignment,
        subject=assignment.subject,
        academic_class=assignment.academic_class,
        session=assignment.session,
        term=assignment.term,
        question_bank=question_bank,
        is_free_test=(exam_type == CBTExamType.FREE_TEST),
    )
    ensure_default_blueprint(exam)

    try:
        extracted_text = extract_text_from_document(import_row.source_file.path)
        parse_payload = parse_question_document(extracted_text)
        normalized_text = parse_payload.get("normalized_text") or ""
        parsed_questions = list(parse_payload.get("parsed_questions") or [])
        flagged_blocks = list(parse_payload.get("flagged_blocks") or [])
        parser_used = "deterministic"

        repaired_rows = _repair_flagged_blocks_with_openai(
            flagged_blocks=flagged_blocks,
            subject_name=assignment.subject.name,
        )
        if repaired_rows:
            seen_keys = {
                (
                    (row.get("question_type") or "").strip(),
                    re.sub(r"\s+", " ", (row.get("stem") or "").strip()).lower(),
                )
                for row in parsed_questions
            }
            for row in repaired_rows:
                key = (
                    (row.get("question_type") or "").strip(),
                    re.sub(r"\s+", " ", (row.get("stem") or "").strip()).lower(),
                )
                if key in seen_keys or not key[1]:
                    continue
                seen_keys.add(key)
                parsed_questions.append(row)
            if parsed_questions:
                parser_used = "deterministic_plus_openai_repair"
            else:
                parser_used = "openai_repair_only"

        if not parsed_questions:
            parsed_questions = _parse_questions_with_openai(
                extracted_text=normalized_text,
                subject_name=assignment.subject.name,
            )
            parser_used = "openai_fallback_fulltext"
        if not parsed_questions:
            linewise_rows = _parse_questions_linewise(normalized_text)
            linewise_objective = sum(
                1 for row in linewise_rows if row.get("question_type") == "OBJECTIVE"
            )
            if linewise_rows and (
                linewise_objective > 0
                or len(linewise_rows) >= 4
                or _objective_marker_count(normalized_text) >= 8
            ):
                parsed_questions = linewise_rows
                parser_used = "linewise_last_resort"
        if not parsed_questions:
            raise ValidationError(
                "No valid questions detected. Ensure your document uses numbered questions and A-D options, then retry."
            )

        created_count = 0
        objective_count = 0
        theory_count = 0
        for index, row in enumerate(parsed_questions, start=1):
            question_type = row.get("question_type", "OBJECTIVE")
            question = Question.objects.create(
                question_bank=question_bank,
                created_by=actor,
                subject=assignment.subject,
                question_type=(
                    CBTQuestionType.OBJECTIVE
                    if question_type == "OBJECTIVE"
                    else CBTQuestionType.SHORT_ANSWER
                ),
                stem=row["stem"],
                topic="Imported",
                difficulty="MEDIUM",
                marks=1,
                source_type=Question.SourceType.DOCUMENT,
                source_reference=str(import_row.id),
            )
            if question_type == "OBJECTIVE":
                option_order = 1
                for label, text in row.get("options", {}).items():
                    Option.objects.create(
                        question=question,
                        label=label,
                        option_text=text,
                        sort_order=option_order,
                    )
                    option_order += 1
                answer = CorrectAnswer.objects.create(
                    question=question,
                    is_finalized=False,
                )
                correct_label = (row.get("correct_label") or "").strip().upper()
                if correct_label:
                    option_qs = question.options.filter(label=correct_label)
                    if option_qs.exists():
                        answer.correct_options.set(option_qs)
                        answer.is_finalized = True
                        answer.save(update_fields=["is_finalized", "updated_at"])
                objective_count += 1
            else:
                CorrectAnswer.objects.create(
                    question=question,
                    is_finalized=False,
                )
                theory_count += 1
            ExamQuestion.objects.create(
                exam=exam,
                question=question,
                sort_order=index,
                marks=question.marks,
            )
            created_count += 1

        if exam_type != CBTExamType.FREE_TEST:
            if objective_count <= 0:
                created_count += 1
                fallback_objective = Question.objects.create(
                    question_bank=question_bank,
                    created_by=actor,
                    subject=assignment.subject,
                    question_type=CBTQuestionType.OBJECTIVE,
                    stem="Objective Question 1",
                    topic="Imported",
                    difficulty="MEDIUM",
                    marks=1,
                    source_type=Question.SourceType.DOCUMENT,
                    source_reference=f"{import_row.id}:FALLBACK_OBJECTIVE",
                )
                for label, option_text, sort_order in (
                    ("A", "Option A", 1),
                    ("B", "Option B", 2),
                    ("C", "Option C", 3),
                    ("D", "Option D", 4),
                ):
                    Option.objects.create(
                        question=fallback_objective,
                        label=label,
                        option_text=option_text,
                        sort_order=sort_order,
                    )
                fallback_answer = CorrectAnswer.objects.create(
                    question=fallback_objective,
                    is_finalized=True,
                )
                fallback_answer.correct_options.set(
                    fallback_objective.options.filter(label="A")
                )
                ExamQuestion.objects.create(
                    exam=exam,
                    question=fallback_objective,
                    sort_order=created_count,
                    marks=fallback_objective.marks,
                )
                objective_count += 1

            if theory_count <= 0:
                created_count += 1
                fallback_theory = Question.objects.create(
                    question_bank=question_bank,
                    created_by=actor,
                    subject=assignment.subject,
                    question_type=CBTQuestionType.SHORT_ANSWER,
                    stem="Theory Question 1",
                    topic="Imported",
                    difficulty="MEDIUM",
                    marks=1,
                    source_type=Question.SourceType.DOCUMENT,
                    source_reference=f"{import_row.id}:FALLBACK_THEORY",
                )
                CorrectAnswer.objects.create(
                    question=fallback_theory,
                    is_finalized=False,
                )
                ExamQuestion.objects.create(
                    exam=exam,
                    question=fallback_theory,
                    sort_order=created_count,
                    marks=fallback_theory.marks,
                )
                theory_count += 1

        if exam_type == CBTExamType.FREE_TEST:
            flow_type = "OBJECTIVE_ONLY"
        else:
            flow_type = "OBJECTIVE_THEORY"

        blueprint = ensure_default_blueprint(exam)
        objective_total, theory_total = _default_section_totals(
            exam_type=exam_type,
            flow_type=flow_type,
        )
        if exam_type == CBTExamType.FREE_TEST:
            objective_target = CBTWritebackTarget.NONE
            theory_target = CBTWritebackTarget.NONE
            effective_ca_target = ""
        elif exam_type in {CBTExamType.CA, CBTExamType.PRACTICAL}:
            default_ca_target = (
                CBTWritebackTarget.CA4
                if exam_type == CBTExamType.PRACTICAL
                else CBTWritebackTarget.CA1
            )
            effective_ca_target = (ca_target or default_ca_target).strip()
            if effective_ca_target == CBTWritebackTarget.CA3:
                effective_ca_target = CBTWritebackTarget.CA2
            if effective_ca_target == CBTWritebackTarget.CA2:
                objective_target = CBTWritebackTarget.CA2
                theory_target = CBTWritebackTarget.CA3
                objective_total = Decimal("10.00")
                theory_total = Decimal("10.00")
            else:
                objective_target = effective_ca_target
                theory_target = effective_ca_target
        else:
            objective_target = CBTWritebackTarget.OBJECTIVE
            theory_target = CBTWritebackTarget.THEORY
            effective_ca_target = ""
        blueprint.section_config = {
            "flow_type": flow_type,
            "objective_count": objective_count,
            "theory_count": theory_count,
            "theory_response_mode": "PAPER",
            "ca_target": effective_ca_target,
            "is_free_test": exam_type == CBTExamType.FREE_TEST,
            "manual_score_split": False,
            "objective_target_max": str(objective_total),
            "theory_target_max": str(theory_total),
        }
        blueprint.objective_writeback_target = objective_target
        blueprint.theory_enabled = theory_count > 0
        blueprint.theory_writeback_target = theory_target
        blueprint.save(
            update_fields=[
                "section_config",
                "objective_writeback_target",
                "theory_enabled",
                "theory_writeback_target",
                "updated_at",
            ]
        )
        _apply_exam_row_marks(
            exam=exam,
            objective_total=objective_total,
            theory_total=theory_total,
        )

        import_row.exam = exam
        import_row.extraction_status = CBTDocumentStatus.SUCCESS
        import_row.extracted_text = extracted_text
        debug_metadata = _write_import_debug_artifacts(
            import_row=import_row,
            raw_text=extracted_text,
            normalized_text=normalized_text,
            parsed_rows=parsed_questions,
            flagged_blocks=flagged_blocks,
        )
        import_row.parse_summary = {
            "parser_used": parser_used,
            "normalized_char_count": len(normalized_text),
            "flagged_block_count": len(flagged_blocks),
            "used_line_fallback": bool(parse_payload.get("used_line_fallback")),
            "low_confidence_detected": bool(parse_payload.get("low_confidence")),
            "question_count": created_count,
            "objective_count": objective_count,
            "theory_count": theory_count,
            "generated_at": timezone.now().isoformat(),
            **debug_metadata,
        }
        import_row.error_message = ""
        import_row.save(
            update_fields=[
                "exam",
                "extraction_status",
                "extracted_text",
                "parse_summary",
                "error_message",
                "updated_at",
            ]
        )
        return exam, import_row, created_count
    except ValidationError as exc:
        import_row.extraction_status = CBTDocumentStatus.FAILED
        import_row.error_message = str(exc)
        import_row.save(
            update_fields=["extraction_status", "error_message", "updated_at"]
        )
        raise
    except Exception as exc:
        import_row.extraction_status = CBTDocumentStatus.FAILED
        import_row.error_message = str(exc)
        import_row.save(
            update_fields=["extraction_status", "error_message", "updated_at"]
        )
        raise ValidationError(
            "Could not process uploaded file. Verify format and content."
        ) from exc


OBJECTIVE_TYPES = {CBTQuestionType.OBJECTIVE, CBTQuestionType.MULTI_SELECT}
THEORY_TYPES = {
    CBTQuestionType.SHORT_ANSWER,
}
STRUCTURED_AUTO_TYPES = {
    CBTQuestionType.LABELING,
    CBTQuestionType.MATCHING,
    CBTQuestionType.ORDERING,
}
DECIMAL_2 = Decimal("0.01")
SIMULATION_STUDENT_COMPLETED_STATUSES = {
    CBTSimulationAttemptStatus.AUTO_CAPTURED,
    CBTSimulationAttemptStatus.VERIFY_PENDING,
    CBTSimulationAttemptStatus.VERIFIED,
    CBTSimulationAttemptStatus.RUBRIC_PENDING,
    CBTSimulationAttemptStatus.RUBRIC_SCORED,
    CBTSimulationAttemptStatus.IMPORTED,
}


def _queue_attempt_snapshot(attempt, *, event_type):
    try:
        queue_exam_attempt_sync(attempt=attempt, event_type=event_type)
    except Exception:
        return None
    return None


def _queue_simulation_snapshot(record, *, event_type):
    try:
        queue_simulation_attempt_sync(record=record, event_type=event_type)
    except Exception:
        return None
    return None


def _to_decimal(value):
    try:
        if value is None or value == "":
            return Decimal("0.00")
        return Decimal(str(value)).quantize(DECIMAL_2)
    except (InvalidOperation, TypeError) as exc:
        raise ValidationError("Invalid score value.") from exc


def _target_max(target):
    if target in {
        CBTWritebackTarget.CA1,
        CBTWritebackTarget.CA2,
        CBTWritebackTarget.CA3,
        CBTWritebackTarget.CA4,
    }:
        return Decimal("10.00")
    if target == CBTWritebackTarget.OBJECTIVE:
        return Decimal("40.00")
    if target == CBTWritebackTarget.THEORY:
        return Decimal("20.00")
    return None


def _section_target_max(blueprint, *, section_key):
    config = getattr(blueprint, "section_config", {}) or {}
    if not isinstance(config, dict):
        return None
    raw_value = config.get(section_key)
    if raw_value in (None, ""):
        return None
    try:
        parsed = Decimal(str(raw_value)).quantize(DECIMAL_2)
    except (InvalidOperation, TypeError, ValueError):
        return None
    if parsed <= 0:
        return None
    return parsed


def _normalize_score(raw_score, raw_max_score, target, *, target_max_override=None):
    raw = _to_decimal(raw_score)
    raw_max = _to_decimal(raw_max_score)
    target_max = None
    if target_max_override not in (None, ""):
        try:
            target_max = _to_decimal(target_max_override)
        except ValidationError:
            target_max = None
    if target_max is None:
        target_max = _target_max(target)
    if target_max is None:
        return raw.quantize(DECIMAL_2)
    if raw_max <= 0:
        return Decimal("0.00")
    return ((raw / raw_max) * target_max).quantize(DECIMAL_2)


def _as_local_datetime(value):
    if value is None:
        return None
    if timezone.is_naive(value):
        value = timezone.make_aware(value, timezone.get_current_timezone())
    return timezone.localtime(value)


def exam_schedule_anchor(exam):
    return _as_local_datetime(exam.schedule_start or exam.activated_at or exam.created_at)


def exam_occurs_on_day(exam, target_day):
    if target_day is None:
        return True

    start = _as_local_datetime(exam.schedule_start)
    end = _as_local_datetime(exam.schedule_end)
    anchor = start or _as_local_datetime(exam.activated_at) or _as_local_datetime(exam.updated_at)

    if exam.status == CBTExamStatus.CLOSED:
        return bool(anchor and anchor.date() == target_day)

    if start and end:
        return start.date() <= target_day <= end.date()
    if start:
        return start.date() == target_day
    return bool(anchor and anchor.date() == target_day)


def _close_expired_exam(exam, *, now=None):
    now = now or timezone.now()
    if (
        exam.status != CBTExamStatus.ACTIVE
        or not exam.is_time_based
        or not exam.schedule_end
        or now <= exam.schedule_end
    ):
        return False

    exam.status = CBTExamStatus.CLOSED
    exam.open_now = False
    exam.save(update_fields=["status", "open_now", "updated_at"])
    ExamReviewAction.objects.create(
        exam=exam,
        actor=None,
        from_status=CBTExamStatus.ACTIVE,
        to_status=CBTExamStatus.CLOSED,
        action="AUTO_CLOSE",
        comment="Closed automatically after schedule window ended.",
    )
    return True


def close_expired_exams(*, now=None, exam_ids=None):
    now = now or timezone.now()
    exams = Exam.objects.filter(
        status=CBTExamStatus.ACTIVE,
        is_time_based=True,
        schedule_end__lt=now,
    ).order_by("schedule_end", "id")
    if exam_ids:
        exams = exams.filter(id__in=exam_ids)

    closed_ids = []
    for exam in exams:
        if _close_expired_exam(exam, now=now):
            closed_ids.append(exam.id)
    return closed_ids


def _is_exam_open_now(exam, now=None):
    now = now or timezone.now()
    if _close_expired_exam(exam, now=now):
        return False
    if exam.status != CBTExamStatus.ACTIVE:
        return False
    if not exam.is_time_based:
        return True
    if exam.open_now:
        if exam.schedule_end and now > exam.schedule_end:
            return False
        return True
    if not exam.schedule_start or not exam.schedule_end:
        return False
    return exam.schedule_start <= now <= exam.schedule_end


def _student_enrolled_for_exam(student, exam):
    if not student.has_role(ROLE_STUDENT):
        return False

    class_enrolled = StudentClassEnrollment.objects.filter(
        student=student,
        academic_class_id__in=exam.academic_class.cohort_class_ids(),
        session=exam.session,
        is_active=True,
    ).exists()
    if not class_enrolled:
        return False

    any_subject_for_session = StudentSubjectEnrollment.objects.filter(
        student=student,
        session=exam.session,
        is_active=True,
    ).exists()
    if not any_subject_for_session:
        return True
    return StudentSubjectEnrollment.objects.filter(
        student=student,
        subject=exam.subject,
        session=exam.session,
        is_active=True,
    ).exists()


def _allowed_attempt_count(blueprint):
    if blueprint.allow_retake:
        return max(blueprint.max_attempts, 1)
    return 1


def student_exam_authorization_reason(*, student, exam, now=None):
    now = now or timezone.now()
    _close_expired_exam(exam, now=now)
    if exam.status != CBTExamStatus.ACTIVE:
        return "Exam is not active."
    if not _is_exam_open_now(exam, now=now):
        return "Exam is outside schedule window."
    if not _student_enrolled_for_exam(student, exam):
        return "You are not authorized for this exam."

    if ExamAttempt.objects.filter(
        exam=exam,
        student=student,
        is_locked=True,
    ).exists():
        return "CBT access is locked for this exam. Contact IT Manager."

    blueprint = getattr(exam, "blueprint", None) or ensure_default_blueprint(exam)
    attempt_qs = ExamAttempt.objects.filter(exam=exam, student=student)

    if not blueprint.allow_retake and attempt_qs.exclude(status=CBTAttemptStatus.IN_PROGRESS).exists():
        return "Attempt already submitted."

    if ExamAttempt.objects.filter(
        exam=exam,
        student=student,
        status=CBTAttemptStatus.IN_PROGRESS,
    ).exists():
        return ""

    current_attempts = attempt_qs.count()
    if current_attempts >= _allowed_attempt_count(blueprint):
        return "Maximum attempt limit reached."
    return ""


def student_available_exams(student):
    now = timezone.now()
    today = timezone.localdate(now)
    close_expired_exams(now=now)
    exams = (
        Exam.objects.select_related(
            "subject",
            "academic_class",
            "session",
            "term",
            "blueprint",
        )
        .filter(status__in=[CBTExamStatus.ACTIVE, CBTExamStatus.CLOSED])
        .order_by("academic_class__code", "schedule_start", "subject__name", "title")
    )
    payload = []
    for exam in exams:
        if not exam_occurs_on_day(exam, today):
            continue
        if not _student_enrolled_for_exam(student, exam):
            continue

        attempt_qs = ExamAttempt.objects.filter(exam=exam, student=student)
        latest_attempt = attempt_qs.order_by("-updated_at").first()
        blueprint = getattr(exam, "blueprint", None) or ensure_default_blueprint(exam)
        is_done = bool(
            latest_attempt and latest_attempt.status in {CBTAttemptStatus.SUBMITTED, CBTAttemptStatus.FINALIZED}
        )
        is_in_progress = bool(latest_attempt and latest_attempt.status == CBTAttemptStatus.IN_PROGRESS)
        can_start = False
        reason = ""

        if is_done:
            status_label = "Done"
            reason = "Completed."
        elif exam.status == CBTExamStatus.CLOSED:
            status_label = "Closed"
            reason = "Exam is closed."
        else:
            reason = student_exam_authorization_reason(student=student, exam=exam, now=now)
            can_start = not bool(reason)
            if can_start:
                status_label = "In Progress" if is_in_progress else "Open"
            elif exam.schedule_start and now < exam.schedule_start:
                start_label = _as_local_datetime(exam.schedule_start)
                reason = f"Exam opens at {start_label.strftime('%H:%M')}." if start_label else "Exam is outside schedule window."
                status_label = "Not Yet"
            else:
                status_label = "Unavailable"

        payload.append(
            {
                "exam": exam,
                "can_start": can_start,
                "reason": reason,
                "attempts_used": attempt_qs.count(),
                "latest_attempt": latest_attempt,
                "max_attempts": _allowed_attempt_count(blueprint),
                "is_done": is_done,
                "is_in_progress": is_in_progress,
                "is_closed": exam.status == CBTExamStatus.CLOSED,
                "status_label": status_label,
            }
        )
    return payload


def _next_attempt_number(student, exam):
    latest = (
        ExamAttempt.objects.filter(student=student, exam=exam)
        .aggregate(max_value=Max("attempt_number"))
        .get("max_value")
    )
    return int(latest or 0) + 1


def _ordered_exam_question_rows(exam, *, shuffle_questions=False):
    question_rows = list(
        exam.exam_questions.select_related("question")
        .prefetch_related("question__options")
        .order_by("sort_order")
    )
    if shuffle_questions and len(question_rows) > 1:
        random.shuffle(question_rows)
    return question_rows


def ordered_exam_simulation_rows(exam):
    return list(
        exam.exam_simulations.select_related("simulation_wrapper").order_by("sort_order")
    )


def _build_option_order_map(question_rows, *, shuffle_options=False):
    option_order_map = {}
    for row in question_rows:
        option_ids = list(row.question.options.order_by("sort_order", "label").values_list("id", flat=True))
        if shuffle_options and len(option_ids) > 1:
            random.shuffle(option_ids)
        option_order_map[str(row.id)] = option_ids
    return option_order_map


def ensure_simulation_records_for_attempt(attempt):
    simulation_rows = ordered_exam_simulation_rows(attempt.exam)
    if not simulation_rows:
        return []
    existing_ids = set(
        attempt.simulation_attempts.values_list("exam_simulation_id", flat=True)
    )
    to_create = []
    for row in simulation_rows:
        if row.id in existing_ids:
            continue
        to_create.append(
            SimulationAttemptRecord(
                attempt=attempt,
                exam_simulation=row,
                status=CBTSimulationAttemptStatus.NOT_STARTED,
            )
        )
    if to_create:
        SimulationAttemptRecord.objects.bulk_create(to_create)
    return list(
        attempt.simulation_attempts.select_related(
            "exam_simulation",
            "exam_simulation__simulation_wrapper",
        ).order_by("exam_simulation__sort_order")
    )


def ordered_attempt_simulation_records(attempt):
    return list(
        attempt.simulation_attempts.select_related(
            "exam_simulation",
            "exam_simulation__simulation_wrapper",
            "verified_by",
            "imported_by",
        ).order_by("exam_simulation__sort_order")
    )


@transaction.atomic
def get_or_start_attempt(*, student, exam):
    if not student.has_role(ROLE_STUDENT):
        raise ValidationError("Only students can start CBT exams.")
    reason = student_exam_authorization_reason(student=student, exam=exam)
    if reason:
        raise ValidationError(reason)

    existing = (
        ExamAttempt.objects.select_related("exam", "exam__blueprint")
        .filter(
            student=student,
            exam=exam,
            status=CBTAttemptStatus.IN_PROGRESS,
        )
        .order_by("-created_at")
        .first()
    )
    if existing:
        ensure_simulation_records_for_attempt(existing)
        return existing, False

    blueprint = getattr(exam, "blueprint", None) or ensure_default_blueprint(exam)
    question_rows = _ordered_exam_question_rows(
        exam,
        shuffle_questions=blueprint.shuffle_questions,
    )
    option_order_map = _build_option_order_map(
        question_rows,
        shuffle_options=blueprint.shuffle_options,
    )
    attempt = ExamAttempt.objects.create(
        exam=exam,
        student=student,
        attempt_number=_next_attempt_number(student, exam),
        writeback_metadata={
            "question_order": [row.id for row in question_rows],
            "option_order_map": option_order_map,
        },
    )
    ExamAttemptAnswer.objects.bulk_create(
        [
            ExamAttemptAnswer(
                attempt=attempt,
                exam_question=row,
            )
            for row in question_rows
        ]
    )
    ensure_simulation_records_for_attempt(attempt)
    _queue_attempt_snapshot(attempt, event_type="ATTEMPT_STARTED")
    return attempt, True


def ordered_attempt_answers(attempt):
    order = (attempt.writeback_metadata or {}).get("question_order") or []
    answers = list(
        attempt.answers.select_related("exam_question", "exam_question__question")
        .prefetch_related("selected_options", "exam_question__question__options")
    )
    if not order:
        return sorted(answers, key=lambda row: row.exam_question.sort_order)
    index_map = {int(value): idx for idx, value in enumerate(order)}
    return sorted(
        answers,
        key=lambda row: index_map.get(row.exam_question_id, 10_000 + row.exam_question.sort_order),
    )


def option_list_for_attempt_answer(answer):
    option_order_map = (answer.attempt.writeback_metadata or {}).get("option_order_map", {})
    ordered_ids = option_order_map.get(str(answer.exam_question_id), [])
    option_qs = answer.exam_question.question.options.all()
    options = list(option_qs)
    if not ordered_ids:
        return sorted(options, key=lambda row: (row.sort_order, row.label))
    index_map = {int(value): idx for idx, value in enumerate(ordered_ids)}
    return sorted(options, key=lambda row: index_map.get(row.id, 10_000 + row.sort_order))


def _resolve_attempt_answer(attempt, exam_question_id):
    return ExamAttemptAnswer.objects.select_related("exam_question", "exam_question__question").get(
        attempt=attempt,
        exam_question_id=exam_question_id,
    )


@transaction.atomic
def save_attempt_answer(
    *,
    attempt,
    exam_question_id,
    selected_option_ids=None,
    response_text="",
    response_payload=None,
    is_flagged=None,
):
    if attempt.status != CBTAttemptStatus.IN_PROGRESS:
        raise ValidationError("Attempt is no longer editable.")
    answer = _resolve_attempt_answer(attempt, exam_question_id)
    question = answer.exam_question.question
    response_payload = response_payload or {}
    if question.question_type in OBJECTIVE_TYPES:
        valid_options = list(
            question.options.filter(id__in=(selected_option_ids or [])).values_list("id", flat=True)
        )
        answer.selected_options.set(valid_options)
        answer.response_text = ""
        answer.response_payload = {}
    else:
        answer.selected_options.clear()
        answer.response_text = (response_text or "").strip()
        answer.response_payload = response_payload
    if is_flagged is not None:
        answer.is_flagged = bool(is_flagged)
    answer.save(update_fields=["response_text", "response_payload", "is_flagged", "updated_at"])
    attempt.last_activity_at = timezone.now()
    attempt.save(update_fields=["last_activity_at", "updated_at"])
    return answer


def _answer_is_correct(answer):
    question = answer.exam_question.question
    if question.question_type not in OBJECTIVE_TYPES:
        return None
    selected = set(answer.selected_options.values_list("id", flat=True))
    correct_answer = getattr(question, "correct_answer", None)
    if not correct_answer:
        return False
    correct = set(correct_answer.correct_options.values_list("id", flat=True))
    if question.question_type == CBTQuestionType.OBJECTIVE:
        return len(selected) == 1 and selected == correct
    return selected == correct and len(correct) > 0


def _result_field_for_target(target):
    mapping = {
        CBTWritebackTarget.CA1: "ca1",
        CBTWritebackTarget.CA2: "ca2",
        CBTWritebackTarget.CA3: "ca3",
        CBTWritebackTarget.CA4: "ca4",
        CBTWritebackTarget.OBJECTIVE: "objective",
        CBTWritebackTarget.THEORY: "theory",
    }
    return mapping.get(target, "")


def _get_or_create_score_row(attempt):
    exam = attempt.exam
    sheet, _ = ResultSheet.objects.get_or_create(
        academic_class=exam.academic_class,
        subject=exam.subject,
        session=exam.session,
        term=exam.term,
        defaults={"created_by": exam.created_by},
    )
    score, _ = StudentSubjectScore.objects.get_or_create(
        result_sheet=sheet,
        student=attempt.student,
    )
    return sheet, score


def _apply_writeback(*, attempt, target, score_value, component_key):
    if target == CBTWritebackTarget.NONE:
        return {"skipped": True, "reason": "Target disabled."}
    field = _result_field_for_target(target)
    if not field:
        return {"skipped": True, "reason": "Unsupported target."}
    sheet, score = _get_or_create_score_row(attempt)
    if sheet.status == ResultSheetStatus.PUBLISHED:
        return {"skipped": True, "reason": "Result sheet already published."}
    previous = _to_decimal(getattr(score, field))
    normalized = _to_decimal(score_value)
    setattr(score, field, normalized)
    score.lock_components(field)
    score.save()
    return {
        "sheet_id": str(sheet.id),
        "field": field,
        "before": str(previous),
        "after": str(normalized),
        "component": component_key,
    }


def _mark_objective_component(attempt):
    answers = (
        attempt.answers.select_related("exam_question", "exam_question__question")
        .prefetch_related("selected_options", "exam_question__question__correct_answer__correct_options")
    )
    objective_raw = Decimal("0.00")
    objective_max = Decimal("0.00")
    theory_max = Decimal("0.00")
    for answer in answers:
        question = answer.exam_question.question
        marks = _to_decimal(answer.exam_question.marks)
        if question.question_type in OBJECTIVE_TYPES:
            objective_max += marks
            is_correct = _answer_is_correct(answer)
            answer.is_correct = bool(is_correct)
            answer.auto_score = marks if is_correct else Decimal("0.00")
            objective_raw += answer.auto_score
            answer.save(update_fields=["is_correct", "auto_score", "updated_at"])
        elif question.question_type in STRUCTURED_AUTO_TYPES:
            objective_max += marks
            correct_answer = getattr(question, "correct_answer", None)
            expected = re.sub(r"\s+", " ", (getattr(correct_answer, "note", "") or "")).strip().lower()
            observed = re.sub(
                r"\s+",
                " ",
                ((answer.response_text or "") or (answer.response_payload or {}).get("raw", "")),
            ).strip().lower()
            is_correct = bool(expected and observed and observed == expected)
            answer.is_correct = is_correct
            answer.auto_score = marks if is_correct else Decimal("0.00")
            objective_raw += answer.auto_score
            answer.save(update_fields=["is_correct", "auto_score", "updated_at"])
        else:
            theory_max += marks
            answer.is_correct = None
            answer.auto_score = Decimal("0.00")
            answer.save(update_fields=["is_correct", "auto_score", "updated_at"])
    blueprint = getattr(attempt.exam, "blueprint", None) or ensure_default_blueprint(attempt.exam)
    objective_target_override = _section_target_max(
        blueprint,
        section_key="objective_target_max",
    )
    attempt.objective_raw_score = objective_raw.quantize(DECIMAL_2)
    attempt.objective_max_score = objective_max.quantize(DECIMAL_2)
    attempt.objective_score = _normalize_score(
        objective_raw,
        objective_max,
        blueprint.objective_writeback_target,
        target_max_override=objective_target_override,
    )
    attempt.theory_max_score = theory_max.quantize(DECIMAL_2)
    return blueprint


def has_pending_required_simulation(attempt):
    if not attempt.exam.exam_simulations.exists():
        return False
    ensure_simulation_records_for_attempt(attempt)
    return attempt.simulation_attempts.filter(
        exam_simulation__is_required=True,
    ).exclude(
        status__in=SIMULATION_STUDENT_COMPLETED_STATUSES,
    ).exists()


@transaction.atomic
def submit_attempt(*, attempt):
    if attempt.status == CBTAttemptStatus.FINALIZED:
        raise ValidationError("Attempt is already finalized.")
    if has_pending_required_simulation(attempt):
        raise ValidationError("Complete all required simulation tasks before submitting.")
    blueprint = _mark_objective_component(attempt)
    has_theory_questions = attempt.answers.filter(
        exam_question__question__question_type__in=THEORY_TYPES
    ).exists()
    attempt.auto_marking_completed = True
    if not has_theory_questions or not blueprint.theory_enabled:
        attempt.theory_marking_completed = True
        attempt.theory_raw_score = Decimal("0.00")
        attempt.theory_score = Decimal("0.00")
        if not has_theory_questions:
            attempt.theory_max_score = Decimal("0.00")
    else:
        attempt.theory_marking_completed = False
    if not attempt.submitted_at:
        attempt.submitted_at = timezone.now()
    attempt.status = CBTAttemptStatus.SUBMITTED
    if is_free_test_exam(attempt.exam):
        objective_writeback = {
            "skipped": True,
            "reason": "Free test does not write into graded records.",
        }
    else:
        objective_writeback = _apply_writeback(
            attempt=attempt,
            target=blueprint.objective_writeback_target,
            score_value=attempt.objective_score,
            component_key="objective",
        )
    attempt.total_score = (attempt.objective_score + attempt.theory_score).quantize(DECIMAL_2)
    metadata = attempt.writeback_metadata or {}
    metadata["objective_writeback"] = objective_writeback
    attempt.writeback_metadata = metadata
    attempt.writeback_completed = bool(attempt.theory_marking_completed)
    attempt.save(
        update_fields=[
            "status",
            "submitted_at",
            "objective_raw_score",
            "objective_max_score",
            "objective_score",
            "theory_raw_score",
            "theory_max_score",
            "theory_score",
            "total_score",
            "auto_marking_completed",
            "theory_marking_completed",
            "writeback_completed",
            "writeback_metadata",
            "updated_at",
        ]
    )
    _queue_attempt_snapshot(attempt, event_type="ATTEMPT_SUBMITTED")
    return attempt


def _teacher_can_mark_attempt(actor, attempt):
    if can_manage_all_cbt(actor):
        return True
    return TeacherSubjectAssignment.objects.filter(
        teacher=actor,
        subject=attempt.exam.subject,
        academic_class=attempt.exam.academic_class,
        session=attempt.exam.session,
        term=attempt.exam.term,
        is_active=True,
    ).exists()


def teacher_can_manage_attempt(actor, attempt):
    return _teacher_can_mark_attempt(actor, attempt)


@transaction.atomic
def apply_theory_scores(*, attempt, actor, score_payload):
    if not _teacher_can_mark_attempt(actor, attempt):
        raise ValidationError("You are not authorized to mark this attempt.")
    if attempt.status not in {CBTAttemptStatus.SUBMITTED, CBTAttemptStatus.FINALIZED}:
        raise ValidationError("Attempt must be submitted before theory marking.")
    blueprint = getattr(attempt.exam, "blueprint", None) or ensure_default_blueprint(attempt.exam)
    if not blueprint.theory_enabled:
        raise ValidationError("Theory marking is disabled for this exam.")

    theory_answers = list(
        attempt.answers.select_related("exam_question", "exam_question__question").filter(
            exam_question__question__question_type__in=THEORY_TYPES
        )
    )
    if not theory_answers:
        attempt.theory_marking_completed = True
        attempt.theory_raw_score = Decimal("0.00")
        attempt.theory_score = Decimal("0.00")
        attempt.total_score = (attempt.objective_score + attempt.theory_score).quantize(DECIMAL_2)
        attempt.writeback_completed = True
        attempt.save(
            update_fields=[
                "theory_marking_completed",
                "theory_raw_score",
                "theory_score",
                "total_score",
                "writeback_completed",
                "updated_at",
            ]
        )
        _queue_attempt_snapshot(attempt, event_type="ATTEMPT_THEORY_MARKED")
        return attempt

    now = timezone.now()
    total_raw = Decimal("0.00")
    total_max = Decimal("0.00")
    all_marked = True
    for answer in theory_answers:
        marks = _to_decimal(answer.exam_question.marks)
        total_max += marks
        posted_value = score_payload.get(str(answer.id))
        if posted_value in (None, ""):
            all_marked = False
            continue
        parsed = _to_decimal(posted_value)
        if parsed < Decimal("0.00") or parsed > marks:
            raise ValidationError(
                f"Score for question {answer.exam_question.sort_order} must be between 0 and {marks}."
            )
        answer.teacher_score = parsed
        answer.teacher_marked_by = actor
        answer.teacher_marked_at = now
        answer.save(update_fields=["teacher_score", "teacher_marked_by", "teacher_marked_at", "updated_at"])

    for answer in theory_answers:
        if answer.teacher_score is None:
            all_marked = False
            continue
        total_raw += _to_decimal(answer.teacher_score)

    attempt.theory_raw_score = total_raw.quantize(DECIMAL_2)
    attempt.theory_max_score = total_max.quantize(DECIMAL_2)
    attempt.theory_marking_completed = all_marked
    if all_marked:
        theory_target_override = _section_target_max(
            blueprint,
            section_key="theory_target_max",
        )
        attempt.theory_score = _normalize_score(
            attempt.theory_raw_score,
            attempt.theory_max_score,
            blueprint.theory_writeback_target,
            target_max_override=theory_target_override,
        )
        if is_free_test_exam(attempt.exam):
            writeback = {
                "skipped": True,
                "reason": "Free test does not write into graded records.",
            }
        else:
            writeback_score = attempt.theory_score
            writeback_component = "theory"
            if (
                blueprint.theory_writeback_target == blueprint.objective_writeback_target
                and blueprint.theory_writeback_target
                in {
                    CBTWritebackTarget.CA1,
                    CBTWritebackTarget.CA2,
                    CBTWritebackTarget.CA3,
                    CBTWritebackTarget.CA4,
                }
            ):
                writeback_score = (attempt.objective_score + attempt.theory_score).quantize(DECIMAL_2)
                writeback_component = "objective_theory_combined"
            writeback = _apply_writeback(
                attempt=attempt,
                target=blueprint.theory_writeback_target,
                score_value=writeback_score,
                component_key=writeback_component,
            )
        metadata = attempt.writeback_metadata or {}
        metadata["theory_writeback"] = writeback
        attempt.writeback_metadata = metadata
        attempt.writeback_completed = True
    attempt.total_score = (attempt.objective_score + attempt.theory_score).quantize(DECIMAL_2)
    attempt.save(
        update_fields=[
            "theory_raw_score",
            "theory_max_score",
            "theory_score",
            "theory_marking_completed",
            "total_score",
            "writeback_completed",
            "writeback_metadata",
            "updated_at",
        ]
    )
    _queue_attempt_snapshot(attempt, event_type="ATTEMPT_THEORY_MARKED")
    return attempt


def attempt_deadline(attempt):
    blueprint = getattr(attempt.exam, "blueprint", None) or ensure_default_blueprint(attempt.exam)
    duration_minutes = blueprint.duration_minutes + int(attempt.extra_time_minutes or 0)
    deadline = attempt.started_at + timezone.timedelta(minutes=duration_minutes)
    if attempt.exam.schedule_end and int(attempt.extra_time_minutes or 0) <= 0:
        deadline = min(deadline, attempt.exam.schedule_end)
    return deadline


def finalize_attempt(attempt):
    if attempt.status == CBTAttemptStatus.FINALIZED:
        return attempt
    blueprint = getattr(attempt.exam, "blueprint", None) or ensure_default_blueprint(attempt.exam)
    if attempt.status == CBTAttemptStatus.IN_PROGRESS:
        submit_attempt(attempt=attempt)
    if not blueprint.allow_retake:
        attempt.status = CBTAttemptStatus.FINALIZED
        attempt.finalized_at = timezone.now()
        attempt.save(update_fields=["status", "finalized_at", "updated_at"])
        _queue_attempt_snapshot(attempt, event_type="ATTEMPT_FINALIZED")
    return attempt


def finalize_cbt_attempts_on_logout(student):
    attempts = (
        ExamAttempt.objects.select_related("exam", "exam__blueprint")
        .filter(
            student=student,
            status__in=[CBTAttemptStatus.IN_PROGRESS, CBTAttemptStatus.SUBMITTED],
        )
        .filter(
            Q(exam__blueprint__finalize_on_logout=True) | Q(exam__blueprint__isnull=True)
        )
    )
    finalized_count = 0
    for attempt in attempts:
        previous_status = attempt.status
        finalize_attempt(attempt)
        attempt.refresh_from_db(fields=["status"])
        if attempt.status == CBTAttemptStatus.FINALIZED and previous_status != CBTAttemptStatus.FINALIZED:
            finalized_count += 1
    return finalized_count


def theory_marking_queryset_for_user(user):
    attempts = (
        ExamAttempt.objects.select_related(
            "exam",
            "exam__subject",
            "exam__academic_class",
            "exam__session",
            "exam__term",
            "student",
        )
        .filter(
            status__in=[CBTAttemptStatus.SUBMITTED, CBTAttemptStatus.FINALIZED],
            exam__blueprint__theory_enabled=True,
        )
        .filter(
            answers__exam_question__question__question_type__in=THEORY_TYPES
        )
        .distinct()
        .order_by("-updated_at")
    )
    if can_manage_all_cbt(user):
        return attempts
    assignments = list(
        TeacherSubjectAssignment.objects.filter(
        teacher=user,
        is_active=True,
    ).values("subject_id", "academic_class_id", "session_id", "term_id")
    )
    if not assignments:
        return attempts.none()
    combined_q = Q()
    for assignment in assignments:
        combined_q |= Q(
            exam__subject_id=assignment["subject_id"],
            exam__academic_class_id=assignment["academic_class_id"],
            exam__session_id=assignment["session_id"],
            exam__term_id=assignment["term_id"],
        )
    return attempts.filter(combined_q)


def simulation_marking_queryset_for_user(user):
    records = (
        SimulationAttemptRecord.objects.select_related(
            "attempt",
            "attempt__student",
            "attempt__exam",
            "attempt__exam__subject",
            "attempt__exam__academic_class",
            "exam_simulation",
            "exam_simulation__simulation_wrapper",
            "verified_by",
            "imported_by",
        )
        .exclude(status=CBTSimulationAttemptStatus.NOT_STARTED)
        .order_by("-updated_at")
    )
    if can_manage_all_cbt(user):
        return records
    assignments = list(
        TeacherSubjectAssignment.objects.filter(
            teacher=user,
            is_active=True,
        ).values("subject_id", "academic_class_id", "session_id", "term_id")
    )
    if not assignments:
        return records.none()
    combined_q = Q()
    for assignment in assignments:
        combined_q |= Q(
            attempt__exam__subject_id=assignment["subject_id"],
            attempt__exam__academic_class_id=assignment["academic_class_id"],
            attempt__exam__session_id=assignment["session_id"],
            attempt__exam__term_id=assignment["term_id"],
        )
    return records.filter(combined_q)


def _resolve_simulation_record(*, attempt, exam_simulation):
    if exam_simulation.exam_id != attempt.exam_id:
        raise ValidationError("Simulation does not belong to this attempt exam.")
    record, _ = SimulationAttemptRecord.objects.get_or_create(
        attempt=attempt,
        exam_simulation=exam_simulation,
        defaults={"status": CBTSimulationAttemptStatus.NOT_STARTED},
    )
    return record


def _validate_simulation_score(value, max_score):
    score = _to_decimal(value)
    max_value = _to_decimal(max_score)
    if score < Decimal("0.00") or score > max_value:
        raise ValidationError(f"Simulation score must be between 0 and {max_value}.")
    return score


def _payload_value(payload, path):
    current = payload
    for key in path:
        if not isinstance(current, dict):
            return None
        current = current.get(key)
    return current


def _find_xapi_statement(payload):
    if not isinstance(payload, dict):
        return {}
    statement_paths = (
        ("statement",),
        ("xapi", "statement"),
        ("xAPI", "statement"),
        ("xapi_statement",),
        ("data", "statement"),
        ("detail", "statement"),
    )
    for path in statement_paths:
        statement = _payload_value(payload, path)
        if isinstance(statement, dict):
            return statement
    return {}


def _extract_score_from_statement(statement, max_score):
    if not isinstance(statement, dict):
        return None
    result = statement.get("result")
    if not isinstance(result, dict):
        return None
    score = result.get("score")
    if not isinstance(score, dict):
        return None

    max_value = _to_decimal(max_score)
    scaled = score.get("scaled")
    if scaled not in (None, ""):
        scaled_decimal = _to_decimal(scaled)
        if scaled_decimal < Decimal("0.00") or scaled_decimal > Decimal("1.00"):
            raise ValidationError("xAPI scaled score must be between 0 and 1.")
        return (scaled_decimal * max_value).quantize(DECIMAL_2)

    raw = score.get("raw")
    raw_max = score.get("max")
    if raw not in (None, "") and raw_max not in (None, ""):
        raw_decimal = _to_decimal(raw)
        raw_max_decimal = _to_decimal(raw_max)
        if raw_max_decimal <= 0:
            return Decimal("0.00")
        return ((raw_decimal / raw_max_decimal) * max_value).quantize(DECIMAL_2)

    if raw not in (None, ""):
        raw_decimal = _to_decimal(raw)
        if raw_decimal <= Decimal("1.00"):
            return (raw_decimal * max_value).quantize(DECIMAL_2)
        return raw_decimal.quantize(DECIMAL_2)

    return None


def resolve_auto_simulation_score_from_payload(payload, max_score):
    if not isinstance(payload, dict):
        raise ValidationError("Invalid simulation callback payload.")

    if payload.get("score") not in (None, ""):
        score = _validate_simulation_score(payload.get("score"), max_score)
        return score, "direct_score_field"

    statement = _find_xapi_statement(payload)
    statement_score = _extract_score_from_statement(statement, max_score)
    if statement_score is not None:
        score = _validate_simulation_score(statement_score, max_score)
        return score, "xapi_statement_result"

    candidate_paths = (
        ("result", "score", "scaled"),
        ("result", "score", "raw"),
        ("score", "scaled"),
        ("score", "raw"),
    )
    for path in candidate_paths:
        value = _payload_value(payload, path)
        if value in (None, ""):
            continue
        if path[-1] == "scaled":
            scaled_score = (_to_decimal(value) * _to_decimal(max_score)).quantize(DECIMAL_2)
            score = _validate_simulation_score(scaled_score, max_score)
            return score, "scaled_payload_value"
        score = _validate_simulation_score(value, max_score)
        return score, "raw_payload_value"

    raise ValidationError(
        "No score found in callback payload. Ensure simulation sends score or xAPI statement."
    )


@transaction.atomic
def capture_auto_simulation_score(*, attempt, exam_simulation, payload):
    wrapper = exam_simulation.simulation_wrapper
    if wrapper.score_mode != CBTSimulationScoreMode.AUTO:
        raise ValidationError("Simulation is not configured for AUTO scoring.")
    score, extraction_method = resolve_auto_simulation_score_from_payload(
        payload,
        exam_simulation.effective_max_score,
    )
    record = _resolve_simulation_record(attempt=attempt, exam_simulation=exam_simulation)
    callback_payload = deepcopy(payload) if isinstance(payload, dict) else {"payload": payload}
    callback_payload["_score_extraction"] = {
        "method": extraction_method,
        "final_score": str(score),
    }
    record.raw_score = score
    record.final_score = score
    record.callback_payload = callback_payload
    record.status = CBTSimulationAttemptStatus.AUTO_CAPTURED
    record.verify_comment = ""
    record.save(
        update_fields=[
            "raw_score",
            "final_score",
            "callback_payload",
            "status",
            "verify_comment",
            "updated_at",
        ]
    )
    _queue_simulation_snapshot(record, event_type="SIMULATION_AUTO_CAPTURED")
    return record


@transaction.atomic
def submit_verify_simulation_evidence(
    *,
    attempt,
    exam_simulation,
    evidence_file=None,
    evidence_note="",
    payload=None,
):
    wrapper = exam_simulation.simulation_wrapper
    if wrapper.score_mode != CBTSimulationScoreMode.VERIFY:
        raise ValidationError("Simulation is not configured for VERIFY mode.")
    record = _resolve_simulation_record(attempt=attempt, exam_simulation=exam_simulation)
    if wrapper.evidence_required and not evidence_file and not record.evidence_file:
        raise ValidationError("Evidence upload is required for this simulation.")
    if evidence_file is not None:
        record.evidence_file = evidence_file
    record.evidence_note = (evidence_note or "").strip()
    record.callback_payload = payload if isinstance(payload, dict) else {"payload": payload or {}}
    record.status = CBTSimulationAttemptStatus.VERIFY_PENDING
    record.save(
        update_fields=[
            "evidence_file",
            "evidence_note",
            "callback_payload",
            "status",
            "updated_at",
        ]
    )
    _queue_simulation_snapshot(record, event_type="SIMULATION_VERIFY_PENDING")
    return record


@transaction.atomic
def submit_rubric_simulation_start(
    *,
    attempt,
    exam_simulation,
    evidence_file=None,
    evidence_note="",
    payload=None,
):
    wrapper = exam_simulation.simulation_wrapper
    if wrapper.score_mode != CBTSimulationScoreMode.RUBRIC:
        raise ValidationError("Simulation is not configured for RUBRIC mode.")
    record = _resolve_simulation_record(attempt=attempt, exam_simulation=exam_simulation)
    if wrapper.evidence_required and not evidence_file and not record.evidence_file:
        raise ValidationError("Evidence upload is required for this simulation.")
    if evidence_file is not None:
        record.evidence_file = evidence_file
    record.evidence_note = (evidence_note or "").strip()
    record.callback_payload = payload if isinstance(payload, dict) else {"payload": payload or {}}
    record.status = CBTSimulationAttemptStatus.RUBRIC_PENDING
    record.save(
        update_fields=[
            "evidence_file",
            "evidence_note",
            "callback_payload",
            "status",
            "updated_at",
        ]
    )
    _queue_simulation_snapshot(record, event_type="SIMULATION_RUBRIC_PENDING")
    return record


@transaction.atomic
def teacher_verify_simulation_score(*, record, actor, verified_score, comment=""):
    if not _teacher_can_mark_attempt(actor, record.attempt):
        raise ValidationError("You are not authorized to verify this simulation.")
    if record.exam_simulation.simulation_wrapper.score_mode != CBTSimulationScoreMode.VERIFY:
        raise ValidationError("This simulation is not in VERIFY mode.")
    score = _validate_simulation_score(
        verified_score,
        record.exam_simulation.effective_max_score,
    )
    record.raw_score = score
    record.final_score = score
    record.verify_comment = (comment or "").strip()
    record.verified_by = actor
    record.verified_at = timezone.now()
    record.status = CBTSimulationAttemptStatus.VERIFIED
    record.save(
        update_fields=[
            "raw_score",
            "final_score",
            "verify_comment",
            "verified_by",
            "verified_at",
            "status",
            "updated_at",
        ]
    )
    _queue_simulation_snapshot(record, event_type="SIMULATION_VERIFIED")
    return record


@transaction.atomic
def teacher_score_rubric_simulation(*, record, actor, rubric_scores, comment=""):
    if not _teacher_can_mark_attempt(actor, record.attempt):
        raise ValidationError("You are not authorized to score this simulation.")
    if record.exam_simulation.simulation_wrapper.score_mode != CBTSimulationScoreMode.RUBRIC:
        raise ValidationError("This simulation is not in RUBRIC mode.")
    if not isinstance(rubric_scores, dict) or not rubric_scores:
        raise ValidationError("Provide rubric criteria scores.")
    parsed_scores = {}
    for key, value in rubric_scores.items():
        score = _to_decimal(value)
        if score < Decimal("0.00") or score > Decimal("100.00"):
            raise ValidationError("Rubric criteria scores must be between 0 and 100.")
        parsed_scores[key] = score
    average_percent = (
        sum(parsed_scores.values()) / Decimal(len(parsed_scores))
    ).quantize(DECIMAL_2)
    max_score = _to_decimal(record.exam_simulation.effective_max_score)
    final_score = ((average_percent / Decimal("100.00")) * max_score).quantize(DECIMAL_2)
    record.rubric_breakdown = {
        "criteria_scores": {key: str(value) for key, value in parsed_scores.items()},
        "average_percent": str(average_percent),
        "max_score": str(max_score),
    }
    record.raw_score = average_percent
    record.final_score = final_score
    record.verify_comment = (comment or "").strip()
    record.verified_by = actor
    record.verified_at = timezone.now()
    record.status = CBTSimulationAttemptStatus.RUBRIC_SCORED
    record.save(
        update_fields=[
            "rubric_breakdown",
            "raw_score",
            "final_score",
            "verify_comment",
            "verified_by",
            "verified_at",
            "status",
            "updated_at",
        ]
    )
    _queue_simulation_snapshot(record, event_type="SIMULATION_RUBRIC_SCORED")
    return record


@transaction.atomic
def import_simulation_score_to_results(
    *,
    record,
    actor,
    writeback_target="",
    manual_score=None,
):
    if not _teacher_can_mark_attempt(actor, record.attempt):
        raise ValidationError("You are not authorized to import this simulation score.")
    target = writeback_target or record.exam_simulation.writeback_target
    if target == CBTWritebackTarget.NONE:
        raise ValidationError("Choose a valid writeback target before importing.")

    source_score = record.final_score
    if manual_score not in (None, ""):
        source_score = _to_decimal(manual_score)
    if source_score is None:
        raise ValidationError("No simulation score available to import.")

    simulation_max = _to_decimal(record.exam_simulation.effective_max_score)
    source_score = _validate_simulation_score(source_score, simulation_max)
    normalized = _normalize_score(source_score, simulation_max, target)
    writeback = _apply_writeback(
        attempt=record.attempt,
        target=target,
        score_value=normalized,
        component_key="simulation",
    )
    if writeback.get("skipped"):
        raise ValidationError(writeback["reason"])

    record.imported_target = target
    record.imported_score = normalized
    record.imported_by = actor
    record.imported_at = timezone.now()
    record.status = CBTSimulationAttemptStatus.IMPORTED
    record.save(
        update_fields=[
            "imported_target",
            "imported_score",
            "imported_by",
            "imported_at",
            "status",
            "updated_at",
        ]
    )

    metadata = record.attempt.writeback_metadata or {}
    imports = metadata.get("simulation_imports") or []
    imports.append(
        {
            "record_id": str(record.id),
            "exam_simulation_id": str(record.exam_simulation_id),
            "target": target,
            "source_score": str(source_score),
            "imported_score": str(normalized),
            "imported_by": actor.username,
            "imported_at": record.imported_at.isoformat(),
            "writeback": writeback,
        }
    )
    metadata["simulation_imports"] = imports
    record.attempt.writeback_metadata = metadata
    record.attempt.save(update_fields=["writeback_metadata", "updated_at"])
    _queue_simulation_snapshot(record, event_type="SIMULATION_IMPORTED")
    _queue_attempt_snapshot(record.attempt, event_type="ATTEMPT_SIMULATION_WRITEBACK")
    return writeback


LOCKDOWN_EVENT_TYPES = {
    "VISIBILITY_CHANGE",
    "FOCUS_LOSS",
    "TAB_SWITCH",
    "MULTIPLE_TAB",
    "COPY_ATTEMPT",
    "PASTE_ATTEMPT",
    "FULLSCREEN_EXIT",
    "CAMERA_BLOCKED",
    "INACTIVITY_TIMEOUT",
}


def lockdown_enabled():
    return settings.FEATURE_FLAGS.get("LOCKDOWN_ENABLED", False)


@transaction.atomic
def register_lockdown_heartbeat(*, attempt, tab_token, request, client_state=None):
    if not lockdown_enabled():
        return {"ok": True, "lockdown_enabled": False}
    if attempt.is_locked:
        return {"locked": True}
    if attempt.status != CBTAttemptStatus.IN_PROGRESS:
        return {"ok": True, "attempt_closed": True}

    token = (tab_token or "").strip()[:120]
    if not token:
        raise ValidationError("Missing tab token.")

    if attempt.active_tab_token and attempt.active_tab_token != token:
        record_lockdown_violation(
            attempt=attempt,
            event_type="MULTIPLE_TAB",
            request=request,
            details={"client_state": client_state or {}},
        )
        return {"locked": True}

    now = timezone.now()
    attempt.active_tab_token = token
    attempt.last_heartbeat_at = now
    attempt.last_activity_at = now
    attempt.save(
        update_fields=[
            "active_tab_token",
            "last_heartbeat_at",
            "last_activity_at",
            "updated_at",
        ]
    )
    return {"ok": True}


@transaction.atomic
def record_lockdown_violation(*, attempt, event_type, request, details=None):
    if not lockdown_enabled():
        return attempt
    if event_type not in LOCKDOWN_EVENT_TYPES:
        raise ValidationError("Unsupported lockdown event type.")
    if attempt.is_locked:
        return attempt

    details = details or {}
    metadata = {
        "event_type": event_type,
        "attempt_id": str(attempt.id),
        "exam_id": str(attempt.exam_id),
        "student_id": str(attempt.student_id),
        "device": request.META.get("HTTP_USER_AGENT", ""),
        "details": details,
    }
    log_lockdown_violation(
        actor=attempt.student,
        request=request,
        metadata=metadata,
    )

    should_lock_for_review = event_type in {"FULLSCREEN_EXIT", "CAMERA_BLOCKED"}
    if attempt.status == CBTAttemptStatus.IN_PROGRESS and not should_lock_for_review:
        submit_attempt(attempt=attempt)

    attempt.is_locked = True
    attempt.lock_reason = event_type
    attempt.locked_at = timezone.now()
    attempt.allow_resume_by_it = bool(should_lock_for_review)
    attempt.active_tab_token = ""
    update_fields = [
        "is_locked",
        "lock_reason",
        "locked_at",
        "allow_resume_by_it",
        "active_tab_token",
        "updated_at",
    ]
    if should_lock_for_review:
        if attempt.status == CBTAttemptStatus.FINALIZED and not attempt.finalized_at:
            attempt.finalized_at = timezone.now()
            update_fields.append("finalized_at")
    else:
        attempt.status = CBTAttemptStatus.FINALIZED
        if not attempt.finalized_at:
            attempt.finalized_at = timezone.now()
        update_fields.extend(["status", "finalized_at"])
    attempt.save(update_fields=update_fields)
    _queue_attempt_snapshot(attempt, event_type="ATTEMPT_LOCKDOWN_LOCKED")
    return attempt


@transaction.atomic
def it_unlock_attempt(*, attempt, actor, allow_resume=False, extra_time_minutes=0, request=None):
    if not actor.has_role(ROLE_IT_MANAGER):
        raise ValidationError("Only IT Manager can unlock CBT attempts.")

    extra_time_minutes = int(extra_time_minutes or 0)
    if extra_time_minutes < 0:
        raise ValidationError("Extra time cannot be negative.")

    attempt.is_locked = False
    attempt.lock_reason = ""
    attempt.locked_at = None
    attempt.active_tab_token = ""
    if extra_time_minutes:
        attempt.extra_time_minutes = int(attempt.extra_time_minutes or 0) + extra_time_minutes

    if allow_resume:
        attempt.allow_resume_by_it = True
        attempt.status = CBTAttemptStatus.IN_PROGRESS
        attempt.finalized_at = None
        attempt.submitted_at = None
    else:
        attempt.allow_resume_by_it = False
        if attempt.status == CBTAttemptStatus.IN_PROGRESS:
            submit_attempt(attempt=attempt)
        attempt.status = CBTAttemptStatus.FINALIZED
        if not attempt.finalized_at:
            attempt.finalized_at = timezone.now()

    attempt.save(
        update_fields=[
            "is_locked",
            "lock_reason",
            "locked_at",
            "active_tab_token",
            "extra_time_minutes",
            "allow_resume_by_it",
            "status",
            "finalized_at",
            "submitted_at",
            "updated_at",
        ]
    )
    _queue_attempt_snapshot(attempt, event_type="ATTEMPT_UNLOCKED_BY_IT")
    if request is not None:
        log_event(
            category=AuditCategory.LOCKDOWN,
            event_type="LOCKDOWN_UNLOCK",
            status=AuditStatus.SUCCESS,
            actor=actor,
            request=request,
            metadata={
                "attempt_id": str(attempt.id),
                "allow_resume": bool(allow_resume),
                "extra_time_minutes_added": extra_time_minutes,
            },
        )
    return attempt
